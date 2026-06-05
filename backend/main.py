import gc
import json
import os
import re
import time
import uuid
import shutil
import base64
import threading
import contextvars
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any, Callable, Optional, List

from dotenv import load_dotenv

load_dotenv()

from fastapi import FastAPI, File, UploadFile, HTTPException, Header, Query, Request, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, Response
from PIL import Image, ImageChops, ImageFilter, ImageOps, ImageEnhance, ImageStat
import pytesseract
from pytesseract import Output
import fitz  # PyMuPDF
import batch_jobs as batches
import md_tools
import table_ocr
import vision_llm
import scan_pages
import tesseract_layout
import paddle_ocr
from fpdf import FPDF
import uvicorn
from markitdown import MarkItDown
import google.generativeai as genai
import pdfplumber
from docx import Document
from docx.shared import Pt

# ── Config ────────────────────────────────────────────────────────────────────
# Vacío = sin clave (recomendado en local). Pon OCR_API_KEY=sci-ocr-2024 si quieres proteger el servidor.
API_KEY              = (os.getenv("OCR_API_KEY") or "").strip()
GEMINI_API_KEY       = os.getenv("GEMINI_API_KEY", "")
UPLOAD_DIR           = (Path("uploads")).resolve()
OUTPUT_DIR           = (Path("outputs")).resolve()
MAX_UPLOAD_MB        = max(1, int(os.getenv("MAX_UPLOAD_MB", "300")))
MAX_UPLOAD_BYTES     = MAX_UPLOAD_MB * 1024 * 1024
MARKITDOWN_MIN_CHARS = 100
SOFT_PAGE_WARNING    = 30
SOFT_SIZE_WARNING_MB = 20
OCR_RENDER_DPI       = int(os.getenv("OCR_RENDER_DPI", "320"))
OCR_UPSCALE_MIN_EDGE = int(os.getenv("OCR_UPSCALE_MIN_EDGE", "1800"))
OCR_CHUNK_PAGES      = max(2, min(24, int(os.getenv("OCR_CHUNK_PAGES", "8"))))
OCR_MIN_PAGE_CHARS   = 24
OCR_HEAVY_ASYNC_MB   = max(1.0, float(os.getenv("OCR_HEAVY_ASYNC_MB", "15")))
OCR_HEAVY_ASYNC_PAGES = max(5, int(os.getenv("OCR_HEAVY_ASYNC_PAGES", "35")))
OCR_MAX_CONCURRENT_JOBS = max(1, min(3, int(os.getenv("OCR_MAX_CONCURRENT_JOBS", "1"))))
OCR_JOB_TTL_HOURS    = max(1, int(os.getenv("OCR_JOB_TTL_HOURS", "48")))
OCR_ASYNC_ALWAYS     = os.getenv("OCR_ASYNC_ALWAYS", "0").strip().lower() in ("1", "true", "yes")
OCR_BALANCED_GOOD_SCORE = float(os.getenv("OCR_BALANCED_GOOD_SCORE", "36"))
OCR_TESSERACT_WORKERS = max(1, min(8, int(os.getenv("OCR_TESSERACT_WORKERS", "6"))))
OCR_GEMINI_WORKERS    = max(1, min(4, int(os.getenv("OCR_GEMINI_WORKERS", "2"))))
GEMINI_MODEL          = (os.getenv("GEMINI_MODEL") or "gemini-2.0-flash").strip()
OCR_SUPPLEMENTARY_PASS = os.getenv("OCR_SUPPLEMENTARY_PASS", "0").strip().lower() in ("1", "true", "yes")
OCR_SUPPLEMENTARY_MIN_SCORE = float(os.getenv("OCR_SUPPLEMENTARY_MIN_SCORE", "42"))
OCR_HYBRID_PARALLEL   = os.getenv("OCR_HYBRID_PARALLEL", "1").strip().lower() in ("1", "true", "yes")
OCR_LARGE_PDF_PAGES   = max(10, int(os.getenv("OCR_LARGE_PDF_PAGES", "50")))
OCR_LARGE_PDF_GEMINI_MAX = max(5, int(os.getenv("OCR_LARGE_PDF_GEMINI_MAX", "45")))
OCR_LARGE_DPI         = int(os.getenv("OCR_LARGE_DPI", "220"))
OCR_LARGE_UPSCALE     = int(os.getenv("OCR_LARGE_UPSCALE", "1400"))
OCR_BULK_FIX_ROTATION = os.getenv("OCR_BULK_FIX_ROTATION", "1").strip().lower() in ("1", "true", "yes")
OCR_FIX_ROTATION = os.getenv("OCR_FIX_ROTATION", "0").strip().lower() in ("1", "true", "yes")
OCR_TESSERACT_ONLY = os.getenv("OCR_TESSERACT_ONLY", "1").strip().lower() in ("1", "true", "yes")
OCR_USE_LAYOUT = os.getenv("OCR_USE_LAYOUT", "1").strip().lower() in ("1", "true", "yes")
OCR_VISION_ENABLED    = os.getenv("OCR_VISION_ENABLED", "0").strip().lower() not in ("0", "false", "no")
OCR_VISION_TABLES_ONLY = os.getenv("OCR_VISION_TABLES_ONLY", "0").strip().lower() in ("1", "true", "yes")
OCR_VISION_DOCS_ONLY  = os.getenv("OCR_VISION_DOCS_ONLY", "0").strip().lower() in ("1", "true", "yes")
OCR_VISION_LEGAL_ONLY = os.getenv("OCR_VISION_LEGAL_ONLY", "0").strip().lower() in ("1", "true", "yes")
OCR_PADDLE_ENABLED = os.getenv("OCR_PADDLE_ENABLED", "0").strip().lower() in ("1", "true", "yes")
OCR_PADDLE_LAYOUT = os.getenv("OCR_PADDLE_LAYOUT", "1").strip().lower() in ("1", "true", "yes")
OCR_PADDLE_LANG = (os.getenv("OCR_PADDLE_LANG") or "es").strip().lower()
ANNOTATIONS_SECTION   = "## Firmas, sellos y anotaciones"
JOBS_DIR = (Path("jobs")).resolve()
SOURCES_DIR = (UPLOAD_DIR / "sources").resolve()
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)
JOBS_DIR.mkdir(exist_ok=True)
SOURCES_DIR.mkdir(exist_ok=True)
HISTORY_PATH = JOBS_DIR / "history.json"
HISTORY_MAX_ITEMS = 30

_queue_lock = threading.Lock()
_job_run_semaphore = threading.Semaphore(OCR_MAX_CONCURRENT_JOBS)
_pending_job_ids: list[str] = []
_running_job_id: Optional[str] = None
_job_controls: dict[str, "JobControl"] = {}
_workers_started: set[str] = set()

_progress_callback: contextvars.ContextVar[Optional[Callable[[int, int, str], None]]] = contextvars.ContextVar(
    "ocr_progress_callback", default=None
)
_job_control_ctx: contextvars.ContextVar[Optional["JobControl"]] = contextvars.ContextVar(
    "ocr_job_control", default=None
)
_ocr_profile_ctx: contextvars.ContextVar[dict[str, Any]] = contextvars.ContextVar(
    "ocr_profile", default={}
)
_vision_provider_ctx: contextvars.ContextVar[Optional[str]] = contextvars.ContextVar(
    "vision_provider", default=None
)
_rotation_fixes_ctx: contextvars.ContextVar[list[int]] = contextvars.ContextVar(
    "rotation_fixes", default=[]
)
_last_progress_write: dict[str, float] = {}


class JobCancelled(Exception):
    """El usuario canceló el trabajo OCR."""


class JobControl:
    def __init__(self, job_id: str) -> None:
        self.job_id = job_id
        self.cancel_event = threading.Event()
        self.pause_event = threading.Event()
        self.resume_event = threading.Event()
        self.resume_event.set()
        self.yield_on_pause = False
        self.slot_yielded = False
        self.holds_semaphore = False

    def request_pause(self, *, yield_slot: bool = False) -> None:
        self.yield_on_pause = yield_slot
        self.pause_event.set()
        self.resume_event.clear()

    def request_resume(self) -> None:
        self.pause_event.clear()
        self.resume_event.set()

    def request_cancel(self) -> None:
        self.cancel_event.set()
        self.resume_event.set()

    def _release_run_slot(self, job_id: str) -> None:
        global _running_job_id
        if not self.yield_on_pause or self.slot_yielded or not self.holds_semaphore:
            return
        with _queue_lock:
            if _running_job_id != job_id:
                return
            _running_job_id = None
            if job_id not in _pending_job_ids:
                _pending_job_ids.append(job_id)
            self.slot_yielded = True
            self.holds_semaphore = False
        _job_run_semaphore.release()
        sync_queue_positions()
        kick_next_queue_worker()

    def _reacquire_run_slot(self, job_id: str) -> None:
        if not self.slot_yielded:
            return
        _job_run_semaphore.acquire()
        self.holds_semaphore = True
        with _queue_lock:
            if job_id in _pending_job_ids:
                _pending_job_ids.remove(job_id)
            global _running_job_id
            _running_job_id = job_id
        self.slot_yielded = False
        self.yield_on_pause = False

    def wait_if_paused(self, job_id: str, state_base: dict[str, Any]) -> None:
        if not self.pause_event.is_set():
            return
        self._release_run_slot(job_id)
        detail = (
            "En pausa — otro trabajo puede usar el procesador. Pulsa Reanudar."
            if self.slot_yielded
            else "En pausa. Pulsa Reanudar para continuar."
        )
        write_job_state(job_id, {
            **state_base,
            "status": "paused",
            "detail": detail,
        })
        while self.pause_event.is_set() and not self.cancel_event.is_set():
            self.resume_event.wait(timeout=0.4)
        if self.cancel_event.is_set():
            raise JobCancelled()
        self._reacquire_run_slot(job_id)
        write_job_state(job_id, {
            **state_base,
            "status": "processing",
            "detail": "Reanudando procesamiento…",
        })


def check_job_control_or_raise() -> None:
    ctrl = _job_control_ctx.get()
    if ctrl and ctrl.cancel_event.is_set():
        raise JobCancelled()


def report_ocr_progress(done: int, total: int, detail: str) -> None:
    check_job_control_or_raise()
    ctrl = _job_control_ctx.get()
    if ctrl:
        state = read_job_state(ctrl.job_id) or {}
        base = {
            k: state[k]
            for k in (
                "source_filename",
                "lang",
                "mode",
                "pages_total",
                "queue_position",
                "created_at",
            )
            if k in state
        }
        ctrl.wait_if_paused(ctrl.job_id, base)
        check_job_control_or_raise()
    cb = _progress_callback.get()
    if not cb or total <= 0:
        return
    pct = max(0, min(100, int((done / total) * 100)))
    try:
        cb(done, total, detail, pct)
    except TypeError:
        cb(done, total, detail)


def release_image(image: Optional[Image.Image]) -> None:
    if image is None:
        return
    try:
        image.close()
    except Exception:
        pass


def job_state_path(job_id: str) -> Path:
    return JOBS_DIR / f"{job_id}.json"


def write_job_state(job_id: str, payload: dict[str, Any]) -> None:
    data = dict(payload)
    data["job_id"] = job_id
    data["updated_at"] = datetime.now(timezone.utc).isoformat()
    job_state_path(job_id).write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")


def read_job_state(job_id: str) -> Optional[dict[str, Any]]:
    path = job_state_path(job_id)
    if not path.exists():
        return None
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        return data if isinstance(data, dict) else None
    except (json.JSONDecodeError, OSError):
        return None


def cleanup_stale_jobs() -> None:
    cutoff = datetime.now(timezone.utc).timestamp() - (OCR_JOB_TTL_HOURS * 3600)
    for path in JOBS_DIR.glob("*.json"):
        if path.name == "history.json":
            continue
        try:
            if path.stat().st_mtime < cutoff:
                path.unlink(missing_ok=True)
        except OSError:
            pass


def load_job_history() -> list[dict[str, Any]]:
    if not HISTORY_PATH.exists():
        return []
    try:
        data = json.loads(HISTORY_PATH.read_text(encoding="utf-8"))
        return data if isinstance(data, list) else []
    except (json.JSONDecodeError, OSError):
        return []


def history_entry_from_job_state(job_id: str, state: dict[str, Any]) -> Optional[dict[str, Any]]:
    status = state.get("status")
    if status not in ("completed", "failed", "cancelled"):
        return None
    result = state.get("result") if isinstance(state.get("result"), dict) else {}
    pf = state.get("page_from")
    pt = state.get("page_to")
    source_filename = state.get("source_filename") or result.get("source_filename") or job_id
    if status == "completed" and pf and pt and not str(source_filename).endswith(")"):
        label = f"{source_filename} ({pf}-{pt})"
    else:
        label = source_filename
    return {
        "job_id": job_id,
        "source_filename": label,
        "status": status,
        "completed_at": state.get("updated_at") or state.get("created_at") or "",
        "pages_total": state.get("pages_total") or result.get("total_pages") or 0,
        "page_from": pf,
        "page_to": pt,
        "batch_id": state.get("batch_id") or result.get("batch_id"),
        "method": result.get("method"),
        "downloads": result.get("downloads"),
        "download_names": result.get("download_names"),
        "batch_downloads": result.get("batch_downloads"),
        "full_document": bool(result.get("full_document")),
    }


def get_job_history() -> list[dict[str, Any]]:
    """Últimos trabajos: historial guardado + trabajos completados en disco."""
    by_id: dict[str, dict[str, Any]] = {}
    for entry in load_job_history():
        job_id = entry.get("job_id")
        if job_id:
            by_id[job_id] = entry

    for path in sorted(JOBS_DIR.glob("*.json"), key=lambda item: item.stat().st_mtime, reverse=True):
        if path.name == "history.json":
            continue
        job_id = path.stem
        if job_id in by_id:
            continue
        state = read_job_state(job_id)
        if not state:
            continue
        entry = history_entry_from_job_state(job_id, state)
        if entry:
            by_id[job_id] = entry

    items = sorted(
        by_id.values(),
        key=lambda item: item.get("completed_at") or "",
        reverse=True,
    )
    return items[:HISTORY_MAX_ITEMS]


def append_job_history(entry: dict[str, Any]) -> None:
    history = load_job_history()
    history = [item for item in history if item.get("job_id") != entry.get("job_id")]
    history.insert(0, entry)
    HISTORY_PATH.write_text(
        json.dumps(history[:HISTORY_MAX_ITEMS], ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def public_job_summary(state: dict[str, Any]) -> dict[str, Any]:
    status = state.get("status", "")
    queue_position = state.get("queue_position")
    if queue_position is None:
        queue_position = 0 if status in ("processing", "paused") else 1
    return {
        "job_id": state.get("job_id"),
        "status": status,
        "progress": state.get("progress", 0),
        "detail": state.get("detail", ""),
        "pages_done": state.get("pages_done", 0),
        "pages_total": state.get("pages_total", 0),
        "source_filename": state.get("source_filename", ""),
        "page_from": state.get("page_from"),
        "page_to": state.get("page_to"),
        "mode": state.get("mode"),
        "lang": state.get("lang"),
        "queue_position": queue_position,
        "created_at": state.get("created_at"),
        "updated_at": state.get("updated_at"),
        "error": state.get("error"),
    }


def _parse_job_updated_at(state: dict[str, Any]) -> Optional[datetime]:
    raw = state.get("updated_at") or state.get("created_at")
    if not raw:
        return None
    try:
        return datetime.fromisoformat(str(raw).replace("Z", "+00:00"))
    except ValueError:
        return None


def _job_upload_missing(state: dict[str, Any]) -> bool:
    upload_raw = (state.get("upload_path") or "").strip()
    if not upload_raw:
        return True
    return not Path(upload_raw).exists()


def reconcile_job_queue() -> dict[str, int]:
    """Sincroniza cola en memoria con trabajos en disco; limpia huérfanos."""
    global _running_job_id
    now = datetime.now(timezone.utc)
    stats = {"requeued": 0, "cancelled": 0, "stale_running_cleared": 0}

    for path in sorted(JOBS_DIR.glob("*.json")):
        if path.name == "history.json":
            continue
        job_id = path.stem
        state = read_job_state(job_id)
        if not state:
            continue
        status = state.get("status", "")
        if status not in ("queued", "processing", "paused"):
            continue

        if _job_upload_missing(state):
            write_job_state(job_id, {
                **state,
                "status": "cancelled",
                "detail": "Trabajo huérfano (archivo no encontrado). Vuelve a subir el PDF.",
                "error": "upload_path missing",
            })
            stats["cancelled"] += 1
            with _queue_lock:
                _pending_job_ids[:] = [jid for jid in _pending_job_ids if jid != job_id]
                if _running_job_id == job_id:
                    _running_job_id = None
            _job_controls.pop(job_id, None)
            continue

        updated = _parse_job_updated_at(state)
        age_min = ((now - updated).total_seconds() / 60.0) if updated else 9999.0

        with _queue_lock:
            in_memory = job_id in _pending_job_ids or _running_job_id == job_id

        if status == "processing" and not in_memory and age_min > 8:
            write_job_state(job_id, {
                **state,
                "status": "failed",
                "detail": "Procesamiento interrumpido. Sube el archivo de nuevo.",
                "error": "stale processing",
            })
            stats["stale_running_cleared"] += 1
            continue

        if status in ("queued", "paused") and not in_memory:
            if age_min > 180:
                write_job_state(job_id, {
                    **state,
                    "status": "cancelled",
                    "detail": "Expirado en cola. Vuelve a subir el documento.",
                })
                stats["cancelled"] += 1
                continue
            with _queue_lock:
                if job_id not in _pending_job_ids:
                    _pending_job_ids.append(job_id)
                    _job_controls.setdefault(job_id, JobControl(job_id))
            with _queue_lock:
                pos = len(_pending_job_ids) + (1 if _running_job_id else 0)
            write_job_state(job_id, {
                **state,
                "status": "queued" if status != "paused" else "paused",
                "detail": "Reencolado — iniciará en breve.",
                "queue_position": pos,
            })
            stats["requeued"] += 1
            spawn_ocr_worker(job_id)

    with _queue_lock:
        if _running_job_id:
            run_state = read_job_state(_running_job_id)
            worker_alive = _running_job_id in _workers_started
            if not run_state or (
                run_state.get("status") not in ("processing", "paused") and not worker_alive
            ):
                _running_job_id = None
                stats["stale_running_cleared"] += 1

    stats["unstuck"] = unstick_queue_if_deadlocked()
    return stats


def unstick_queue_if_deadlocked(*, force: bool = False) -> int:
    """Si hay trabajos esperando turno pero nadie procesa, reinicia el semáforo y reencola."""
    global _job_run_semaphore, _running_job_id
    now = datetime.now(timezone.utc)
    with _queue_lock:
        if _running_job_id:
            return 0
        pending = list(_pending_job_ids)

    stuck_ids: list[str] = []
    for job_id in pending:
        state = read_job_state(job_id) or {}
        if state.get("status") not in ("queued", "paused"):
            continue
        detail = state.get("detail") or ""
        if not force and "Esperando turno" not in detail and state.get("status") != "queued":
            continue
        if force:
            stuck_ids.append(job_id)
            continue
        updated = _parse_job_updated_at(state)
        age_sec = ((now - updated).total_seconds() if updated else 9999.0)
        if age_sec > 20:
            stuck_ids.append(job_id)

    if not stuck_ids:
        return 0

    with _queue_lock:
        if _running_job_id:
            return 0
        for job_id in stuck_ids:
            _workers_started.discard(job_id)
        _job_run_semaphore = threading.Semaphore(OCR_MAX_CONCURRENT_JOBS)
        _running_job_id = None

    for job_id in stuck_ids:
        spawn_ocr_worker(job_id)
    return len(stuck_ids)


def list_active_jobs() -> list[dict[str, Any]]:
    reconcile_job_queue()
    with _queue_lock:
        running = _running_job_id
        pending = list(_pending_job_ids)
    items: list[dict[str, Any]] = []
    seen: set[str] = set()

    if running:
        state = read_job_state(running)
        if state and state.get("status") not in ("cancelled", "failed", "completed"):
            row = public_job_summary(state)
            row["queue_position"] = 0
            row["is_active"] = True
            row["is_running"] = True
            row["in_queue"] = False
            items.append(row)
            seen.add(running)

    for index, job_id in enumerate(pending, start=1):
        if job_id in seen:
            continue
        state = read_job_state(job_id)
        if not state or state.get("status") in ("cancelled", "failed", "completed"):
            continue
        row = public_job_summary(state)
        row["queue_position"] = index if running else index
        if not running and index == 1:
            row["queue_position"] = 1
        row["is_active"] = False
        row["is_running"] = False
        row["in_queue"] = True
        items.append(row)
        seen.add(job_id)

    for path in sorted(JOBS_DIR.glob("*.json"), key=lambda p: p.stat().st_mtime, reverse=True):
        if path.name == "history.json":
            continue
        job_id = path.stem
        if job_id in seen:
            continue
        state = read_job_state(job_id)
        if not state or state.get("status") not in ("queued", "processing", "paused"):
            continue
        row = public_job_summary(state)
        row["queue_position"] = row.get("queue_position") or 99
        row["is_active"] = state.get("status") == "processing"
        row["is_running"] = job_id == running or state.get("status") == "processing"
        row["in_queue"] = state.get("status") in ("queued", "paused")
        items.append(row)
        seen.add(job_id)

    items.sort(key=lambda row: (row.get("queue_position", 99), row.get("created_at") or ""))
    return items


def sync_queue_positions() -> None:
    with _queue_lock:
        running = _running_job_id
        pending = list(_pending_job_ids)
    offset = 1 if running else 0
    for index, job_id in enumerate(pending, start=1):
        state = read_job_state(job_id)
        if not state or state.get("status") in ("cancelled", "failed", "completed"):
            continue
        pos = index + offset
        prev_pos = state.get("queue_position")
        if prev_pos == pos and state.get("status") != "queued":
            continue
        write_job_state(job_id, {
            **state,
            "queue_position": pos,
            "status": state.get("status") if state.get("status") == "paused" else "queued",
            "detail": (
                f"En pausa · posición {pos} en cola."
                if state.get("status") == "paused"
                else f"En cola · posición {pos}."
                + (" Hay otro documento en curso." if running else " Iniciará en breve.")
            ),
        })


def kick_next_queue_worker() -> None:
    with _queue_lock:
        pending = list(_pending_job_ids)
    for job_id in pending:
        state = read_job_state(job_id)
        if not state or state.get("status") in ("cancelled", "failed", "completed", "paused"):
            continue
        spawn_ocr_worker(job_id)
        break


def spawn_ocr_worker(job_id: str) -> None:
    """Arranca el OCR en un hilo; el semáforo serializa si hay varios trabajos."""
    with _queue_lock:
        if job_id in _workers_started:
            return
        _workers_started.add(job_id)

    def runner() -> None:
        global _running_job_id
        state = read_job_state(job_id) or {}
        if state.get("status") in ("cancelled", "paused"):
            _workers_started.discard(job_id)
            return
        write_job_state(job_id, {
            **state,
            "status": "queued",
            "detail": "Esperando turno del procesador…",
        })
        control = _job_controls.get(job_id) or JobControl(job_id)
        _job_controls[job_id] = control
        _job_run_semaphore.acquire()
        control.holds_semaphore = True
        try:
            state = read_job_state(job_id) or {}
            if state.get("status") in ("cancelled", "paused"):
                return
            with _queue_lock:
                if job_id in _pending_job_ids:
                    _pending_job_ids.remove(job_id)
                _running_job_id = job_id
            state_now = read_job_state(job_id) or {}
            write_job_state(job_id, {
                **state_now,
                "status": "processing",
                "progress": 0,
                "detail": "Procesando documento…",
                "queue_position": 0,
                "pages_done": state_now.get("pages_done", 0),
                "pages_total": state_now.get("pages_total", 0),
            })
            sync_queue_positions()
            execute_ocr_job(job_id)
        finally:
            with _queue_lock:
                if _running_job_id == job_id:
                    _running_job_id = None
            if control.holds_semaphore:
                _job_run_semaphore.release()
                control.holds_semaphore = False
            _workers_started.discard(job_id)
            kick_next_queue_worker()

    threading.Thread(target=runner, daemon=True, name=f"ocr-worker-{job_id}").start()


def prepare_work_pdf(
    upload_path: Path,
    page_from: int,
    page_to: Optional[int],
) -> tuple[Path, int, int, bool]:
    """Devuelve (ruta de trabajo, page_from, page_to, si hay que borrar el slice temporal)."""
    if upload_path.suffix.lower() != ".pdf":
        return upload_path, 1, 1, False
    total = count_pdf_pages(upload_path)
    pf, pt = batches.normalize_page_range(page_from, page_to, total)
    if pf == 1 and pt == total:
        return upload_path, pf, pt, False
    slice_path = (UPLOAD_DIR / f"slice_{uuid.uuid4().hex[:8]}.pdf").resolve()
    batches.extract_pdf_slice(upload_path.resolve(), pf, pt, slice_path)
    return slice_path, pf, pt, True


def should_delete_upload_after_job(state: dict[str, Any], upload_path: Path) -> bool:
    if state.get("keep_source"):
        return False
    try:
        return upload_path.resolve().parent != SOURCES_DIR.resolve()
    except OSError:
        return True


def enqueue_ocr_job(
    job_id: str,
    upload_path: Path,
    *,
    lang: str,
    mode: str,
    original_filename: str,
    pages_total: int,
    page_from: int = 1,
    page_to: Optional[int] = None,
    batch_id: Optional[str] = None,
    keep_source: bool = False,
    range_slice: Optional[bool] = None,
    processing_mode: Optional[str] = None,
    vision_provider: Optional[str] = None,
) -> int:
    control = JobControl(job_id)
    with _queue_lock:
        _job_controls[job_id] = control
        if _running_job_id == job_id:
            return 0
        if job_id in _pending_job_ids:
            queue_position = _pending_job_ids.index(job_id) + 1
            if _running_job_id:
                queue_position += 1
            spawn_ocr_worker(job_id)
            return queue_position
        _pending_job_ids.append(job_id)
        queue_position = len(_pending_job_ids)
        if _running_job_id:
            queue_position += 1
    if upload_path.suffix.lower() == ".pdf":
        pf, pt = batches.normalize_page_range(page_from, page_to, count_pdf_pages(upload_path))
    else:
        pf, pt = 1, 1
    range_pages = pt - pf + 1
    part_label = f"páginas {pf}-{pt}"
    effective_range_slice = bool(range_slice) if range_slice is not None else bool(batch_id)
    effective_processing_mode = processing_mode or ("ranges" if effective_range_slice else "full")
    created_at = datetime.now(timezone.utc).isoformat()
    write_job_state(job_id, {
        "status": "queued",
        "progress": 0,
        "detail": (
            f"En cola · {part_label} (posición {queue_position}). "
            + ("Otro documento en curso." if queue_position > 1 else "Iniciará en breve.")
        ),
        "pages_done": 0,
        "pages_total": range_pages,
        "page_from": pf,
        "page_to": pt,
        "part_label": part_label,
        "batch_id": batch_id,
        "keep_source": keep_source,
        "range_slice": effective_range_slice,
        "processing_mode": effective_processing_mode,
        "source_filename": original_filename,
        "lang": lang,
        "mode": mode,
        "vision_provider": normalize_vision_provider(vision_provider) or "auto",
        "upload_path": str(upload_path.resolve()),
        "queue_position": queue_position,
        "created_at": created_at,
    })
    if batch_id:
        try:
            batches.add_batch_part(JOBS_DIR, batch_id, job_id=job_id, page_from=pf, page_to=pt)
        except KeyError:
            pass
    sync_queue_positions()
    spawn_ocr_worker(job_id)
    return queue_position


def recover_jobs_on_startup() -> None:
    for path in sorted(JOBS_DIR.glob("*.json")):
        if path.name == "history.json":
            continue
        try:
            state = json.loads(path.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError):
            continue
        if not isinstance(state, dict):
            continue
        job_id = state.get("job_id") or path.stem
        status = state.get("status", "")
        upload_raw = state.get("upload_path", "")
        upload_path = Path(upload_raw) if upload_raw else None

        if status in ("processing", "queued", "paused") and upload_path and upload_path.exists():
            with _queue_lock:
                already = job_id in _pending_job_ids or _running_job_id == job_id or job_id in _workers_started
            if not already:
                enqueue_ocr_job(
                    job_id,
                    upload_path,
                    lang=state.get("lang", "spa"),
                    mode=state.get("mode", "balanced"),
                    original_filename=state.get("source_filename", "documento"),
                    pages_total=int(state.get("pages_total") or 1),
                    page_from=int(state.get("page_from") or 1),
                    page_to=int(state.get("page_to")) if state.get("page_to") is not None else None,
                    batch_id=state.get("batch_id"),
                    keep_source=bool(state.get("keep_source")),
                )
                refreshed = read_job_state(job_id) or state
                write_job_state(job_id, {
                    **refreshed,
                    "detail": "Reencolado tras reinicio del servidor",
                })
            continue

        if status == "processing":
            write_job_state(job_id, {
                **state,
                "status": "failed",
                "detail": "Interrumpido por reinicio del servidor",
                "error": "Reinicio del servidor durante el procesamiento",
            })


def should_process_async(file_path: Path, size_bytes: int, *, force_async: bool, force_sync: bool) -> bool:
    if force_sync:
        return False
    if force_async or OCR_ASYNC_ALWAYS:
        return True
    if size_bytes >= int(OCR_HEAVY_ASYNC_MB * 1024 * 1024):
        return True
    if file_path.suffix.lower() == ".pdf" and count_pdf_pages(file_path) >= OCR_HEAVY_ASYNC_PAGES:
        return True
    return False


def gc_every_n_pages(page_index: int, interval: int) -> None:
    if interval > 0 and (page_index + 1) % interval == 0:
        gc.collect()


def is_large_document(page_count: int) -> bool:
    return page_count >= OCR_LARGE_PDF_PAGES


def gemini_allowed_for_pages(page_count: int) -> bool:
    return vision_allowed_for_pages(page_count)


def vision_allowed_for_pages(page_count: int) -> bool:
    if not OCR_VISION_ENABLED:
        return False
    return bool(vision_llm.configured_providers()) and page_count <= OCR_LARGE_PDF_GEMINI_MAX


def vision_tables_allowed() -> bool:
    if not OCR_VISION_TABLES_ONLY:
        return False
    return bool(vision_llm.providers_for_tables())


def vision_docs_allowed() -> bool:
    if not OCR_VISION_DOCS_ONLY:
        return False
    return bool(vision_llm.providers_for_documents())


def vision_legal_allowed() -> bool:
    if not OCR_VISION_LEGAL_ONLY:
        return False
    return bool(vision_llm.providers_for_legal())


def build_ocr_profile(mode: str, page_count: int) -> dict[str, Any]:
    mode = normalize_ocr_mode(mode)
    if not is_large_document(page_count):
        return {
            "bulk": False,
            "page_count": page_count,
            "mode": mode,
            "dpi": ocr_dpi_for_mode(mode),
            "min_edge": upscale_min_edge_for_mode(mode),
            "use_gemini": False if OCR_TESSERACT_ONLY else vision_allowed_for_pages(page_count),
            "use_vision": False if OCR_TESSERACT_ONLY else vision_allowed_for_pages(page_count),
            "supplementary": OCR_SUPPLEMENTARY_PASS and mode != "fast",
            "skip_osd": True,
            "fix_rotation": OCR_FIX_ROTATION,
        }
    bulk_mode = mode
    bulk_dpi = ocr_dpi_for_mode(mode)
    if mode == "balanced":
        bulk_dpi = max(OCR_LARGE_DPI, ocr_dpi_for_mode("balanced"))
    elif mode == "fast":
        bulk_dpi = min(OCR_LARGE_DPI, ocr_dpi_for_mode("fast"))
    elif mode == "quality":
        bulk_dpi = ocr_dpi_for_mode("quality")
    bulk_edge = OCR_LARGE_UPSCALE if mode != "quality" else min(OCR_LARGE_UPSCALE + 400, OCR_UPSCALE_MIN_EDGE)
    return {
        "bulk": True,
        "page_count": page_count,
        "mode": bulk_mode,
        "dpi": bulk_dpi,
        "min_edge": bulk_edge,
        "use_gemini": False if OCR_TESSERACT_ONLY else gemini_allowed_for_pages(page_count),
        "use_vision": False if OCR_TESSERACT_ONLY else vision_allowed_for_pages(page_count),
        "supplementary": OCR_SUPPLEMENTARY_PASS and mode != "fast",
        "skip_osd": True,
        "fix_rotation": OCR_FIX_ROTATION and OCR_BULK_FIX_ROTATION and mode != "fast",
    }


def build_ocr_profile_for_range(mode: str, range_pages: int) -> dict[str, Any]:
    """Trozos por rango: equilibrado = DPI moderado + OCR rápido; calidad = máximo detalle."""
    mode = normalize_ocr_mode(mode)
    effective = mode
    if mode == "balanced":
        return {
            "bulk": False,
            "range_slice": True,
            "balanced_fast": range_pages > 35,
            "page_count": range_pages,
            "mode": "balanced",
            "dpi": ocr_dpi_for_mode("balanced"),
            "min_edge": upscale_min_edge_for_mode("balanced"),
            "use_gemini": False if OCR_TESSERACT_ONLY else gemini_allowed_for_pages(range_pages),
            "use_vision": False if OCR_TESSERACT_ONLY else vision_allowed_for_pages(range_pages),
            "supplementary": OCR_SUPPLEMENTARY_PASS,
            "skip_osd": True,
            "fix_rotation": OCR_FIX_ROTATION,
        }
    if mode == "fast":
        return {
            "bulk": False,
            "range_slice": True,
            "page_count": range_pages,
            "mode": "fast",
            "dpi": ocr_dpi_for_mode("fast"),
            "min_edge": upscale_min_edge_for_mode("fast"),
            "use_gemini": False,
            "use_vision": False,
            "supplementary": False,
            "skip_osd": True,
            "fix_rotation": False,
        }
    return {
        "bulk": False,
        "range_slice": True,
        "page_count": range_pages,
        "mode": "quality",
        "dpi": ocr_dpi_for_mode("quality"),
        "min_edge": upscale_min_edge_for_mode("quality"),
        "use_gemini": False if OCR_TESSERACT_ONLY else gemini_allowed_for_pages(range_pages),
        "use_vision": False if OCR_TESSERACT_ONLY else vision_allowed_for_pages(range_pages),
        "supplementary": OCR_SUPPLEMENTARY_PASS,
        "skip_osd": True,
        "fix_rotation": OCR_FIX_ROTATION,
    }


def get_ocr_profile() -> dict[str, Any]:
    return _ocr_profile_ctx.get() or {}


def set_ocr_profile(profile: dict[str, Any]) -> contextvars.Token:
    return _ocr_profile_ctx.set(profile)


def normalize_vision_provider(provider: Optional[str]) -> Optional[str]:
    choice = (provider or "").strip().lower()
    if not choice or choice == "auto":
        return None
    if choice in vision_llm.configured_providers():
        return choice
    return None


def get_vision_provider() -> Optional[str]:
    return _vision_provider_ctx.get()


def set_vision_provider(provider: Optional[str]) -> contextvars.Token:
    return _vision_provider_ctx.set(provider)


def _is_vision_ocr_source(source: str) -> bool:
    base = (source or "").split("-", 1)[0]
    if base == "parallel" or base.startswith("parallel:"):
        return True
    return base in ("gemini", "openai", "anthropic", "vision")


def _vision_source_providers(source: str) -> list[str]:
    if not source:
        return []
    base = source.split("-", 1)[0]
    if base.startswith("parallel:"):
        return [p for p in base.split(":", 1)[-1].split("+") if p]
    if base in ("gemini", "openai", "anthropic"):
        return [base]
    return []


def throttle_job_progress_write(job_id: str, interval_sec: float = 1.5) -> bool:
    if not job_id:
        return True
    now = time.monotonic()
    last = _last_progress_write.get(job_id, 0.0)
    if now - last >= interval_sec:
        _last_progress_write[job_id] = now
        return True
    return False


def normalize_ocr_mode(mode: Optional[str]) -> str:
    m = (mode or "balanced").strip().lower()
    if m in ("fast", "rapid", "speed", "rapido"):
        return "fast"
    if m in ("quality", "high", "best", "max", "calidad"):
        return "quality"
    return "balanced"


def ocr_dpi_for_mode(mode: str) -> int:
    if mode == "fast":
        return int(os.getenv("OCR_DPI_FAST", "200"))
    if mode == "balanced":
        return int(os.getenv("OCR_DPI_BALANCED", "220"))
    return int(os.getenv("OCR_DPI_QUALITY", str(OCR_RENDER_DPI)))


def upscale_min_edge_for_mode(mode: str) -> int:
    if mode == "fast":
        return int(os.getenv("OCR_UPSCALE_FAST", "1200"))
    if mode == "balanced":
        return int(os.getenv("OCR_UPSCALE_BALANCED", "1200"))
    return OCR_UPSCALE_MIN_EDGE

TESSERACT_CMD = os.getenv("TESSERACT_CMD", r"C:\Program Files\Tesseract-OCR\tesseract.exe")
pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD


def local_ocr_engines() -> list[str]:
    engines = ["tesseract"]
    if OCR_PADDLE_ENABLED and paddle_ocr.paddle_ocr_available():
        engines.append("paddle")
    return engines


def paddle_ocr_active() -> bool:
    return OCR_PADDLE_ENABLED and paddle_ocr.paddle_ocr_available()


def hybrid_ocr_parallel_enabled() -> bool:
    """Paddle no es thread-safe en Windows; serializar páginas escaneadas."""
    if paddle_ocr_active():
        return False
    return OCR_HYBRID_PARALLEL


def paddle_layout_active() -> bool:
    return paddle_ocr_active() and OCR_PADDLE_LAYOUT


def run_paddle_layout_page_markdown(
    image: Image.Image,
    lang: str,
    *,
    mode: str,
) -> tuple[str, dict[str, Any]]:
    """OCR Paddle con bloques y el mismo post-procesado que Tesseract."""
    layout = paddle_ocr.extract_page_layout(
        image,
        lang=OCR_PADDLE_LANG or lang,
        line_filter=is_stamp_ocr_gibberish,
    )
    cleaned = clean_text(layout.get("text") or "")
    if not cleaned:
        return "", layout
    markdown = sanitize_document_markdown(cleaned, aggressive=False)
    profile = get_ocr_profile()
    effective_mode = profile.get("mode", mode)
    if OCR_SUPPLEMENTARY_PASS and markdown:
        markdown = enrich_markdown_with_annotations(
            markdown, image, lang=lang, mode=effective_mode
        )
    return markdown, layout


def run_paddle_page_markdown(
    image: Image.Image,
    lang: str,
    *,
    mode: str,
) -> tuple[str, float]:
    """OCR Paddle con el mismo post-procesado que Tesseract (sin pasada de sellos)."""
    raw, confidence = paddle_ocr.extract_text_with_paddle(
        image,
        lang=OCR_PADDLE_LANG or lang,
        line_filter=is_stamp_ocr_gibberish,
    )
    cleaned = clean_text(raw)
    if not cleaned:
        return "", 0.0
    markdown = sanitize_document_markdown(cleaned, aggressive=False)
    score = text_quality_score(markdown) + confidence * 1.35
    return markdown, score


def pick_best_local_ocr_markdown(
    image: Image.Image,
    lang: str,
    *,
    mode: str,
) -> tuple[str, str]:
    """Ejecuta Tesseract y Paddle (si está activo) y devuelve (markdown, motor)."""
    tesseract_md = flatten_illegal_pipe_tables(
        ocr_image_to_markdown(image, lang=lang, mode=mode)
    )
    tesseract_score = text_quality_score(tesseract_md)
    best_md, best_source, best_score = tesseract_md, "tesseract", tesseract_score

    if paddle_ocr_active():
        try:
            paddle_md, paddle_score = run_paddle_page_markdown(image, lang=lang, mode=mode)
            if paddle_md and paddle_score > best_score:
                profile = get_ocr_profile()
                effective_mode = profile.get("mode", mode)
                best_md = enrich_markdown_with_annotations(
                    paddle_md, image, lang=lang, mode=effective_mode
                )
                best_source = "paddle"
        except Exception:
            pass

    return best_md, best_source

if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)

# ── App ───────────────────────────────────────────────────────────────────────
app = FastAPI(title="SCI OCR API", version="1.0.0")


@app.on_event("startup")
def _log_ocr_capabilities() -> None:
    if OCR_TESSERACT_ONLY:
        engines = local_ocr_engines()
        paddle_note = " paddle_serial=1" if paddle_ocr_active() else ""
        print(
            f"[SCI OCR] motores={','.join(engines)} layout={OCR_USE_LAYOUT} "
            f"rotacion_auto={OCR_FIX_ROTATION} sellos_aparte={OCR_SUPPLEMENTARY_PASS}{paddle_note}"
        )
    else:
        doc_on = OCR_VISION_DOCS_ONLY and bool(vision_llm.providers_for_documents())
        legal_on = OCR_VISION_LEGAL_ONLY and bool(vision_llm.providers_for_legal())
        tbl_on = OCR_VISION_TABLES_ONLY and bool(vision_llm.providers_for_tables())
        print(
            f"[SCI OCR] legal_ia={legal_on} docs_ia={doc_on} tablas_ia={tbl_on} "
            f"vision_global={OCR_VISION_ENABLED}"
        )


def cancel_orphan_jobs_on_disk() -> int:
    """Cancela trabajos en disco sin archivo PDF asociado."""
    global _running_job_id
    cancelled = 0
    for path in JOBS_DIR.glob("*.json"):
        if path.name == "history.json":
            continue
        job_id = path.stem
        state = read_job_state(job_id)
        if not state:
            continue
        if state.get("status") not in ("queued", "processing", "paused"):
            continue
        if not _job_upload_missing(state):
            continue
        write_job_state(job_id, {
            **state,
            "status": "cancelled",
            "detail": "Trabajo huérfano. Sube el PDF de nuevo.",
            "error": "upload_path missing",
        })
        cancelled += 1
        with _queue_lock:
            _pending_job_ids[:] = [jid for jid in _pending_job_ids if jid != job_id]
            if _running_job_id == job_id:
                _running_job_id = None
        _workers_started.discard(job_id)
    return cancelled


@app.on_event("startup")
def on_startup() -> None:
    cleanup_stale_jobs()
    cancel_orphan_jobs_on_disk()
    reconcile_job_queue()
    recover_jobs_on_startup()


app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.middleware("http")
async def limit_upload_size(request: Request, call_next):
    if request.method == "POST":
        content_length = request.headers.get("content-length")
        if content_length and int(content_length) > MAX_UPLOAD_BYTES:
            return JSONResponse(
                {"detail": f"Archivo demasiado grande. Límite: {MAX_UPLOAD_MB} MB."},
                status_code=413
            )
    return await call_next(request)

frontend_path = Path("../frontend")
if frontend_path.exists():
    app.mount("/static", StaticFiles(directory=str(frontend_path)), name="static")

# ── Auth ──────────────────────────────────────────────────────────────────────
def api_key_required() -> bool:
    return bool(API_KEY)


def check_api_key(x_api_key: Optional[str]):
    if not api_key_required():
        return
    if (x_api_key or "").strip() != API_KEY:
        raise HTTPException(status_code=403, detail="API Key inválida")


@app.get("/auth/check")
async def auth_check(x_api_key: Optional[str] = Header(None, alias="X-API-Key")):
    if not api_key_required():
        return {"required": False, "ok": True}
    return {"required": True, "ok": (x_api_key or "").strip() == API_KEY}

# ── Normalización y salida estructurada ───────────────────────────────────────
def count_pdf_pages(file_path: Path) -> int:
    if file_path.suffix.lower() != ".pdf":
        return 1
    try:
        with fitz.open(str(file_path)) as doc:
            return max(len(doc), 1)
    except Exception:
        return 1

def normalize_markdown(text: str) -> str:
    normalized = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    normalized = normalized.replace("\f", "\n\n---\n\n")
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()

def normalize_whitespace(text: str) -> str:
    text = (text or "").replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _spanish_word_ratio(text: str) -> float:
    tokens = re.findall(r"[a-záéíóúñA-ZÁÉÍÓÚÑ]{3,}", text)
    if not tokens:
        return 0.0
    good = sum(1 for t in tokens if re.search(r"[aeiouáéíóú]", t, re.I))
    return good / len(tokens)


def minimum_pdf_text_chars(file_path: Path) -> int:
    pages = count_pdf_pages(file_path)
    return min(MARKITDOWN_MIN_CHARS, max(36, pages * 18))


_LEGAL_ORDINAL_HEADING = re.compile(
    r"^(?:PRIMERO|SEGUNDO|TERCERO|CUARTO|QUINTO|SEXTO|S[EÉ]PTIMO|OCTAVO|NOVENO|D[EÉ]CIMO)\b",
    re.I,
)
_CONTRACT_CLAUSE_HEADING = re.compile(
    r"^(?:CL[AÁ]USULA|CAP[IÍ]TULO|ART[IÍ]CULO|T[IÍ]TULO|ANEXO)\s+"
    r"(?:PRIMERA?|SEGUNDA?|TERCERA?|CUARTA?|QUINTA?|SEXTA?|S[EÉ]PTIMA?|OCTAVA?|NOVENA?|D[EÉ]CIMA?|\d+\.?)\b",
    re.I,
)
_NAME_LIKE_HEADING = re.compile(
    r"^[A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s\.]{3,}$"
)
_STAMP_DATETIME = re.compile(r"\b\d{2}/\d{2}/\d{2}\s+\d{2}:\d{2}:\d{2}\b")
_CERT_SERIAL_DOTTED = re.compile(r"\b(?:\d{2}\.){6,}\d{2,}(?:\.\d{2,})*\b")
_HEX_BYTE_CHAIN = re.compile(r"\b(?:[0-9a-f]{2}\s+){12,}[0-9a-f]{2}\b", re.I)
_LEGAL_STAMP_LINE = re.compile(
    r"^(?:PRIMER|SEGUNDO|TERCER|CUARTO|QUINTO|SEXTO|AGRAVIO|RECURSO|AMPARO|"
    r"R\s*E\s*S\s*U\s*E\s*L\s*V\s*E|FUENTE|REGISTRO\s+DIGITAL)",
    re.I,
)
_SENTENCE_LIKE = re.compile(
    r"\b(?:atentamente|comparezco|solicito|continuaci[oó]n|fundamento|representada|jurisdiccional)\b",
    re.I,
)


def _contains_long_base64_blob(text: str, min_len: int = 44) -> bool:
    compact = re.sub(r"\s+", "", text or "")
    for match in re.finditer(rf"[A-Za-z0-9+/]{{{min_len},}}={{0,2}}", compact):
        chunk = match.group(0)
        if any(c in chunk for c in "+/="):
            return True
    return False


def is_cryptographic_evidence_page(text: str) -> bool:
    sample = (text or "")[:1200]
    return bool(
        re.search(r"EVIDENCIA\s+CRIPTOGR", sample, re.I)
        and re.search(r"(?:FIRMANTE|CADENA\s+DE\s+FIRMA|AUTORIDAD\s+CERTIFICADORA)", sample, re.I)
    )


def summarize_crypto_evidence_page(text: str) -> str:
    """Resume página de firma electrónica sin volcar cadenas hex/OCSP."""
    def pick(pattern: str, label: str) -> Optional[str]:
        match = re.search(pattern, text, re.I | re.S)
        if not match:
            return None
        value = normalize_whitespace(match.group(1))
        return f"- **{label}:** {value}" if value else None

    lines = [
        "### Evidencia criptográfica (resumen)",
        "",
        pick(r"Archivo\s+Firmado:\s*\n?\s*(\S+)", "Archivo firmado"),
        pick(r"FIRMANTE\s+Nombre:\s*\n\s*([^\n]+)", "Firmante"),
        pick(r"FIRMA\s+No\.\s*serie:\s*\n\s*([^\n]+)", "No. serie"),
        pick(r"Algoritmo:\s*\n?\s*([^\n]+)", "Algoritmo"),
        pick(
            r"Fecha:\s*\n?\s*\(UTC/?\s*CDMX\)\s*\n?\s*([^\n]+)",
            "Fecha de firma",
        ),
        pick(r"Status:\s*\n?\s*([^\n]+)", "Estatus"),
        pick(r"Validez:\s*\n?\s*([^\n]+)", "Validez"),
        pick(r"Nombre del respondedor:\s*\n?\s*([^\n]+)", "OCSP"),
        pick(r"Identificador de la respuesta TSP:\s*\n?\s*(\S+)", "TSP"),
    ]
    body = "\n".join(line for line in lines if line)
    if len(body.splitlines()) <= 2:
        body = "### Evidencia criptográfica\n\n> Página de firma electrónica detectada."
    body += (
        "\n\n> Cadena de firma, sellos de tiempo y tablas hex omitidos; "
        "solo metadatos legibles."
    )
    return body


def _is_watermark_stamp_line(line: str) -> bool:
    stripped = (line or "").strip()
    if not stripped:
        return False
    if _LEGAL_STAMP_LINE.match(stripped):
        return False
    if _STAMP_DATETIME.search(stripped) and len(stripped) < 40:
        return True
    if _CERT_SERIAL_DOTTED.search(stripped) and len(stripped) < 100:
        return True
    if _HEX_BYTE_CHAIN.search(stripped):
        return True
    letters = re.sub(r"[^A-Za-zÁÉÍÓÚáéíóúÑñ]", "", stripped)
    if (
        letters
        and letters.isupper()
        and 2 <= len(stripped.split()) <= 5
        and len(stripped) < 56
        and not re.search(r"[a-záéíóúñ]", stripped)
        and not any(t in stripped.upper() for t in (" DE ", " DEL ", " LA ", " EL ", " Y "))
    ):
        return True
    return False


def strip_electronic_watermarks(text: str) -> str:
    """Quita sellos electrónicos del Poder Judicial (nombre, fecha, serie cert.)."""
    if not text:
        return ""
    cleaned_lines: list[str] = []
    for line in text.replace("\r\n", "\n").split("\n"):
        stripped = line.strip()
        if not stripped:
            cleaned_lines.append("")
            continue
        if _is_watermark_stamp_line(stripped):
            continue
        row = _STAMP_DATETIME.sub(" ", stripped)
        row = _CERT_SERIAL_DOTTED.sub(" ", row)
        row = re.sub(
            r"\b[A-ZÁÉÍÓÚÑ]{2,}(?:\s+[A-ZÁÉÍÓÚÑ]{2,}){1,4}\b",
            lambda m: "" if _is_watermark_stamp_line(m.group(0)) else m.group(0),
            row,
        )
        row = re.sub(r"\s{2,}", " ", row).strip()
        if row:
            cleaned_lines.append(row)
    merged = "\n".join(cleaned_lines)
    merged = re.sub(r"\n{3,}", "\n\n", merged)
    return merged.strip()


def normalize_digital_text(text: str) -> str:
    """Texto seleccionable del PDF: conservar saltos de línea, sin reescritura."""
    text = strip_electronic_watermarks(text or "")
    if not text:
        return ""
    lines = [line.rstrip() for line in text.replace("\r\n", "\n").split("\n")]
    return normalize_markdown("\n".join(lines).strip())


def polish_digital_legal_text(text: str) -> str:
    """Conserva párrafos del PDF digital sin fragmentar en encabezados falsos."""
    text = strip_electronic_watermarks(text or "")
    if not text:
        return ""
    blocks = re.split(r"\n\s*\n", text.replace("\r\n", "\n").strip())
    output: list[str] = []
    for block in blocks:
        lines = [line.strip() for line in block.split("\n") if line.strip()]
        if not lines:
            continue
        if len(lines) == 1:
            line = lines[0]
            if _LEGAL_ORDINAL_HEADING.match(line) or _CONTRACT_CLAUSE_HEADING.match(line):
                output.append(f"## {line.rstrip(':').strip()}")
            elif (
                line.isupper()
                and len(line) <= 42
                and line.rstrip(":").count(" ") <= 4
                and not _SENTENCE_LIKE.search(line)
            ):
                output.append(f"## {line.rstrip(':').strip()}")
            else:
                output.append(line)
            continue
        output.append(" ".join(lines))
    return normalize_markdown("\n\n".join(output))


def is_likely_heading_line(line: str) -> bool:
    stripped = (line or "").strip()
    if not stripped or len(stripped) > 100:
        return False
    if re.search(r"[a-záéíóúñ]", stripped):
        return False
    if _SENTENCE_LIKE.search(stripped):
        return False
    if stripped.endswith(":") and len(stripped) <= 72 and stripped.isupper():
        return True
    if _LEGAL_ORDINAL_HEADING.match(stripped):
        return True
    if _CONTRACT_CLAUSE_HEADING.match(stripped):
        return True
    letters = re.sub(r"[^A-Za-zÁÉÍÓÚáéíóúÑñ]", "", stripped)
    if not letters or not letters.isupper() or len(letters) < 4:
        return False
    if len(stripped) > 48:
        return False
    tokens = stripped.split()
    if 2 <= len(tokens) <= 5 and _NAME_LIKE_HEADING.match(stripped):
        return False
    return len(tokens) <= 6


def structure_plain_text_as_markdown(text: str) -> str:
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not text:
        return ""

    if re.search(r"^#{1,6}\s|^\|.+\||^\s*[-*]\s+\S|^\d+\.\s+\S", text, re.MULTILINE):
        return normalize_markdown(text)

    blocks = re.split(r"\n\s*\n", text)
    output: list[str] = []

    for block in blocks:
        lines = [line.strip() for line in block.split("\n") if line.strip()]
        if not lines:
            continue

        if len(lines) == 1:
            line = lines[0]
            if is_likely_heading_line(line):
                output.append(f"## {line.rstrip(':').strip()}")
            elif re.match(r"^\d+[\.\)]\s+", line):
                output.append(re.sub(r"^(\d+)[\.\)]\s+", r"\1. ", line))
            elif re.match(r"^[-•*]\s+", line):
                output.append("- " + re.sub(r"^[-•*]\s+", "", line))
            else:
                output.append(line)
            continue

        if all(re.match(r"^\d+[\.\)]\s+", line) for line in lines):
            output.extend(re.sub(r"^(\d+)[\.\)]\s+", r"\1. ", line) for line in lines)
            continue

        if all(re.match(r"^[-•*]\s+", line) for line in lines):
            output.extend("- " + re.sub(r"^[-•*]\s+", "", line) for line in lines)
            continue

        output.append("\n".join(lines))

    return normalize_markdown("\n\n".join(output))


def ensure_readable_markdown(text: str) -> str:
    normalized = normalize_markdown(text or "")
    if not normalized:
        return ""
    if re.search(r"^#{1,6}\s|^\|.+\||^\d+\.\s+\S", normalized, re.MULTILINE):
        return normalized
    return structure_plain_text_as_markdown(normalized)


def extract_digital_page_markdown(page) -> str:
    """Extrae texto digital del PDF conservando párrafos legibles."""
    try:
        flowing = normalize_whitespace(page.get_text("text", sort=True) or "")
    except (TypeError, ValueError):
        flowing = ""
    if is_cryptographic_evidence_page(flowing):
        return summarize_crypto_evidence_page(flowing)
    if len(flowing) >= OCR_MIN_PAGE_CHARS:
        return normalize_digital_text(flowing)

    blocks = page.get_text("blocks", sort=True) or []
    rows: list[tuple[float, float, str]] = []
    for block in blocks:
        if len(block) < 7 or block[6] != 0:
            continue
        text = normalize_whitespace(str(block[4]))
        if text:
            rows.append((float(block[1]), float(block[0]), text))

    rows.sort(key=lambda item: (round(item[0] / 4), item[1]))
    parts: list[str] = []
    for _, _, text in rows:
        if _is_watermark_stamp_line(text):
            continue
        parts.append(text)
    merged = "\n\n".join(parts).strip()
    if merged:
        return normalize_digital_text(merged)
    return ""


def render_pdf_page_image(page, dpi: int) -> Image.Image:
    """Renderiza la página aplicando la rotación del PDF."""
    matrix = fitz.Matrix(dpi / 72.0, dpi / 72.0).prerotate(page.rotation)
    pix = page.get_pixmap(matrix=matrix, alpha=False)
    return Image.frombytes("RGB", [pix.width, pix.height], pix.samples)


def _osd_probe_image(image: Image.Image, max_edge: int = 1200) -> Image.Image:
    edge = min(image.size) if image.size else 0
    if edge and edge > max_edge:
        scale = max_edge / float(edge)
        return image.resize(
            (max(1, int(image.width * scale)), max(1, int(image.height * scale))),
            Image.Resampling.LANCZOS,
        )
    return image


def _quick_orientation_score(image: Image.Image, lang: str = "spa") -> float:
    try:
        text = pytesseract.image_to_string(
            image, lang=lang, config="--oem 1 --psm 6 -c preserve_interword_spaces=1"
        )
        return text_quality_score(text)
    except Exception:
        return 0.0


def detect_and_correct_rotation(image: Image.Image, *, lang: str = "spa") -> tuple[Image.Image, int]:
    """Corrige orientación solo con OSD de alta confianza (evita voltear páginas verticales)."""
    if not OCR_FIX_ROTATION:
        return ImageOps.exif_transpose(image), 0
    try:
        probe = _osd_probe_image(image)
        osd = pytesseract.image_to_osd(probe, config="--psm 0")
        rotate_match = re.search(r"Rotate:\s+(\d+)", osd)
        conf_match = re.search(r"Orientation confidence:\s+([\d.]+)", osd)
        rotate = int(rotate_match.group(1)) % 360 if rotate_match else 0
        confidence = float(conf_match.group(1)) if conf_match else 0.0
        if rotate and confidence >= 3.5:
            return image.rotate(-rotate, expand=True, fillcolor=255), rotate
    except Exception:
        pass
    return ImageOps.exif_transpose(image), 0


def prepare_page_image_for_ocr(page, dpi: int, *, lang: str = "spa") -> Image.Image:
    image = render_pdf_page_image(page, dpi)
    profile = get_ocr_profile()
    if profile.get("fix_rotation", OCR_FIX_ROTATION):
        corrected, degrees = detect_and_correct_rotation(image, lang=lang)
        if degrees:
            fixes = list(_rotation_fixes_ctx.get() or [])
            fixes.append(page.number + 1)
            _rotation_fixes_ctx.set(fixes)
        image = corrected
    else:
        image = ImageOps.exif_transpose(image)
    if OCR_TESSERACT_ONLY:
        gray = ImageOps.grayscale(image)
        return ImageOps.autocontrast(gray, cutoff=1)
    page_kind = scan_pages.classify_scan_page(image, lang=lang)
    if page_kind == scan_pages.PAGE_CHECKLIST or scan_pages.page_needs_legal_vision(
        image, lang=lang, page_kind=page_kind
    ):
        return ImageOps.exif_transpose(image)
    return suppress_visual_noise_for_ocr(image)


def suppress_visual_noise_for_ocr(image: Image.Image) -> Image.Image:
    """Atenúa sellos, marcas de agua, resaltado y tinta de anotaciones antes del OCR principal."""
    rgb = ImageOps.exif_transpose(image).convert("RGB")
    r, g, b = rgb.split()
    gray = ImageOps.grayscale(rgb)

    red_ink = ImageChops.subtract(r, ImageChops.lighter(g, b))
    red_mask = red_ink.point(lambda v: 255 if v > 55 else 0)
    blue_ink = ImageChops.subtract(b, ImageChops.lighter(r, g))
    blue_mask = blue_ink.point(lambda v: 255 if v > 50 else 0)
    yellow = ImageChops.subtract(ImageChops.lighter(r, g), b)
    yellow_mask = yellow.point(lambda v: 255 if v > 40 else 0)

    blurred = gray.filter(ImageFilter.GaussianBlur(radius=6))
    diff = ImageChops.difference(gray, blurred)
    wm_mask = diff.point(lambda v: 255 if v < 12 else 0)
    mid_gray = gray.point(lambda v: 255 if 100 < v < 225 else 0)
    watermark_mask = ImageChops.multiply(wm_mask, mid_gray)

    noise_mask = ImageChops.lighter(
        ImageChops.lighter(red_mask, blue_mask),
        ImageChops.lighter(yellow_mask, watermark_mask),
    )
    text_layer = gray.point(lambda v: 0 if v < 145 else 255)
    clean = ImageChops.composite(Image.new("L", gray.size, 255), text_layer, noise_mask)
    return ImageOps.autocontrast(clean, cutoff=2)


def ocr_page_with_fallback(
    image: Image.Image, lang: str, mode: str
) -> tuple[str, str, Optional[dict[str, Any]]]:
    """Devuelve (markdown, origen, layout?) con origen tesseract|paddle|tesseract-table."""
    profile = get_ocr_profile()
    effective_mode = profile.get("mode", mode)
    page_kind = scan_pages.classify_scan_page(image, lang=lang)

    if OCR_TESSERACT_ONLY:
        is_table_page = (
            page_kind == scan_pages.PAGE_SPREADSHEET
            and scan_pages.page_allows_tesseract_table(page_kind)
        )
        if is_table_page:
            try:
                table_md = table_ocr.ocr_image_to_table_markdown(image, lang=lang)
                if table_md and text_quality_score(table_md) >= 20:
                    return table_md, "tesseract-table", None
            except Exception:
                pass
        if paddle_layout_active():
            markdown, layout = run_paddle_layout_page_markdown(
                image, lang=lang, mode=effective_mode
            )
            return markdown, "paddle", layout
        if paddle_ocr_active():
            markdown, source = pick_best_local_ocr_markdown(
                image, lang=lang, mode=effective_mode
            )
            return markdown, source, None
        tesseract_md = ocr_image_to_markdown(image, lang=lang, mode=effective_mode)
        return flatten_illegal_pipe_tables(tesseract_md), "tesseract", None

    if page_kind == scan_pages.PAGE_GENERIC:
        visual_kind = scan_pages.guess_visual_page_kind(image, lang=lang)
        if visual_kind:
            page_kind = visual_kind
    is_table_page = page_kind == scan_pages.PAGE_SPREADSHEET and scan_pages.page_allows_tesseract_table(page_kind)
    vision_ok = bool(vision_llm.configured_providers()) and OCR_VISION_ENABLED
    tables_vision_ok = vision_tables_allowed()
    legal_vision_ok = vision_legal_allowed() and scan_pages.page_needs_legal_vision(
        image, lang=lang, page_kind=page_kind
    )
    docs_vision_ok = vision_docs_allowed() and scan_pages.page_needs_docs_vision(
        image, lang=lang, page_kind=page_kind
    )
    use_vision = bool(profile.get("use_vision", profile.get("use_gemini", False))) and vision_ok
    is_doc_page = page_kind in (
        scan_pages.PAGE_RECEIPT,
        scan_pages.PAGE_CREDENTIAL,
        scan_pages.PAGE_FORM,
        scan_pages.PAGE_CHECKLIST,
        scan_pages.PAGE_CV,
        scan_pages.PAGE_CONTRACT,
    )

    def _usable_table(md: str) -> bool:
        if not md:
            return False
        lines = [ln for ln in md.splitlines() if ln.strip().startswith("|")]
        if lines:
            cols = lines[0].count("|") - 1
            if cols > 10 and not table_ocr._looks_like_payroll_table(md):
                return False
        return (
            table_ocr.score_table_markdown(md) >= 28
            and table_ocr._table_content_quality(md) >= 0.35
            and text_quality_score(md) >= 20
        )

    def _usable_text(md: str) -> bool:
        if _vision_refused(md):
            return False
        plain = strip_markdown_syntax(md)
        min_chars = 12 if is_doc_page else OCR_MIN_PAGE_CHARS
        min_score = 14 if is_doc_page else 18
        return bool(md) and text_quality_score(md) >= min_score and len(plain) >= min_chars

    def _try_legal_vision() -> tuple[str, str]:
        if not legal_vision_ok and not docs_vision_ok:
            return "", ""
        try:
            if legal_vision_ok:
                pool = vision_llm.providers_for_legal()
            else:
                pool = vision_llm.providers_for_documents()
            preview_text = scan_pages._preview_text(image, lang=lang)
            vision_kind = page_kind
            if (
                scan_pages._CHECKLIST_MARKERS.search(preview_text)
                or scan_pages._looks_like_employment_checklist(preview_text)
            ):
                vision_kind = scan_pages.PAGE_CHECKLIST
            elif vision_kind == scan_pages.PAGE_GENERIC and not legal_vision_ok:
                preview_words = len(
                    re.findall(
                        r"\b[\wáéíóúÁÉÍÓÚñ]{3,}\b",
                        preview_text,
                    )
                )
                if (
                    preview_words < 25
                    and not scan_pages._looks_like_employment_checklist(preview_text)
                    and not scan_pages._CHECKLIST_MARKERS.search(preview_text)
                    and not scan_pages._looks_like_contract(preview_text)
                ):
                    vision_kind = scan_pages.PAGE_CREDENTIAL
            elif vision_kind == scan_pages.PAGE_GENERIC and scan_pages._looks_like_contract(preview_text):
                vision_kind = scan_pages.PAGE_CONTRACT

            contract_like = vision_kind in (
                scan_pages.PAGE_CONTRACT,
                scan_pages.PAGE_GENERIC,
            )
            if contract_like and "openai" in pool:
                preferred = "openai"
            else:
                preferred = "anthropic" if "anthropic" in pool else (pool[0] if pool else None)

            def _attempt(
                *,
                prompt_override: str | None = None,
                kind: str = vision_kind,
                prov: str | None = preferred,
            ) -> tuple[str, str]:
                md, source, _ = try_vision_page(
                    image,
                    page_kind=kind,
                    docs_only=True,
                    provider=prov,
                    provider_pool=pool,
                    prompt_override=prompt_override,
                )
                if _vision_refused(md):
                    return "", ""
                if _usable_text(md):
                    return md, source or "vision-legal"
                return "", ""

            prompt_chain: list[str | None] = [None]
            if contract_like:
                prompt_chain.extend([
                    CONTRACT_LITERAL_OCR_PROMPT,
                    OCR_RAW_PIXEL_PROMPT,
                    MRZ_LITERAL_PROMPT,
                ])
            elif vision_kind == scan_pages.PAGE_CREDENTIAL:
                prompt_chain.append(MRZ_LITERAL_PROMPT)

            for prompt_override in prompt_chain:
                vision_md, source = _attempt(prompt_override=prompt_override)
                if vision_md:
                    return vision_md, source, None

            alt_providers = [p for p in pool if p != preferred]
            for alt in alt_providers:
                for prompt_override in prompt_chain:
                    vision_md, source = _attempt(prompt_override=prompt_override, prov=alt)
                    if vision_md:
                        label = source or "vision-legal"
                        return vision_md, f"{label}:{alt}", None
        except Exception:
            pass
        return "", "", None

    legal_md, legal_source = _try_legal_vision()
    if legal_md:
        return legal_md, legal_source, None

    if is_table_page and tables_vision_ok:
        try:
            table_providers = vision_llm.providers_for_tables()
            preferred = "anthropic" if "anthropic" in table_providers else (table_providers[0] if table_providers else None)
            vision_md, source, _ = try_vision_page(
                image,
                table_mode=True,
                tables_only=True,
                provider=preferred,
            )
            if _usable_table(vision_md):
                return vision_md, source or "vision-table", None
        except Exception:
            pass

    if is_table_page and use_vision and vision_ok:
        try:
            vision_md, source, _ = try_vision_page(image, table_mode=True)
            if _usable_table(vision_md):
                return vision_md, source or "vision-table", None
        except Exception:
            pass

    if is_table_page:
        try:
            table_md = table_ocr.ocr_image_to_table_markdown(image, lang=lang)
            if _usable_table(table_md):
                return table_md, "tesseract-table", None
        except Exception:
            pass

    if use_vision and vision_ok:
        try:
            vision_md, source, _ = try_vision_page(image, table_mode=is_table_page)
            if vision_md:
                if is_table_page and _usable_table(vision_md):
                    return vision_md, source or "vision-table", None
                if _usable_text(vision_md):
                    return vision_md, source or "vision", None
        except Exception:
            pass

    tesseract_md = ocr_image_to_markdown(image, lang=lang, mode=effective_mode)
    tesseract_md = flatten_illegal_pipe_tables(tesseract_md)
    if is_table_page and not _usable_text(tesseract_md):
        try:
            fallback_table = table_ocr.ocr_cluster_table(image, lang=lang)
            if _usable_table(fallback_table) and table_ocr.score_table_markdown(fallback_table) > table_ocr.score_table_markdown(tesseract_md) + 8:
                return fallback_table, "tesseract-table", None
        except Exception:
            pass
    return tesseract_md, "tesseract", None

def strip_markdown_syntax(markdown: str) -> str:
    text = normalize_markdown(markdown)
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*[-*]\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*\d+\.\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*\|?(?:\s*:?-+:?\s*\|)+\s*$", "", text, flags=re.MULTILINE)
    text = text.replace("|", " ")
    text = text.replace("**", "").replace("__", "")
    text = text.replace("*", "").replace("_", "")
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)
    text = re.sub(r"\n\s*---\s*\n", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

def split_markdown_pages(markdown: str) -> list[dict[str, Any]]:
    normalized = normalize_markdown(markdown)
    if not normalized:
        return []
    chunks = [chunk.strip() for chunk in re.split(r"\n\s*---\s*\n", normalized) if chunk.strip()]
    if len(chunks) <= 1:
        return []
    return [
        {
            "index": index,
            "markdown": chunk,
            "text": strip_markdown_syntax(chunk),
        }
        for index, chunk in enumerate(chunks)
    ]

def join_pages_markdown(pages: list[dict[str, Any]]) -> str:
    return "\n\n---\n\n".join(page["markdown"].strip() for page in pages if page.get("markdown", "").strip()).strip()

def join_pages_text(pages: list[dict[str, Any]]) -> str:
    if not pages:
        return ""
    if len(pages) == 1:
        return pages[0]["text"].strip()

    blocks = []
    total = len(pages)
    for index, page in enumerate(pages, 1):
        body = page.get("text", "").strip()
        if body:
            blocks.append(f"--- Página {index}/{total} ---\n{body}")
    return "\n\n".join(blocks).strip()

def append_warning(warnings: list[str], warning: str):
    if warning and warning not in warnings:
        warnings.append(warning)


def text_quality_score(text: str) -> float:
    plain = strip_markdown_syntax(text)
    if not plain:
        return 0.0

    total = len(plain)
    useful = len(re.findall(r"[a-zA-ZáéíóúÁÉÍÓÚñÑüÜ0-9\s,\.;:\!\?\(\)\"\'\-\/%$#@&]", plain))
    useful_ratio = useful / total if total else 0.0
    words = len(re.findall(r"\b[\wáéíóúÁÉÍÓÚñÑüÜ]+\b", plain))
    lines = len([line for line in plain.split("\n") if line.strip()])

    return useful_ratio * 55 + min(total, 3000) / 60 + min(words, 400) / 10 + min(lines, 80) / 4

def normalize_table_cell(value: Any) -> str:
    cell = normalize_whitespace(str(value or ""))
    return cell.replace("\n", " ").replace("|", "/").strip()

def table_to_markdown(table: list[list[Any]]) -> str:
    rows = []
    for row in table or []:
        normalized_row = [normalize_table_cell(cell) for cell in row]
        if any(cell for cell in normalized_row):
            rows.append(normalized_row)

    if len(rows) < 2:
        return ""

    col_count = max(len(row) for row in rows)
    padded = [row + [""] * (col_count - len(row)) for row in rows]
    header = padded[0]
    if not any(header):
        header = [f"Columna {idx + 1}" for idx in range(col_count)]

    separator = ["---"] * col_count
    body = padded[1:]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    lines.extend("| " + " | ".join(row) + " |" for row in body)
    return "\n".join(lines)

def extract_pdf_tables_markdown(file_path: Path) -> dict[int, list[str]]:
    table_map: dict[int, list[str]] = {}
    strategies = (
        {
            "vertical_strategy": "lines",
            "horizontal_strategy": "lines",
            "snap_tolerance": 3,
            "intersection_tolerance": 3,
        },
        {
            "vertical_strategy": "lines_strict",
            "horizontal_strategy": "lines_strict",
            "snap_tolerance": 2,
            "intersection_tolerance": 2,
        },
        {
            "vertical_strategy": "text",
            "horizontal_strategy": "text",
            "snap_tolerance": 4,
            "join_tolerance": 4,
        },
    )
    try:
        with pdfplumber.open(str(file_path)) as pdf:
            for page_index, page in enumerate(pdf.pages):
                best_tables: list[str] = []
                best_score = 0.0
                for settings in strategies:
                    extracted = page.extract_tables(settings) or []
                    page_tables = [table_to_markdown(table) for table in extracted]
                    page_tables = [table for table in page_tables if table]
                    for candidate in page_tables:
                        score = table_ocr.score_table_markdown(candidate)
                        if score > best_score:
                            best_score = score
                            best_tables = [candidate]
                        elif score == best_score and score > 0:
                            best_tables.append(candidate)
                if best_tables:
                    table_map[page_index] = best_tables
    except Exception:
        return {}
    return table_map


def iter_pdf_page_chunks(pdf_path: Path, dpi: int, chunk_size: int = OCR_CHUNK_PAGES):
    doc = fitz.open(str(pdf_path))
    try:
        total = len(doc)
        for chunk_start in range(0, total, chunk_size):
            chunk_end = min(chunk_start + chunk_size, total)
            batch: list[tuple[int, Image.Image]] = []
            for page_index in range(chunk_start, chunk_end):
                page = doc.load_page(page_index)
                image = prepare_page_image_for_ocr(page, dpi)
                batch.append((page_index, image))
            yield batch
    finally:
        doc.close()


def flatten_illegal_pipe_tables(markdown: str, *, allow_payroll: bool = False) -> str:
    """Convierte rejillas OCR basura en texto; conserva tablas pequeñas legibles (acta, recibo)."""
    if allow_payroll or not markdown or "|" not in markdown:
        return markdown
    lines = markdown.splitlines()
    out: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip().startswith("|"):
            out.append(line)
            i += 1
            continue
        block: list[str] = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            block.append(lines[i])
            i += 1
        block_text = "\n".join(block)
        cols = block[0].count("|") - 1 if block else 0
        keep_table = False
        if table_ocr._looks_like_payroll_table(block_text) and cols <= 10:
            keep_table = True
        elif cols <= 5 and table_ocr._table_content_quality(block_text) >= 0.48:
            keep_table = True
        elif cols <= 4 and table_ocr.score_table_markdown(block_text) >= 35:
            keep_table = True
        if keep_table:
            out.extend(block)
            out.append("")
            continue
        for row in block:
            if "---" in row:
                continue
            cells = [c.strip() for c in row.strip().strip("|").split("|") if c.strip()]
            if cells:
                out.append(" ".join(cells))
        out.append("")
    return re.sub(r"\n{3,}", "\n\n", "\n".join(out)).strip()


def _markdown_has_garbage_pipe_grid(markdown: str) -> bool:
    for line in markdown.splitlines():
        if line.strip().startswith("|") and line.count("|") > 12:
            return True
    return False


def image_to_png_bytes(image: Image.Image) -> bytes:
    buffer = BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()

def build_result(
    method: str,
    markdown: str,
    *,
    pages: Optional[list[dict[str, Any]]] = None,
    file_path: Optional[Path] = None,
    warnings: Optional[list[str]] = None,
    sanitize_aggressive: Optional[bool] = None,
) -> dict[str, Any]:
    if sanitize_aggressive is None:
        sanitize_aggressive = method not in ("digital_pdf", "markitdown")

    def _sanitize_page_markdown(raw: str, *, digital: bool = False, ocr_sourced: bool = False) -> str:
        aggressive = sanitize_aggressive and not digital and not ocr_sourced
        cleaned = sanitize_document_markdown(
            normalize_markdown(raw),
            aggressive=aggressive,
        )
        return flatten_illegal_pipe_tables(cleaned)

    normalized_pages = []
    for index, page in enumerate(pages or []):
        page_markdown = _sanitize_page_markdown(
            page.get("markdown") or page.get("text") or "",
            digital=bool(page.get("digital")),
            ocr_sourced=bool(page.get("ocr_sourced")),
        )
        if not page_markdown:
            continue
        page_text = strip_markdown_syntax(page.get("text") or page_markdown)
        entry: dict[str, Any] = {
            "index": page.get("index", index),
            "markdown": page_markdown,
            "text": page_text,
        }
        if "digital" in page:
            entry["digital"] = bool(page.get("digital"))
        if "ocr_sourced" in page:
            entry["ocr_sourced"] = bool(page.get("ocr_sourced"))
        if page.get("ocr_source"):
            entry["ocr_source"] = page.get("ocr_source")
        if page.get("blocks"):
            entry["blocks"] = page.get("blocks")
        if page.get("image_size"):
            entry["image_size"] = page.get("image_size")
        normalized_pages.append(entry)

    normalized_markdown = _sanitize_page_markdown(markdown) if markdown else ""
    if not normalized_pages and normalized_markdown:
        normalized_pages = split_markdown_pages(normalized_markdown)

    if normalized_pages:
        normalized_markdown = join_pages_markdown(normalized_pages)
    elif not normalized_markdown and markdown:
        normalized_markdown = _sanitize_page_markdown(markdown)

    plain_text = join_pages_text(normalized_pages) if normalized_pages else strip_markdown_syntax(normalized_markdown)

    total_pages = len(normalized_pages)
    if total_pages == 0:
        if file_path and file_path.suffix.lower() == ".pdf":
            total_pages = count_pdf_pages(file_path)
        else:
            total_pages = 1 if normalized_markdown or plain_text else 0

    if not normalized_pages and normalized_markdown:
        normalized_pages = [{
            "index": 0,
            "markdown": normalized_markdown,
            "text": plain_text,
        }]

    result_warnings = list(warnings or [])
    if file_path and file_path.exists():
        size_mb = file_path.stat().st_size / (1024 * 1024)
        if size_mb >= SOFT_SIZE_WARNING_MB:
            append_warning(
                result_warnings,
                f"Archivo pesado ({size_mb:.1f} MB). Si el resultado sale incompleto, intenta dividir el documento."
            )
        if file_path.suffix.lower() == ".pdf" and total_pages > SOFT_PAGE_WARNING:
            append_warning(
                result_warnings,
                f"PDF largo ({total_pages} páginas). Revisa el resultado por secciones para validar tablas y texto."
            )

    return {
        "method": method,
        "text": plain_text,
        "markdown": normalized_markdown,
        "pages": normalized_pages,
        "total_pages": total_pages,
        "warnings": result_warnings,
    }

def try_digital_pdf(file_path: Path) -> Optional[dict[str, Any]]:
    try:
        table_map = extract_pdf_tables_markdown(file_path)
        pages = []
        total_chars = 0
        text_pages = 0
        table_count = sum(len(tables) for tables in table_map.values())

        with fitz.open(str(file_path)) as doc:
            for page_index, page in enumerate(doc):
                raw_page = page.get_text("text", sort=True) or ""
                text_block = extract_digital_page_markdown(page)
                page_parts = []

                if len(strip_markdown_syntax(text_block)) >= OCR_MIN_PAGE_CHARS:
                    page_parts.append(text_block)
                    text_pages += 1

                page_tables = table_map.get(page_index, [])
                if page_tables and not is_cryptographic_evidence_page(raw_page):
                    if page_parts:
                        page_parts.append("## Tablas detectadas")
                    page_parts.extend(page_tables)

                page_markdown = "\n\n".join(part for part in page_parts if part.strip()).strip()
                if page_markdown:
                    pages.append({
                        "index": page_index,
                        "markdown": page_markdown,
                        "digital": True,
                    })
                    total_chars += len(strip_markdown_syntax(page_markdown))

        if total_chars < minimum_pdf_text_chars(file_path) or text_pages == 0:
            return None

        warnings = []
        if table_count:
            append_warning(warnings, f"Se detectaron {table_count} tablas en el PDF digital.")

        return build_result(
            "digital_pdf",
            join_pages_markdown(pages),
            pages=pages,
            file_path=file_path,
            warnings=warnings,
        )
    except Exception:
        return None


def page_image_coverage(page) -> float:
    """Porción del área de la página cubierta por imágenes embebidas (0–1)."""
    page_area = max(float(page.rect.width * page.rect.height), 1.0)
    covered = 0.0
    seen: set[int] = set()
    for img in page.get_images(full=True):
        xref = int(img[0])
        if xref in seen:
            continue
        seen.add(xref)
        try:
            for rect in page.get_image_rects(xref):
                covered += abs(rect)
        except Exception:
            continue
    return min(1.0, covered / page_area)


def _page_text_block_count(page) -> int:
    blocks = page.get_text("blocks") or []
    return sum(
        1 for b in blocks
        if len(b) >= 7 and int(b[6]) == 0 and normalize_whitespace(str(b[4]))
    )


def page_is_scan_dominant(page) -> bool:
    """Página escaneada o foto a pantalla completa (imagen > texto útil)."""
    coverage = page_image_coverage(page)
    raw = normalize_whitespace(page.get_text("text", sort=True) or "")
    text_blocks = _page_text_block_count(page)
    if coverage >= 0.52:
        return True
    if coverage >= 0.38 and text_blocks <= 3 and len(raw) < 250:
        return True
    if coverage >= 0.22 and len(raw) < 36:
        return True
    return False


def page_has_reliable_digital_text(page, digital_md: str) -> bool:
    """Texto nativo del PDF (Word/exportación), no capa OCR invisible sobre un scan."""
    if page_is_scan_dominant(page):
        return False
    plain = strip_markdown_syntax(digital_md or "")
    raw = normalize_whitespace(page.get_text("text", sort=True) or "")
    if len(plain) < 72 or len(raw) < 48:
        return False
    text_blocks = _page_text_block_count(page)
    if text_blocks >= 3 and _spanish_word_ratio(raw) >= 0.28:
        return True
    if len(plain) >= 280 and _spanish_word_ratio(raw) >= 0.22:
        return True
    return False


def classify_pdf_page(page, digital_md: str) -> str:
    """
    digital — solo extraer texto seleccionable.
    ocr     — renderizar página y OCR (escaneos / fotos).
    both    — texto digital + OCR para fotos embebidas grandes.
    """
    if page_is_scan_dominant(page):
        return "ocr"
    if page_has_reliable_digital_text(page, digital_md):
        if page_image_coverage(page) >= 0.28:
            return "both"
        return "digital"
    return "ocr"


def _hybrid_ocr_scanned_pages_chunk(
    file_path: Path,
    page_indices: list[int],
    *,
    dpi: int,
    lang: str,
    mode: str,
) -> list[tuple[int, str, str, str, Optional[dict[str, Any]]]]:
    """Procesa un bloque de páginas con un solo open del PDF (mucho más rápido)."""
    rows: list[tuple[int, str, str, str, Optional[dict[str, Any]]]] = []
    with fitz.open(str(file_path)) as doc:
        for page_index in page_indices:
            check_job_control_or_raise()
            page = doc.load_page(page_index)
            digital_md = extract_digital_page_markdown(page)
            image = prepare_page_image_for_ocr(page, dpi, lang=lang)
            try:
                ocr_md, source, layout = ocr_page_with_fallback(image, lang=lang, mode=mode)
            finally:
                release_image(image)
            rows.append((page_index, digital_md, ocr_md, source, layout))
    return rows


def try_hybrid_pdf(file_path: Path, lang: str = "spa", *, mode: str = "balanced") -> Optional[dict[str, Any]]:
    """Por página: texto digital si existe; si no, Gemini/OCR con formato legible."""
    mode = normalize_ocr_mode(mode)
    with fitz.open(str(file_path)) as doc:
        total_pages = len(doc)
    profile = get_ocr_profile() or build_ocr_profile(mode, total_pages)
    dpi = int(profile.get("dpi", ocr_dpi_for_mode(mode)))
    try:
        table_map = extract_pdf_tables_markdown(file_path)
        pages: list[dict[str, Any]] = []
        warnings: list[str] = []
        digital_pages = 0
        vision_pages = 0
        ocr_pages = 0
        table_count = sum(len(tables) for tables in table_map.values())

        if profile.get("bulk"):
            append_warning(
                warnings,
                f"PDF grande ({total_pages} páginas): modo rápido automático "
                f"(DPI {dpi}, sin pasada extra de sellos"
                + (", sin IA por página" if not profile.get("use_vision") else "")
                + "). Para máxima calidad en lotes pequeños usa menos de "
                f"{OCR_LARGE_PDF_PAGES} páginas o modo Máxima calidad en trozos.",
            )

        with fitz.open(str(file_path)) as doc:
            if total_pages > OCR_CHUNK_PAGES:
                append_warning(
                    warnings,
                    f"PDF largo procesado por bloques ({total_pages} páginas)."
                )

            page_plan: list[tuple[int, str, str]] = []
            for page_index in range(total_pages):
                page = doc.load_page(page_index)
                digital_md = extract_digital_page_markdown(page)
                page_mode = classify_pdf_page(page, digital_md)
                page_plan.append((page_index, page_mode, digital_md))

        ocr_indices = [idx for idx, mode, _ in page_plan if mode in ("ocr", "both")]
        ocr_results: dict[int, tuple[str, str, str, Optional[dict[str, Any]]]] = {}

        if ocr_indices:
            max_workers = (
                OCR_GEMINI_WORKERS
                if profile.get("use_vision")
                else OCR_TESSERACT_WORKERS
            )
            if paddle_ocr_active():
                max_workers = 1
            max_workers = max(1, min(max_workers, len(ocr_indices)))
            chunk_size = max(4, min(16, (len(ocr_indices) + max_workers - 1) // max_workers))
            chunks = [
                ocr_indices[i : i + chunk_size]
                for i in range(0, len(ocr_indices), chunk_size)
            ]

            def chunk_job(indices: list[int]) -> list[tuple[int, str, str, str]]:
                return _hybrid_ocr_scanned_pages_chunk(
                    file_path, indices, dpi=dpi, lang=lang, mode=mode
                )

            done_count = 0
            if hybrid_ocr_parallel_enabled() and len(chunks) > 1:
                with ThreadPoolExecutor(max_workers=min(max_workers, len(chunks))) as executor:
                    for batch in executor.map(chunk_job, chunks):
                        for page_index, digital_md, ocr_md, source, layout in batch:
                            ocr_results[page_index] = (digital_md, ocr_md, source, layout)
                            done_count += 1
                            report_ocr_progress(
                                done_count,
                                len(ocr_indices),
                                f"Parsing ({done_count} / {len(ocr_indices)} pages)",
                            )
            else:
                for indices in chunks:
                    for page_index, digital_md, ocr_md, source, layout in chunk_job(indices):
                        ocr_results[page_index] = (digital_md, ocr_md, source, layout)
                        done_count += 1
                        report_ocr_progress(
                            done_count,
                            len(ocr_indices),
                            f"Parsing ({done_count} / {len(ocr_indices)} pages)",
                        )

        gc_interval = max(4, OCR_CHUNK_PAGES)
        for page_index, page_mode, digital_md in page_plan:
            check_job_control_or_raise()
            page_parts: list[str] = []
            source = ""
            layout: Optional[dict[str, Any]] = None

            if page_mode == "digital":
                page_parts.append(digital_md)
                digital_pages += 1
            else:
                dig_partial, ocr_md, source, layout = ocr_results.get(
                    page_index, (digital_md, "", "", None)
                )
                if page_mode == "ocr":
                    if ocr_md and len(strip_markdown_syntax(ocr_md)) >= OCR_MIN_PAGE_CHARS:
                        page_parts.append(ocr_md)
                    elif dig_partial:
                        page_parts.append(dig_partial)
                else:
                    if digital_md:
                        page_parts.append(digital_md)
                        digital_pages += 1
                    if ocr_md and len(strip_markdown_syntax(ocr_md)) >= OCR_MIN_PAGE_CHARS:
                        page_parts.append(ocr_md)
                if _is_vision_ocr_source(source):
                    vision_pages += 1
                elif ocr_md:
                    ocr_pages += 1

            report_ocr_progress(
                page_index + 1,
                total_pages,
                f"Parsing ({page_index + 1} / {total_pages} pages)",
            )
            gc_every_n_pages(page_index, gc_interval)

            page_tables = table_map.get(page_index, [])
            page_markdown = "\n\n".join(part for part in page_parts if part.strip()).strip()
            if page_tables and not is_cryptographic_evidence_page(page_markdown):
                existing_score = table_ocr.score_table_markdown(page_markdown)
                extra_tables = [
                    tbl for tbl in page_tables
                    if table_ocr.score_table_markdown(tbl) >= 40
                    and table_ocr._looks_like_payroll_table(tbl)
                ]
                if extra_tables and existing_score < 35:
                    if page_markdown:
                        page_parts.append("## Tablas detectadas")
                    page_parts.extend(extra_tables)
                    page_markdown = "\n\n".join(part for part in page_parts if part.strip()).strip()
            if page_markdown:
                page_entry: dict[str, Any] = {
                    "index": page_index,
                    "markdown": page_markdown,
                    "digital": page_mode == "digital",
                    "ocr_sourced": page_mode in ("ocr", "both"),
                    "ocr_source": source if page_mode != "digital" else "",
                }
                if layout and layout.get("blocks"):
                    page_entry["blocks"] = layout.get("blocks")
                    page_entry["image_size"] = layout.get("image_size")
                pages.append(page_entry)

        if not pages:
            return None

        total_chars = sum(len(strip_markdown_syntax(page["markdown"])) for page in pages)
        if total_chars < minimum_pdf_text_chars(file_path):
            return None

        if table_count:
            append_warning(warnings, f"Se detectaron {table_count} tablas en el documento.")

        if any(
            _markdown_has_garbage_pipe_grid(p.get("markdown") or "")
            for p in pages
        ):
            append_warning(
                warnings,
                "Se detectó una tabla OCR ilegible. Reinicia el servidor OCR y vuelve a procesar el documento.",
            )
        if ocr_pages and (vision_legal_allowed() or vision_docs_allowed()) and vision_pages == 0:
            append_warning(
                warnings,
                "Algunas páginas usaron solo OCR local (sin Claude). Reinicia el servidor y verifica "
                "ANTHROPIC_API_KEY, OCR_VISION_LEGAL_ONLY=1 y OCR_VISION_DOCS_ONLY=1 en backend/.env.",
            )

        if digital_pages == len(pages):
            method = "digital_pdf"
        elif vision_pages > 0:
            method = "vision"
        elif digital_pages > 0:
            method = "digital_pdf"
        else:
            method = "structured"

        return build_result(
            method,
            join_pages_markdown(pages),
            pages=pages,
            file_path=file_path,
            warnings=warnings,
        )
    except Exception:
        return None

# ── Markitdown ────────────────────────────────────────────────────────────────
def try_markitdown(file_path: Path) -> Optional[dict[str, Any]]:
    try:
        md = MarkItDown()
        result = md.convert(str(file_path))
        text = result.text_content.strip()
        if len(text) >= minimum_pdf_text_chars(file_path):
            return build_result(
                "markitdown",
                ensure_readable_markdown(text),
                file_path=file_path,
            )
        return None
    except Exception:
        return None

# ── Gemini OCR ────────────────────────────────────────────────────────────────
_LITERAL_OCR_RULES = (
    "\n\nMODO OCR ESTRICTO — NO eres asistente legal:\n"
    "- Transcribe únicamente el texto impreso visible, tal cual, en orden de lectura.\n"
    "- Prohibido analizar, resumir, interpretar cláusulas o explicar el documento.\n"
    "- Prohibido decir qué tipo de documento es ni ofrecer revisiones o análisis.\n"
    "- Prohibido texto tuyo antes o después (nada de 'Aquí está...', 'Este es un contrato...').\n"
    "Devuelve ÚNICAMENTE la transcripción en Markdown."
)


def _literal_ocr_prompt(body: str) -> str:
    return body.strip() + _LITERAL_OCR_RULES


LITERAL_PAGE_OCR_PROMPT = _literal_ocr_prompt(
    "Transcribe TODO el texto impreso visible en esta imagen, tal cual aparece, "
    "línea por línea y en el orden de lectura.\n\n"
    "Usa Markdown mínimo: # o ## solo si el documento ya tiene títulos claros; "
    "tablas Markdown solo donde haya cuadros visibles; listas y numeración como en el original.\n"
    "Ignora sellos circulares, marcas de agua tenues y firmas manuscritas ilegibles."
)

GEMINI_MARKDOWN_PROMPT = _literal_ocr_prompt(
    "Transcribe a Markdown únicamente el TEXTO ÚTIL del documento (impreso o tipográfico).\n\n"
    "IGNORA por completo (no transcribas, no menciones, no resumas):\n"
    "- Sellos circulares o rectangulares, timbres, estampillas PAGADO, sellos de notaría, "
    "juzgado, empresa o cualquier sello superpuesto.\n"
    "- Marcas de agua, escudos tenues de fondo, logos semitransparentes.\n"
    "- Firmas manuscritas, rubricas, iniciales, anotaciones a pluma o lápiz.\n"
    "- Texto vertical en márgenes (ej. COTEJADO), resaltado amarillo, marcas de revisión.\n"
    "- Sellos digitales CFDI/SAT: cadenas base64 largas, «Sello digital del CFDI», "
    "«Sello del SAT», «Cadena Original del Complemento», números de certificado.\n"
    "- Pie de página legal repetitivo de facturas electrónicas (ej. «representación impresa de un CFDI»).\n\n"
    "Fotografías e imágenes:\n"
    "- Foto genérica (retrato, escena, producto, etc.): pon solo la línea `[Fotografía]` "
    "en el lugar donde aparece en el flujo de lectura.\n"
    "Sí transcribe con fidelidad:\n"
    "- Títulos, párrafos, listas, numeración legal, formularios.\n"
    "- Tablas como tablas Markdown válidas.\n"
    "- Casillas marcadas como [x] o [ ].\n"
    "- Fechas y montos impresos en el cuerpo del documento (no los del sello)."
)

GEMINI_TABLE_PROMPT = _literal_ocr_prompt(
    "Transcribe esta página como una tabla Markdown que replique fielmente el documento.\n\n"
    "Reglas obligatorias:\n"
    "- Detecta todas las columnas y filas visibles (encabezados incluidos).\n"
    "- Usa una tabla Markdown válida con fila separadora | --- |.\n"
    "- Conserva el orden exacto: FECHA, NOMBRE, CONCEPTO, PAGO, IMPORTE, etc.\n"
    "- Mantén fechas, montos ($), comas decimales y celdas vacías.\n"
    "- Una fila del PDF = una fila de la tabla; no mezcles columnas.\n"
    "- No resumas ni omitas filas legibles.\n"
    "- Ignora sellos, marcas de agua y firmas manuscritas."
)

RECEIPT_VISION_PROMPT = _literal_ocr_prompt(
    "Transcribe al pie de la letra este recibo o comprobante de pago (luz, agua, teléfono, predial, etc.).\n\n"
    "Formato obligatorio (como documento escaneado legible):\n"
    "- Título principal con # y el nombre de la empresa (ej. # TELMEX, # CFE).\n"
    "- Dirección, RFC, datos del cliente y domicilio como párrafos o líneas sueltas, en el orden visual.\n"
    "- Montos y resúmenes en tablas Markdown cuando el recibo los muestre en columnas o cuadros:\n"
    "  | Total a Pagar: | $ 1,451.00 |\n"
    "  | --- | --- |\n"
    "- Secciones visibles como ## (ej. ## AVISO IMPORTANTE, ## Resumen del Estado de Cuenta).\n"
    "- Usa tablas Markdown solo para montos (2–3 columnas máximo), nunca para todo el recibo.\n"
    "- Conserva números, fechas, símbolos $ y decimales exactamente.\n"
    "- Ignora marcas de agua (CamScanner), códigos de barras y publicidad de pago.\n"
    "No uses listas con viñetas para sustituir tablas; no inventes datos."
)

FORM_DOCUMENT_VISION_PROMPT = _literal_ocr_prompt(
    "Transcribe al pie de la letra este documento oficial mexicano (acta de nacimiento, CURP, constancia, etc.).\n\n"
    "Conserva títulos, etiquetas en MAYÚSCULAS, valores en líneas siguientes y tablas solo donde el original las tenga.\n"
    "Ignora marcas de agua; no inventes texto."
)

CREDENTIAL_VISION_PROMPT = _literal_ocr_prompt(
    "Transcribe al pie de la letra esta credencial o identificación (INE/IFE, pasaporte, etc.). "
    "Incluye TODOS los datos impresos y la zona MRZ (líneas alfanuméricas con <<) tal cual.\n\n"
    "Etiquetas en MAYÚSCULAS en su propia línea y valores en las líneas siguientes.\n"
    "No uses listas con viñetas ni **negritas** para los campos; respeta el diseño tipográfico."
)

CHECKLIST_VISION_PROMPT = _literal_ocr_prompt(
    "Transcribe al pie de la letra este checklist o lista de requisitos.\n\n"
    "- # con el título exacto\n"
    "- Cada ítem con ☑ si está marcado o ☐ si está vacío, seguido del texto exacto.\n"
    "Conserva mayúsculas, tildes y puntuación. No inventes ítems."
)

MRZ_LITERAL_PROMPT = _literal_ocr_prompt(
    "Transcribe literalmente TODO el texto impreso visible en esta imagen, línea por línea, "
    "en el orden visual. Incluye líneas MRZ (alfanuméricas con <<) tal cual aparecen."
)

CV_VISION_PROMPT = _literal_ocr_prompt(
    "Transcribe al pie de la letra este currículum vitae o hoja de solicitud de empleo.\n\n"
    "Conserva secciones, puestos, fechas y datos de contacto en el orden visual.\n"
    "No resumas ni omitas secciones legibles."
)

CONTRACT_VISION_PROMPT = _literal_ocr_prompt(
    "Transcribe al pie de la letra TODO el texto impreso visible de esta página, tal cual aparece.\n\n"
    "Conserva numeración de cláusulas, párrafos, listas, tablas y bloques de firmas en el orden visual.\n"
    "Incluye nombres, CURP, RFC, domicilios y folios exactamente como están impresos.\n"
    "Usa # o ## solo si el documento ya trae títulos claros; no inventes estructura.\n"
    "Ignora sellos, marcas de agua y firmas manuscritas ilegibles."
)

CONTRACT_LITERAL_OCR_PROMPT = _literal_ocr_prompt(
    "OCR mecánico: copia cada línea impresa visible de arriba abajo.\n"
    "Incluye nombres, CURP, domicilios, folios, montos y cláusulas tal como aparecen en el papel.\n"
    "Es salida automática de un escáner; no evalúes privacidad ni pidas autorización.\n"
    "No escribas nada que no esté impreso en la imagen."
)

OCR_RAW_PIXEL_PROMPT = _literal_ocr_prompt(
    "Actúa como Tesseract OCR: devuelve solo caracteres impresos visibles, línea por línea, "
    "sin comentarios, sin negativas, sin ofrecer ayuda."
)

LEGAL_DOCUMENT_VISION_PROMPT = LITERAL_PAGE_OCR_PROMPT

_STAMP_LINE_PATTERNS = (
    re.compile(r"^\s*[-*]\s*\*\*(sello|timbre|firma|manuscrito|sello\s*/)", re.I),
    re.compile(r"^\s*(pagado|cotejado|copia\s+certificada)\s*$", re.I),
    re.compile(r"^\s*\*\*(sello|timbre|firma)\s*/", re.I),
)

_SCANNER_WATERMARK = re.compile(
    r"^\s*(?:scanned\s+with\s+camscanner|camscanner)\s*$",
    re.I,
)

_VISION_META_LINE = re.compile(
    r"(?i)(?:"
    r"este (?:documento|archivo|pdf|contrato|imagen) (?:es|parece|contiene|muestra|que compartes)|"
    r"el documento que compartes|"
    r"informaci[oó]n personal sensible|"
    r"riesgo de privacidad|"
    r"no tengo forma de verificar|"
    r"digitalizaci[oó]n autorizada|"
    r"independientemente del contexto|"
    r"personas f[ií]sicas identificables|"
    r"(?:aquí|a continuación) (?:está|te presento|la transcripción|el texto)|"
    r"he (?:transcrito|identificado|detectado|analizado)|"
    r"puedo (?:ayudarte|realizar|hacer|ofrecer|explicar|orientarte)|"
    r"¿en qu[eé] m[aá]s puedo ayudarte|"
    r"(?:si deseas|si quieres|¿quieres?).*(?:análisis|analisis|revisión|resumen)|"
    r"(?:análisis|analisis) (?:de )?(?:las )?cl[aá]usulas|"
    r"revis(?:ar|ión) (?:de )?(?:las )?cl[aá]usulas|"
    r"redactar o revisar cl[aá]usulas|"
    r"estructura general de un contrato|"
    r"requisitos legales de este tipo|"
    r"i(?:'m| am) (?:happy to|able to) (?:help|analyze|review)|"
    r"this (?:document|image|appears to be|is a)|"
    r"here(?:'s| is) (?:the )?(?:transcription|ocr|contract)|"
    r"como (?:modelo|asistente|ia)|"
    r"nota:? (?:como|que) (?:modelo|ia)|"
    r"transcripci[oó]n del documento|"
    r"transcribir o reproducir esos datos"
    r")"
)

_VISION_REFUSAL = re.compile(
    r"(?i)(not able to transcribe|cannot transcribe|can't transcribe|"
    r"privacy and security|identity theft|facilitate fraud|"
    r"no puedo transcribir|no estoy en capacidad|por razones de privacidad|"
    r"informaci[oó]n personal sensible|riesgo de privacidad|"
    r"no tengo forma de verificar|personas f[ií]sicas identificables|"
    r"transcribir o reproducir esos datos|"
    r"recomiendo usar.*canales oficiales|"
    r"puedo (?:ayudarte|ofrecerte|explicar|orientarte)|"
    r"¿en qu[eé] m[aá]s puedo ayudarte|"
    r"(?:análisis|analisis) (?:legal|de cl[aá]usulas)|"
    r"redactar o revisar cl[aá]usulas|"
    r"estructura general de un contrato)"
)


def strip_vision_meta_output(text: str) -> str:
    """Quita preámbulos y cierres del modelo que no son transcripción."""
    if not text:
        return ""
    text = text.replace("\r\n", "\n")
    for pivot in (
        r"(?i)\n\s*¿en qu[eé] m[aá]s puedo ayudarte.*",
        r"(?i)\n\s*puedo explicar la estructura.*",
        r"(?i)\n\s*puedo ayudarte a redactar.*",
        r"(?i)\n\s*puedo orientarte sobre.*",
    ):
        text = re.sub(pivot, "", text, flags=re.DOTALL).strip()
    if _vision_refused(text) and len(strip_markdown_syntax(text)) < 80:
        return ""
    lines = text.split("\n")
    start = 0
    for index, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue
        if _VISION_META_LINE.search(stripped) or _VISION_REFUSAL.search(stripped):
            start = index + 1
            continue
        if re.match(r"^(?:---+|\*\*\*|```)", stripped):
            start = index + 1
            continue
        break
    end = len(lines)
    for index in range(len(lines) - 1, start - 1, -1):
        stripped = lines[index].strip()
        if not stripped:
            end = index
            continue
        if _VISION_META_LINE.search(stripped):
            end = index
            continue
        break
    return "\n".join(lines[start:end]).strip()


def _vision_refused(text: str) -> bool:
    if not text:
        return False
    if _VISION_REFUSAL.search(text):
        return True
    lower = text.lower()
    refusal_markers = (
        "información personal sensible",
        "informacion personal sensible",
        "riesgo de privacidad",
        "no tengo forma de verificar",
        "¿en qué más puedo ayudarte",
        "en que más puedo ayudarte",
        "puedo explicar la estructura",
        "puedo ayudarte a redactar",
        "puedo orientarte sobre",
    )
    if any(marker in lower for marker in refusal_markers):
        return True
    plain = strip_markdown_syntax(text)
    if len(plain) < 120:
        meta_lines = sum(1 for line in text.splitlines() if _VISION_META_LINE.search(line.strip()))
        if meta_lines and meta_lines >= max(1, len(text.splitlines()) // 2):
            return True
    return False

_CFDI_BOILERPLATE = re.compile(
    r"representaci[oó]n impresa de un cfdi|persona moral grandes contribuyentes|"
    r"^hoja\s*\d+\s*/\s*\d+\s*$",
    re.I,
)

_CFDI_SEAL_LINE = re.compile(
    r"sello digital del cfdi|sello digital del sat|sello del sat|cadena original del complemento|"
    r"n[uú]mero de serie del certificado|certificado de sello digital|"
    r"cadena original del timbre|complemento de certificaci[oó]n digital",
    re.I,
)

_BASE64_BLOB = re.compile(
    r"(?:[A-Za-z0-9+/=]{48,}|"
    r"(?:[A-Za-z0-9+/]{5,}\s+){3,}[A-Za-z0-9+/=]{5,})"
)


def _remove_base64_blobs(text: str) -> str:
    def _replace(match: re.Match[str]) -> str:
        chunk = match.group(0)
        compact = re.sub(r"\s+", "", chunk)
        if len(compact) < 36:
            return chunk
        if any(c in chunk for c in "+/=") or _base64_char_ratio(compact) > 0.88:
            return " "
        return chunk

    return _BASE64_BLOB.sub(_replace, text)


def _replace_residue_base64(match: re.Match[str]) -> str:
    chunk = match.group(0)
    if any(c in chunk for c in "+/=") or _base64_char_ratio(chunk) > 0.9:
        return " "
    return chunk
_CERT_SERIAL = re.compile(r"\b\d{17,20}\b")
_CFDI_UUID = re.compile(
    r"\b[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}\b",
    re.I,
)
_OBJ_IMPUESTO_REPEAT = re.compile(
    r"(?:objeto de (?:i|í|l|1)mpuesto:?\s*0?2\s*(?:s[iííl1]\s*objeto de impuesto\s*)+)+",
    re.I,
)


def _base64_char_ratio(text: str) -> float:
    if not text:
        return 0.0
    b64 = sum(1 for c in text if c in "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789+/=")
    return b64 / len(text)


def strip_inline_cfdi_noise(line: str) -> str:
    """Quita sellos digitales CFDI/SAT incrustados en una línea con texto útil."""
    text = line
    for pat in (
        r"Sello digital del CFDI:?.*",
        r"Sello digital del SAT:?.*",
        r"Sello del SAT:?.*",
        r"Cadena [Oo]riginal del Complemento de Certificaci[oó]n Digital del SAT.*",
        r"N[uú]mero de serie del Certificado(?: de Sello Digital)?(?: del SAT)?:?.*",
        r"Certificado de Sello Digital:?.*",
        r"Complemento de Certificaci[oó]n Digital del SAT.*",
        r"Este documento es una representaci[oó]n impresa de un CFDI.*",
        r"Persona Moral Grandes Contribuyentes.*",
        r"HOJA\s+\d+\s*/\s*\d+.*",
    ):
        text = re.sub(pat, "", text, flags=re.I)
    text = _OBJ_IMPUESTO_REPEAT.sub(" ", text)
    text = _CFDI_UUID.sub(" ", text)
    text = _CERT_SERIAL.sub(" ", text)
    text = _remove_base64_blobs(text)
    text = re.sub(r"[A-Za-z0-9+/=]{32,}", _replace_residue_base64, text)
    text = re.sub(r"\[\s*m\d?\s*\]", "", text, flags=re.I)
    text = re.sub(r"\|{2,}", " ", text)
    text = re.sub(r"\s{2,}", " ", text).strip(" |.,;:-'\"")
    return text.strip()


def _meaningful_words(text: str) -> list[str]:
    out: list[str] = []
    for t in re.findall(r"\S+", text):
        token = t.strip(".,;:'\"")
        if len(token) < 4:
            continue
        if not re.fullmatch(r"[a-záéíóúñA-ZÁÉÍÓÚÑ0-9]+", token):
            continue
        if not re.search(r"[aeiouáéíóú]", token, re.I):
            continue
        if re.search(r"[a-záéíóúñ]", token) or token.upper() == token:
            out.append(token)
    return out


def _looks_like_seal_token_line(line: str) -> bool:
    tokens = re.findall(r"\S+", line)
    if not tokens:
        return False
    sealish = 0
    for token in tokens:
        t = token.strip("'\".,;")
        if len(t) < 6:
            continue
        letters = sum(1 for c in t if c.isalpha())
        vowels = sum(1 for c in t.lower() if c in "aeiouáéíóú")
        digits = sum(1 for c in t if c.isdigit())
        if letters >= 4 and digits >= 1 and vowels / max(letters, 1) < 0.22:
            sealish += 1
            continue
        if len(t) >= 6 and t.upper() == t and vowels / max(len(t), 1) < 0.12:
            sealish += 1
            continue
        if re.fullmatch(r"[A-Z0-9+/=|<>)($'\"\.:\-\[\]]{8,}", t) and vowels < 2:
            sealish += 1
    if sealish >= 1 and _spanish_word_ratio(line) < 0.42:
        return True
    if ">" in line and len(line) < 55 and _spanish_word_ratio(line) < 0.55:
        if not re.search(r"(?:RFC|TOTAL|IMPORTE|\$|\d{3,}|@)", line, re.I):
            return True
    return False


def is_cfdi_seal_line(line: str) -> bool:
    plain = line.strip()
    if not plain:
        return False
    if _CFDI_SEAL_LINE.search(plain):
        remainder = strip_inline_cfdi_noise(plain)
        if len(remainder) < 28 or _base64_char_ratio(remainder) > 0.45:
            return True
    if len(plain) > 350 and _base64_char_ratio(plain) > 0.38:
        remainder = strip_inline_cfdi_noise(plain)
        if len(remainder) < max(40, len(plain) * 0.07):
            return True
        if _spanish_word_ratio(remainder) < 0.2:
            return True
        return False
    if _CERT_SERIAL.search(plain) and len(plain) < 120 and _CFDI_SEAL_LINE.search(plain):
        return True
    if _contains_long_base64_blob(plain):
        return True
    return False


_KNOWN_DOC_LINE = re.compile(
    r"(?i)(^(nombre|domicilio|curp|vigencia|secci[oó]n|sexo|m[eé]xico)$|"
    r"clave\s+de\s+elector|fecha\s+de\s+nacimiento|a[nñ]o\s+de\s+registro|"
    r"identificador\s+electr[oó]nico|total\s+a\s+pagar|saldo\s+al\s+corte|"
    r"^[A-Z0-9<]{24,}$|☑|☐)"
)


def is_stamp_ocr_gibberish(line: str) -> bool:
    """Detecta líneas típicas de sellos circulares mal OCR-eados."""
    plain = line.strip()
    if len(plain) < 4:
        return True
    if plain.startswith("#"):
        return False
    if _KNOWN_DOC_LINE.search(plain):
        return False

    tokens = re.findall(r"\S+", plain)
    if not tokens:
        return True

    tiny = sum(1 for t in tokens if len(t) <= 3)
    if len(tokens) >= 4 and tiny / len(tokens) > 0.5 and _spanish_word_ratio(plain) < 0.38:
        return True

    if len(tokens) >= 6 and _spanish_word_ratio(plain) < 0.22:
        return True

    if re.search(r"(\b[A-Za-z]{1,2}\b[\s|]{1,3}){5,}", plain) and _spanish_word_ratio(plain) < 0.35:
        return True

    # Mezcla alfanumérica sin palabras (ej. u9I33n4suo > OJNJIDNDIY, a9mri1c100)
    if re.search(r"[A-Za-z]+\d+[A-Za-z0-9]*|\d+[A-Za-z]+", plain) and _spanish_word_ratio(plain) < 0.35:
        real_words = re.findall(r"\b[a-záéíóúñA-ZÁÉÍÓÚÑ]{5,}\b", plain)
        if not real_words or sum(1 for w in real_words if re.search(r"[aeiouáéíóú]", w, re.I)) == 0:
            return True

    if len(plain) >= 12 and _spanish_word_ratio(plain) < 0.15:
        letters = sum(1 for c in plain if c.isalpha())
        digits = sum(1 for c in plain if c.isdigit())
        if letters >= 6 and digits >= 1:
            return True

    # Líneas cortas sin palabra reconocible
    if len(plain) <= 45:
        real_words = _meaningful_words(plain)
        if len(tokens) <= 5 and not real_words:
            if not re.search(r"(?:RFC|IVA|CP|CURP|BBVA|\$|TOTAL|PAGADO)", plain, re.I):
                return True

    if len(plain) < 60 and not _meaningful_words(plain):
        if re.search(r"[A-Za-z0-9+/=]{5,}", plain):
            if not re.search(r"(?:RFC|IVA|CP|CURP|BBVA|TOTAL|PAGADO|\$|\d{3,})", plain, re.I):
                return True

    # Mayúsculas sueltas tipo sello (OJNJIDNDIY, MEUS, ASIN) — no nombres/campos oficiales
    if len(plain) <= 30 and plain.upper() == plain and _spanish_word_ratio(plain) < 0.25:
        if _meaningful_words(plain):
            return False
        if not re.search(r"\d{3,}", plain):
            return True

    return False


_SPANISH_FUNCTION_WORDS = frozenset({
    "a", "al", "de", "del", "el", "en", "la", "las", "lo", "los", "un", "una", "y", "o",
    "que", "por", "con", "su", "se", "es", "como", "para", "mi", "me", "le", "si", "no",
})


def is_garbage_ocr_line(line: str) -> bool:
    """Filtra líneas típicas de sellos mal leídos, CFDI o ruido de escaneo."""
    plain = line.strip()
    if not plain:
        return True
    if is_cfdi_seal_line(plain):
        return True
    if _looks_like_seal_token_line(plain):
        return True
    if is_stamp_ocr_gibberish(plain):
        return True
    if _CFDI_BOILERPLATE.search(plain) and len(plain) < 160:
        return True

    tokens = plain.split()
    if len(tokens) >= 8:
        short = sum(
            1 for t in tokens
            if len(t) <= 2 and t.lower().strip(".,;:") not in _SPANISH_FUNCTION_WORDS
        )
        if short / len(tokens) > 0.48:
            return True
        words = [t for t in tokens if len(t) >= 3 and re.search(r"[a-záéíóúñA-ZÁÉÍÓÚÑ]", t)]
        if len(tokens) >= 8 and len(words) / len(tokens) < 0.28:
            return True

    if len(plain) >= 20:
        letters = sum(1 for c in plain if c.isalpha())
        if letters < 6:
            return True
        vowels = sum(1 for c in plain.lower() if c in "aeiouáéíóú")
        if letters >= 10 and vowels / letters < 0.12:
            return True

    if len(plain) >= 80:
        alnum = sum(1 for c in plain if c.isalnum() or c in "áéíóúñÁÉÍÓÚÑ$.,")
        if alnum / len(plain) < 0.52:
            return True
        if _base64_char_ratio(plain) > 0.42 and _spanish_word_ratio(plain) < 0.25:
            return True

    return False


_ALLOWED_DOC_HEADINGS = re.compile(
    r"(?i)(instituto\s+nacional\s+elector|credencial\s+para\s+votar|"
    r"estados\s+unidos\s+mexicanos|acta\s+de\s+nacimiento|checklist|"
    r"requisitos\s+de\s+contrataci|telmex|telefonos?\s+de\s+mexico|\bcfe\b|"
    r"recibo|comprobante|experiencia\s+laboral|mi\s+perfil|formaci[oó]n|"
    r"datos\s+de\s+(?:la\s+persona|filiaci[oó]n))"
)


def _is_watermark_heading(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("#"):
        return False
    plain = re.sub(r"^#+\s*", "", stripped).strip()
    if not plain or len(plain) > 72:
        return False
    if _ALLOWED_DOC_HEADINGS.search(plain):
        return False
    letters = re.sub(r"[^A-Za-zÁÉÍÓÚáéíóúÑñ]", "", plain)
    if letters and letters.isupper() and 2 <= len(plain.split()) <= 5:
        return True
    if _looks_like_seal_token_line(plain) or is_stamp_ocr_gibberish(plain):
        return True
    return False


def _is_markdown_table_line(line: str) -> bool:
    stripped = (line or "").strip()
    return (
        stripped.startswith("|")
        and stripped.endswith("|")
        and stripped.count("|") >= 3
    )


def clean_ocr_line(line: str, *, aggressive: bool = True) -> Optional[str]:
    stripped = line.strip()
    if not stripped:
        return ""
    if _is_markdown_table_line(stripped):
        return stripped
    if stripped.startswith("#"):
        if _is_watermark_heading(stripped):
            return None
        return stripped
    if any(p.search(stripped) for p in _STAMP_LINE_PATTERNS):
        return None
    if _SCANNER_WATERMARK.match(stripped):
        return None
    if len(stripped) <= 3 and not stripped.isdigit():
        return None
    cleaned = strip_inline_cfdi_noise(stripped)
    if not cleaned:
        return None
    if (
        len(stripped) > 180
        and len(cleaned) >= 35
        and _spanish_word_ratio(cleaned) >= 0.32
        and not _looks_like_seal_token_line(cleaned)
    ):
        return cleaned
    if is_garbage_ocr_line(cleaned) or is_garbage_ocr_line(stripped):
        return None
    if aggressive and len(stripped) > 90 and len(cleaned) < max(22, len(stripped) * 0.14):
        return None
    if aggressive and len(cleaned) >= 120 and _spanish_word_ratio(cleaned) < 0.18:
        return None
    if aggressive and len(cleaned) <= 14 and _spanish_word_ratio(cleaned) < 0.45:
        if not re.search(r"[\$€]|^\d+[.,]\d{2}$|RFC|IVA|CP\b", cleaned, re.I):
            return None
    return cleaned


def _collapse_stamp_runs(lines: list[str]) -> list[str]:
    """Colapsa bloques largos de líneas vacías consecutivas tras filtrar sellos."""
    out: list[str] = []
    blank_run = 0
    for line in lines:
        if not line.strip():
            blank_run += 1
            if blank_run <= 2:
                out.append("")
            continue
        blank_run = 0
        out.append(line)
    return out


def sanitize_document_markdown(markdown: str, *, aggressive: bool = True) -> str:
    """Quita sellos, CFDI digital y ruido de escaneo del Markdown."""
    if not markdown:
        return ""
    text = strip_vision_meta_output(markdown.strip())
    text = re.split(
        r"(?i)\n##\s+Firmas,\s+sellos\s+y\s+anotaciones\b.*",
        text,
        maxsplit=1,
    )[0].strip()
    if aggressive:
        text = _remove_base64_blobs(text)
    else:
        text = strip_electronic_watermarks(text)
    lines: list[str] = []
    skipped_run = 0
    for line in text.splitlines():
        if line.strip().startswith("#"):
            if skipped_run >= 10:
                lines.append("")
                lines.append("> Sello o timbre omitido.")
                lines.append("")
            skipped_run = 0
            cleaned_heading = clean_ocr_line(line, aggressive=aggressive)
            if cleaned_heading is None:
                skipped_run += 1
                continue
            lines.append(cleaned_heading)
            continue
        cleaned = clean_ocr_line(line, aggressive=aggressive)
        if cleaned is None:
            if line.strip():
                skipped_run += 1
            continue
        if aggressive and skipped_run >= 10:
            lines.append("")
            lines.append("> Sello o timbre omitido.")
            lines.append("")
        skipped_run = 0
        if not cleaned.strip():
            lines.append("")
            continue
        lines.append(cleaned)
    cleaned = "\n".join(_collapse_stamp_runs(lines))
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    if aggressive:
        cleaned = re.sub(
            r"(?m)^> Sello o timbre omitido\.\s*\n(?:\s*\n> Sello o timbre omitido\.\s*\n)+",
            "> Sello o timbre omitido.\n\n",
            cleaned,
        )
    if aggressive:
        return ensure_readable_markdown(cleaned)
    return normalize_markdown(cleaned)

def _vision_page_score(text: str, *, table_mode: bool = False) -> float:
    score = text_quality_score(text)
    if table_mode:
        score += table_ocr.score_table_markdown(text) * 0.45
    meta_hits = sum(1 for line in (text or "").splitlines() if _VISION_META_LINE.search(line.strip()))
    score -= meta_hits * 30
    if _VISION_REFUSAL.search(text or ""):
        score -= 200
    return score


def try_vision_page(
    image: Image.Image,
    *,
    table_mode: bool = False,
    tables_only: bool = False,
    docs_only: bool = False,
    page_kind: str = scan_pages.PAGE_GENERIC,
    provider: Optional[str] = None,
    prompt_override: Optional[str] = None,
    provider_pool: Optional[list[str]] = None,
) -> tuple[str, str, dict]:
    if prompt_override:
        prompt = prompt_override
    elif table_mode:
        prompt = GEMINI_TABLE_PROMPT
    elif page_kind == scan_pages.PAGE_RECEIPT:
        prompt = RECEIPT_VISION_PROMPT
    elif page_kind == scan_pages.PAGE_CREDENTIAL:
        prompt = CREDENTIAL_VISION_PROMPT
    elif page_kind == scan_pages.PAGE_FORM:
        prompt = FORM_DOCUMENT_VISION_PROMPT
    elif page_kind == scan_pages.PAGE_CHECKLIST:
        prompt = CHECKLIST_VISION_PROMPT
    elif page_kind == scan_pages.PAGE_CV:
        prompt = CV_VISION_PROMPT
    elif page_kind == scan_pages.PAGE_CONTRACT:
        prompt = CONTRACT_VISION_PROMPT
    else:
        prompt = LITERAL_PAGE_OCR_PROMPT
    if provider_pool is not None:
        pool = provider_pool
    elif tables_only:
        pool = vision_llm.providers_for_tables()
    elif docs_only:
        pool = vision_llm.providers_for_legal() or vision_llm.providers_for_documents()
    else:
        pool = None
    chosen = provider or get_vision_provider()
    md, source, meta = vision_llm.ocr_page_vision(
        image,
        prompt,
        provider=chosen,
        provider_pool=pool,
        sanitize=lambda raw: sanitize_document_markdown(raw, aggressive=False),
        score_fn=lambda raw: _vision_page_score(raw, table_mode=table_mode),
    )
    if not md:
        return "", "", meta
    label = source or "vision"
    if table_mode and label:
        label = f"{label}-table"
    elif docs_only and label:
        label = f"{label}-legal" if page_kind in (
            scan_pages.PAGE_CONTRACT,
            scan_pages.PAGE_GENERIC,
        ) else f"{label}-doc"
    return md, label, meta


def try_gemini_page(image: Image.Image, *, table_mode: bool = False) -> str:
    md, _, _ = try_vision_page(image, table_mode=table_mode)
    return md


def try_vision_pdf(file_path: Path, lang: str = "spa", *, mode: str = "balanced") -> Optional[dict[str, Any]]:
    if not vision_llm.configured_providers():
        return None
    mode = normalize_ocr_mode(mode)
    page_count = count_pdf_pages(file_path) if file_path.suffix.lower() == ".pdf" else 1
    if not vision_allowed_for_pages(page_count):
        return None
    dpi = ocr_dpi_for_mode(mode)
    try:
        suffix = file_path.suffix.lower()
        pages: list[dict[str, Any]] = []
        warnings: list[str] = []
        local_fallback_pages = 0
        providers_used: dict[str, int] = {}
        quota_warnings_seen: set[str] = set()

        def vision_job(item: tuple[int, Image.Image]) -> tuple[int, str, str, dict]:
            page_index, image = item
            md, source, meta = try_vision_page(image)
            return page_index, md, source, meta

        def record_vision_meta(meta: dict) -> None:
            for warning in meta.get("quota_warnings") or []:
                if warning not in quota_warnings_seen:
                    quota_warnings_seen.add(warning)
                    append_warning(warnings, warning)

        if suffix == ".pdf":
            total_pages = count_pdf_pages(file_path)
            if total_pages > OCR_CHUNK_PAGES:
                append_warning(
                    warnings,
                    f"PDF grande procesado por bloques internos de {OCR_CHUNK_PAGES} páginas para mejorar estabilidad."
                )
            for batch in iter_pdf_page_chunks(file_path, dpi=dpi, chunk_size=OCR_CHUNK_PAGES):
                by_idx = {idx: im for idx, im in batch}
                workers = min(OCR_GEMINI_WORKERS, len(batch))
                if workers <= 1:
                    vision_results = [vision_job(item) for item in batch]
                else:
                    with ThreadPoolExecutor(max_workers=workers) as executor:
                        vision_results = list(executor.map(vision_job, batch))

                for page_index, page_markdown, source, meta in sorted(vision_results, key=lambda item: item[0]):
                    record_vision_meta(meta)
                    image = by_idx[page_index]
                    final_md = page_markdown
                    if text_quality_score(page_markdown) < 28 or len(strip_markdown_syntax(page_markdown)) < OCR_MIN_PAGE_CHARS:
                        local_markdown = ocr_image_to_markdown(image, lang=lang, mode=mode)
                        if text_quality_score(local_markdown) > text_quality_score(page_markdown) + 8:
                            final_md = local_markdown
                            local_fallback_pages += 1
                    elif source and _is_vision_ocr_source(source):
                        for prov in _vision_source_providers(source):
                            providers_used[prov] = providers_used.get(prov, 0) + 1
                    if final_md:
                        pages.append({
                            "index": page_index,
                            "markdown": final_md,
                        })
        else:
            image = Image.open(file_path)
            page_markdown, source, meta = try_vision_page(image)
            record_vision_meta(meta)
            if text_quality_score(page_markdown) < 28 or len(strip_markdown_syntax(page_markdown)) < OCR_MIN_PAGE_CHARS:
                local_markdown = ocr_image_to_markdown(image, lang=lang, mode=mode)
                if text_quality_score(local_markdown) > text_quality_score(page_markdown) + 8:
                    page_markdown = local_markdown
                    local_fallback_pages += 1
            elif source and _is_vision_ocr_source(source):
                for prov in _vision_source_providers(source):
                    providers_used[prov] = providers_used.get(prov, 0) + 1
            if page_markdown:
                pages.append({
                    "index": 0,
                    "markdown": page_markdown,
                })

        if not pages:
            return None

        if local_fallback_pages:
            append_warning(
                warnings,
                f"{local_fallback_pages} página(s) se resolvieron con OCR local reforzado porque la salida IA fue débil."
            )

        method = "vision"
        if providers_used:
            method = "parallel" if vision_llm.use_parallel_mode(get_vision_provider()) else max(
                providers_used, key=providers_used.get
            )

        return build_result(method, join_pages_markdown(pages), pages=pages, file_path=file_path, warnings=warnings)
    except Exception as e:
        print(f"Vision OCR error: {e}")
        return None


def try_gemini(file_path: Path, lang: str = "spa", *, mode: str = "balanced") -> Optional[dict[str, Any]]:
    return try_vision_pdf(file_path, lang=lang, mode=mode)

# ── Preprocesado de imagen ────────────────────────────────────────────────────
def upscale_for_ocr(image: Image.Image, *, min_edge: Optional[int] = None) -> Image.Image:
    target = min_edge if min_edge is not None else OCR_UPSCALE_MIN_EDGE
    img = ImageOps.exif_transpose(image).convert("L")
    edge = min(img.size) if img.size else 0
    if edge and edge < target:
        scale = min(2.5, target / float(edge))
        img = img.resize(
            (max(1, int(img.width * scale)), max(1, int(img.height * scale))),
            Image.Resampling.LANCZOS,
        )
    return img

def correct_orientation(image: Image.Image) -> Image.Image:
    corrected, _ = detect_and_correct_rotation(image)
    return corrected

def preprocess_variants(
    image: Image.Image,
    *,
    skip_osd: bool = False,
    min_edge: Optional[int] = None,
    full: bool = False,
) -> dict[str, Image.Image]:
    base = upscale_for_ocr(image, min_edge=min_edge)
    base = ImageOps.autocontrast(base, cutoff=1)
    if not skip_osd and OCR_FIX_ROTATION:
        base = correct_orientation(base)

    soft = ImageEnhance.Contrast(base).enhance(1.25)
    soft = soft.filter(ImageFilter.MedianFilter(size=3))

    strong = ImageEnhance.Contrast(base).enhance(1.85)
    strong = ImageEnhance.Sharpness(strong).enhance(1.7)
    strong = strong.filter(ImageFilter.SHARPEN)

    mean_value = ImageStat.Stat(strong).mean[0] if strong.size else 150
    threshold = max(115, min(190, int(mean_value * 0.92)))
    binary = strong.point(lambda value: 255 if value > threshold else 0, "L")

    variants: dict[str, Image.Image] = {
        "base": base,
        "soft": soft,
        "strong": strong,
        "binary": binary,
    }
    if not full:
        return variants

    variants["invert"] = ImageOps.invert(base)
    fine = ImageEnhance.Contrast(base).enhance(0.88)
    fine = ImageEnhance.Brightness(fine).enhance(1.1)
    fine = fine.filter(ImageFilter.GaussianBlur(radius=0.6))
    variants["fine"] = ImageEnhance.Sharpness(fine).enhance(2.0)
    return variants


def emphasis_stamp_channel(image: Image.Image, channel: str, *, min_edge: Optional[int] = None) -> Image.Image:
    """Resalta tintas rojas o azules típicas de sellos y timbres."""
    rgb = ImageOps.exif_transpose(image).convert("RGB")
    r, g, b = rgb.split()
    if channel == "red":
        gray = ImageChops.subtract(r, ImageChops.lighter(g, b))
    else:
        gray = ImageChops.subtract(b, ImageChops.lighter(r, g))
    gray = ImageChops.add(gray, Image.new("L", gray.size, 48))
    gray = ImageOps.autocontrast(gray, cutoff=2)
    edge = min(gray.size) if gray.size else 0
    target = min_edge if min_edge is not None else OCR_UPSCALE_MIN_EDGE
    if edge and edge < target:
        scale = min(2.5, target / float(edge))
        gray = gray.resize(
            (max(1, int(gray.width * scale)), max(1, int(gray.height * scale))),
            Image.Resampling.LANCZOS,
        )
    return gray


def annotation_preprocess_variants(
    image: Image.Image,
    *,
    min_edge: Optional[int] = None,
    mode: str = "balanced",
) -> dict[str, Image.Image]:
    """Variantes orientadas a texto disperso, sellos y trazos finos (recorte en equilibrado)."""
    full = normalize_ocr_mode(mode) == "quality"
    all_variants = preprocess_variants(image, skip_osd=True, min_edge=min_edge, full=full)
    if full:
        pick = ("fine", "soft", "binary", "strong")
        return {name: all_variants[name] for name in pick if name in all_variants}
    return {
        "binary": all_variants["binary"],
        "soft": all_variants["soft"],
    }

# ── Limpieza de texto ─────────────────────────────────────────────────────────
def clean_text(text: str, *, min_useful_ratio: float = 0.55, min_len_skip: int = 40) -> str:
    lines = text.split("\n")
    cleaned = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if cleaned and cleaned[-1] != "":
                cleaned.append("")
            continue
        total = len(stripped)
        useful = len(re.findall(r'[a-zA-ZáéíóúÁÉÍÓÚñÑüÜ0-9\s,\.;\:\!\?\(\)\"\'\-\/]', stripped))
        ratio = useful / total if total > 0 else 0
        if ratio < min_useful_ratio and total < min_len_skip:
            continue
        line_clean = re.sub(r'[^\x20-\x7EáéíóúÁÉÍÓÚñÑüÜ«»°\n]', '', stripped)
        line_clean = line_clean.strip()
        if line_clean:
            cleaned.append(line_clean)
    return "\n".join(cleaned)

# ── OCR Tesseract ─────────────────────────────────────────────────────────────
LOCAL_OCR_CONFIGS = [
    r"--oem 3 --psm 3 -c preserve_interword_spaces=1",
    r"--oem 3 --psm 4 -c preserve_interword_spaces=1",
    r"--oem 3 --psm 6 -c preserve_interword_spaces=1",
    r"--oem 3 --psm 11 -c preserve_interword_spaces=1",
    r"--oem 3 --psm 1 -c preserve_interword_spaces=1",
]

ANNOTATION_OCR_CONFIGS = [
    r"--oem 3 --psm 11 -c preserve_interword_spaces=1",
    r"--oem 3 --psm 8 -c preserve_interword_spaces=1",
    r"--oem 3 --psm 7 -c preserve_interword_spaces=1",
    r"--oem 3 --psm 13 -c preserve_interword_spaces=1",
]


def annotations_section_present(markdown: str) -> bool:
    return ANNOTATIONS_SECTION.lower() in (markdown or "").lower()


def append_annotations_section(markdown: str, annotations_body: str) -> str:
    body = (annotations_body or "").strip()
    if not body:
        return markdown
    base = normalize_markdown(markdown)
    if annotations_section_present(base):
        return base
    return f"{base}\n\n{ANNOTATIONS_SECTION}\n\n{body}".strip()


def filter_lines_not_in_main(lines: list[str], main_plain: str) -> list[str]:
    main_lower = (main_plain or "").lower()
    filtered: list[str] = []
    for line in lines:
        norm = line.strip()
        if len(norm) < 2:
            continue
        key = norm.lower()
        if key in main_lower:
            continue
        if len(norm) <= 80 and any(main_lower.find(key[i : i + 12]) >= 0 for i in range(0, max(1, len(key) - 11), 4) if len(key) >= 12):
            continue
        filtered.append(norm)
    return filtered


def format_annotations_markdown(lines: list[str]) -> str:
    blocks: list[str] = []
    for line in lines:
        compact = re.sub(r"\s+", " ", line).strip()
        if not compact:
            continue
        if re.match(r"^[A-Z0-9][A-Z0-9\s\-\./]{2,}$", compact) and len(compact) <= 48:
            blocks.append(f"- **Sello / código:** {compact}")
        elif len(compact) <= 42 and re.search(r"[a-záéíóúñ]", compact):
            blocks.append(f"- **Firma / manuscrito:** {compact}")
        elif re.search(r"\d{2,}", compact) and len(compact) <= 32:
            blocks.append(f"- **Sello / código:** {compact}")
        else:
            blocks.append(f"- {compact}")
    return "\n".join(blocks)


def should_run_supplementary_pass(markdown: str, mode: str) -> bool:
    profile = get_ocr_profile()
    if profile.get("supplementary") is False:
        return False
    if not OCR_SUPPLEMENTARY_PASS or normalize_ocr_mode(mode) == "fast":
        return False
    if annotations_section_present(markdown):
        return False
    if OCR_TESSERACT_ONLY:
        return True
    if text_quality_score(markdown) >= OCR_SUPPLEMENTARY_MIN_SCORE:
        return False
    return True


def ocr_supplementary_annotations(image: Image.Image, lang: str, mode: str) -> str:
    """Segunda pasada: sellos, firmas y texto disperso que el bloque principal suele omitir."""
    if not OCR_SUPPLEMENTARY_PASS or mode == "fast":
        return ""
    mode = normalize_ocr_mode(mode)
    min_edge = upscale_min_edge_for_mode(mode)
    best_by_line: dict[str, dict[str, Any]] = {}
    configs = ANNOTATION_OCR_CONFIGS if mode == "quality" else ANNOTATION_OCR_CONFIGS[:2]

    for variant in annotation_preprocess_variants(image, min_edge=min_edge, mode=mode).values():
        for config in configs:
            try:
                candidate = run_ocr_candidate(
                    variant,
                    lang=lang,
                    config=config,
                    min_useful_ratio=0.38,
                    min_len_skip=12,
                    compute_confidence=False,
                )
            except Exception:
                continue
            for line in (candidate.get("text") or "").split("\n"):
                key = line.strip().lower()
                if len(key) < 2:
                    continue
                prev = best_by_line.get(key)
                if not prev or candidate["score"] > prev["score"]:
                    best_by_line[key] = {"line": line.strip(), "score": candidate["score"]}

    ranked = sorted(best_by_line.values(), key=lambda item: item["score"], reverse=True)
    return format_annotations_markdown([item["line"] for item in ranked[:24]])


def enrich_markdown_with_annotations(
    markdown: str,
    image: Image.Image,
    lang: str,
    mode: str,
) -> str:
    if not should_run_supplementary_pass(markdown, mode):
        return markdown
    main_plain = strip_markdown_syntax(markdown)
    extra = ocr_supplementary_annotations(image, lang=lang, mode=mode)
    if not extra:
        return markdown
    unique_lines = filter_lines_not_in_main(extra.split("\n"), main_plain)
    unique_lines = [
        re.sub(r"^-\s*\*\*[^*]+\*\*:\s*", "", line).strip()
        for line in unique_lines
        if line.strip()
    ]
    if not unique_lines:
        return markdown
    body = format_annotations_markdown(unique_lines)
    return append_annotations_section(markdown, body)

def run_ocr_candidate(
    image: Image.Image,
    lang: str,
    config: str,
    *,
    min_useful_ratio: float = 0.55,
    min_len_skip: int = 40,
    compute_confidence: bool = True,
) -> dict[str, Any]:
    if OCR_USE_LAYOUT:
        layout_text = tesseract_layout.extract_text_with_layout(
            image,
            lang,
            config,
            line_filter=is_stamp_ocr_gibberish,
        )
        cleaned = clean_text(
            layout_text,
            min_useful_ratio=min_useful_ratio,
            min_len_skip=min_len_skip,
        )
    else:
        raw_text = pytesseract.image_to_string(image, lang=lang, config=config)
        cleaned = clean_text(raw_text, min_useful_ratio=min_useful_ratio, min_len_skip=min_len_skip)
    if not cleaned:
        return {"text": "", "confidence": 0.0, "score": 0.0}

    confidence = 55.0
    if compute_confidence:
        try:
            data = pytesseract.image_to_data(image, lang=lang, config=config, output_type=Output.DICT)
            confidences = [float(conf) for conf in data.get("conf", []) if conf not in {"-1", "-1.0", ""}]
            if confidences:
                confidence = sum(confidences) / len(confidences)
        except Exception:
            confidence = 0.0

    score = text_quality_score(cleaned) + confidence * 1.35
    return {
        "text": cleaned,
        "confidence": confidence,
        "score": score,
    }

def ocr_image_to_markdown(image: Image.Image, lang: str = "spa", *, mode: str = "balanced") -> str:
    profile = get_ocr_profile()
    mode = profile.get("mode", normalize_ocr_mode(mode))
    min_edge = int(profile.get("min_edge", upscale_min_edge_for_mode(mode)))
    skip_osd = bool(profile.get("skip_osd", False))
    fast_conf = mode != "quality"
    if not OCR_TESSERACT_ONLY:
        page_kind = scan_pages.classify_scan_page(image, lang=lang)
        if not scan_pages.page_needs_docs_vision(image, lang=lang, page_kind=page_kind):
            image = suppress_visual_noise_for_ocr(image)

    def _finish(text: str) -> str:
        base = sanitize_document_markdown(text, aggressive=False)
        try:
            return enrich_markdown_with_annotations(base, image, lang=lang, mode=mode)
        except Exception:
            return base

    if mode == "fast":
        img = upscale_for_ocr(image, min_edge=min_edge)
        img = ImageOps.autocontrast(img, cutoff=1)
        best_candidate = {"text": "", "confidence": 0.0, "score": 0.0}
        for config in LOCAL_OCR_CONFIGS[:2]:
            try:
                candidate = run_ocr_candidate(
                    img, lang=lang, config=config, compute_confidence=False
                )
            except Exception:
                continue
            if candidate["score"] > best_candidate["score"]:
                best_candidate = candidate
        return _finish(best_candidate["text"])

    if mode == "quality" and not profile.get("bulk"):
        best_candidate = {"text": "", "confidence": 0.0, "score": 0.0}
        for variant in preprocess_variants(image, skip_osd=skip_osd, min_edge=min_edge, full=True).values():
            for config in LOCAL_OCR_CONFIGS:
                try:
                    candidate = run_ocr_candidate(variant, lang=lang, config=config)
                except Exception:
                    continue
                if candidate["score"] > best_candidate["score"]:
                    best_candidate = candidate
        return _finish(best_candidate["text"])

    balanced_configs = LOCAL_OCR_CONFIGS[:3]
    best_quick = {"text": "", "confidence": 0.0, "score": 0.0}
    img_quick = upscale_for_ocr(image, min_edge=min_edge)
    img_quick = ImageOps.autocontrast(img_quick, cutoff=1)
    try:
        best_quick = run_ocr_candidate(
            img_quick, lang=lang, config=LOCAL_OCR_CONFIGS[0], compute_confidence=fast_conf
        )
    except Exception:
        pass

    quick_plain = strip_markdown_syntax(normalize_markdown(best_quick["text"]))
    min_chars = OCR_MIN_PAGE_CHARS if profile.get("balanced_fast") else OCR_MIN_PAGE_CHARS * 2
    if best_quick["score"] >= OCR_BALANCED_GOOD_SCORE and len(quick_plain) >= min_chars:
        return _finish(best_quick["text"])

    if profile.get("bulk") and best_quick["text"]:
        return _finish(best_quick["text"])

    best_full = {"text": "", "confidence": 0.0, "score": 0.0}
    if profile.get("balanced_fast"):
        variants = [
            upscale_for_ocr(image, min_edge=min_edge),
            preprocess_variants(image, skip_osd=skip_osd, min_edge=min_edge, full=False).get("soft")
            or upscale_for_ocr(image, min_edge=min_edge),
        ]
        refine_configs = LOCAL_OCR_CONFIGS[:2]
    else:
        variants = list(preprocess_variants(image, skip_osd=skip_osd, min_edge=min_edge, full=False).values())
        refine_configs = LOCAL_OCR_CONFIGS[:2] if profile.get("bulk") else balanced_configs
    if profile.get("bulk"):
        variants = variants[:2]
    for variant in variants:
        for config in refine_configs:
            try:
                candidate = run_ocr_candidate(
                    variant, lang=lang, config=config, compute_confidence=fast_conf
                )
            except Exception:
                continue
            if candidate["score"] > best_full["score"]:
                best_full = candidate

    chosen = best_full if best_full["score"] >= best_quick["score"] else best_quick
    return _finish(chosen["text"])

def process_with_tesseract(file_path: Path, lang: str = "spa", *, mode: str = "balanced") -> dict[str, Any]:
    mode = normalize_ocr_mode(mode)
    dpi = ocr_dpi_for_mode(mode)
    suffix = file_path.suffix.lower()
    pages: list[dict[str, Any]] = []
    warnings: list[str] = []

    def tesseract_job(item: tuple[int, Image.Image]) -> tuple[int, str]:
        page_index, img = item
        text = ocr_image_to_markdown(img, lang=lang, mode=mode)
        return page_index, text

    if suffix == ".pdf":
        total_pages = count_pdf_pages(file_path)
        if total_pages > OCR_CHUNK_PAGES:
            append_warning(
                warnings,
                f"PDF largo procesado por bloques internos de {OCR_CHUNK_PAGES} páginas para ahorrar memoria."
            )
        for batch in iter_pdf_page_chunks(file_path, dpi=dpi, chunk_size=OCR_CHUNK_PAGES):
            workers = min(OCR_TESSERACT_WORKERS, len(batch))
            if workers <= 1:
                results = [tesseract_job(item) for item in batch]
            else:
                with ThreadPoolExecutor(max_workers=workers) as executor:
                    results = list(executor.map(tesseract_job, batch))
            for page_index, text in sorted(results, key=lambda item: item[0]):
                if text:
                    pages.append({
                        "index": page_index,
                        "markdown": text,
                        "ocr_sourced": True,
                    })
    elif suffix in {".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp"}:
        img = Image.open(file_path)
        text = ocr_image_to_markdown(img, lang=lang, mode=mode)
        if text:
            pages.append({
                "index": 0,
                "markdown": text,
                "ocr_sourced": True,
            })
    else:
        raise ValueError(f"Tipo no soportado: {suffix}")
    return build_result("structured", join_pages_markdown(pages), pages=pages, file_path=file_path, warnings=warnings)

# ── Procesador principal ──────────────────────────────────────────────────────
def process_file(
    file_path: Path,
    lang: str = "spa",
    *,
    mode: str = "balanced",
    range_slice: bool = False,
    vision_provider: Optional[str] = None,
) -> dict[str, Any]:
    suffix = file_path.suffix.lower()
    mode = normalize_ocr_mode(mode)
    total_steps = count_pdf_pages(file_path) if suffix == ".pdf" else 1
    if range_slice:
        profile = build_ocr_profile_for_range(mode, total_steps)
    else:
        profile = build_ocr_profile(mode, total_steps)
    rot_token = _rotation_fixes_ctx.set([])
    profile_token = set_ocr_profile(profile)
    vision_token = set_vision_provider(normalize_vision_provider(vision_provider))
    report_ocr_progress(0, total_steps, "Iniciando transcripción del documento…")

    try:
        # 1. PDF híbrido: digital por página + OCR estructurado donde haga falta
        if suffix == ".pdf":
            result = try_hybrid_pdf(file_path, lang=lang, mode=mode)
            if result:
                rotated = _rotation_fixes_ctx.get() or []
                if rotated:
                    pages_str = ", ".join(str(p) for p in sorted(set(rotated))[:12])
                    extra = f" (+{len(rotated) - 12} más)" if len(rotated) > 12 else ""
                    append_warning(
                        result.setdefault("warnings", []),
                        f"Se corrigió la orientación en {len(rotated)} página(s): {pages_str}{extra}.",
                    )
                return result

        # 2. Markitdown — fallback para PDFs digitales
        if suffix == ".pdf":
            result = try_markitdown(file_path)
            if result:
                return result

        # 3. Visión IA (solo si OCR_TESSERACT_ONLY=0)
        if not OCR_TESSERACT_ONLY and vision_allowed_for_pages(total_steps):
            result = try_vision_pdf(file_path, lang=lang, mode=mode)
            if result:
                return result

        # 4. Tesseract — OCR local
        effective = profile.get("mode", mode)
        result = process_with_tesseract(file_path, lang=lang, mode=effective)
        report_ocr_progress(total_steps, total_steps, "OCR finalizado")
        return result
    finally:
        _ocr_profile_ctx.reset(profile_token)
        _vision_provider_ctx.reset(vision_token)
        _rotation_fixes_ctx.reset(rot_token)


def persist_ocr_outputs(job_id: str, result: dict[str, Any], original_filename: str) -> dict[str, Any]:
    save_job_meta(job_id, original_filename)
    base = OUTPUT_DIR / job_id
    save_txt(result["text"], base.with_suffix(".txt"))
    save_md(result["markdown"], base.with_suffix(".md"))
    try:
        save_pdf(result["text"], base.with_suffix(".pdf"))
        pdf_ok = True
    except Exception:
        pdf_ok = False
    try:
        save_docx(result["markdown"], base.with_suffix(".docx"))
        docx_ok = True
    except Exception:
        docx_ok = False
    download_stem = safe_download_stem(original_filename)
    return {
        "txt": f"/download/{job_id}/txt",
        "md": f"/download/{job_id}/md",
        "pdf": f"/download/{job_id}/pdf" if pdf_ok else None,
        "docx": f"/download/{job_id}/docx" if docx_ok else None,
        "download_names": {"md": f"{download_stem}.md"},
    }


def build_ocr_response(
    job_id: str,
    result: dict[str, Any],
    *,
    original_filename: str,
    effective_mode: str,
    async_job: bool = False,
) -> dict[str, Any]:
    downloads = persist_ocr_outputs(job_id, result, original_filename)
    payload = {
        "job_id": job_id,
        "source_filename": original_filename,
        **result,
        "ocr_mode": effective_mode,
        "downloads": {
            "txt": downloads["txt"],
            "md": downloads["md"],
            "pdf": downloads["pdf"],
            "docx": downloads["docx"],
        },
        "download_names": downloads["download_names"],
    }
    if async_job:
        payload["async"] = True
    return payload


def execute_ocr_job(job_id: str) -> None:
    state = read_job_state(job_id)
    if not state:
        return
    if state.get("status") == "cancelled":
        return

    upload_raw = state.get("upload_path", "")
    upload_path = Path(upload_raw).resolve() if upload_raw else None
    if not upload_path or not upload_path.exists():
        write_job_state(job_id, {
            **state,
            "status": "failed",
            "detail": "Archivo de entrada no encontrado",
            "error": "upload_path missing",
        })
        batch_id_fail = state.get("batch_id")
        if batch_id_fail:
            batches.update_batch_part_status(JOBS_DIR, batch_id_fail, job_id, "failed")
        return

    lang = state.get("lang", "spa")
    mode = state.get("mode", "balanced")
    vision_provider = state.get("vision_provider", "auto")
    original_filename = state.get("source_filename", "documento")
    effective_mode = normalize_ocr_mode(mode)
    page_from = int(state.get("page_from") or 1)
    page_to = state.get("page_to")
    page_to = int(page_to) if page_to is not None else None
    batch_id = state.get("batch_id")
    write_job_state(job_id, {
        **state,
        "status": "processing",
        "progress": 0,
        "detail": "Preparando páginas…",
        "pages_done": 0,
        "pages_total": int(state.get("pages_total") or 0),
        "queue_position": 0,
    })
    work_path, pf, pt, delete_slice = prepare_work_pdf(upload_path, page_from, page_to)
    total_steps = pt - pf + 1

    control = _job_controls.get(job_id) or JobControl(job_id)
    _job_controls[job_id] = control

    def on_progress(done: int, total: int, detail: str, pct: int) -> None:
        if done > 0 and done % 3 != 0 and not throttle_job_progress_write(job_id):
            return
        current = read_job_state(job_id) or {}
        write_job_state(job_id, {
            **current,
            "status": "processing",
            "progress": pct,
            "detail": detail,
            "pages_done": done,
            "pages_total": total,
            "source_filename": original_filename,
            "lang": lang,
            "mode": effective_mode,
            "queue_position": 0,
        })

    progress_token = _progress_callback.set(on_progress)
    control_token = _job_control_ctx.set(control)
    try:
        write_job_state(job_id, {
            **state,
            "status": "processing",
            "progress": 0,
            "detail": "Procesando documento…",
            "pages_done": 0,
            "pages_total": total_steps,
            "queue_position": 0,
        })
        result = process_file(
            work_path,
            lang=lang,
            mode=effective_mode,
            range_slice=bool(state.get("range_slice")),
            vision_provider=vision_provider,
        )
        check_job_control_or_raise()
        processing_mode = state.get("processing_mode") or "ranges"
        upload_total = count_pdf_pages(upload_path) if upload_path.suffix.lower() == ".pdf" else 1
        is_full_document = processing_mode == "full" or (pf == 1 and pt >= upload_total)
        if is_full_document:
            part_filename = original_filename
        else:
            part_hdr = f"## {original_filename} — páginas {pf}-{pt}\n\n"
            result["markdown"] = part_hdr + (result.get("markdown") or "")
            result["text"] = f"--- Páginas {pf}-{pt} ---\n" + (result.get("text") or "")
            part_filename = f"{Path(original_filename).stem}_p{pf}-{pt}.pdf"
        report_ocr_progress(total_steps, total_steps, "Guardando archivos de salida…")
        response = build_ocr_response(
            job_id,
            result,
            original_filename=part_filename,
            effective_mode=effective_mode,
            async_job=True,
        )
        response["page_from"] = pf
        response["page_to"] = pt
        response["batch_id"] = batch_id
        response["processing_mode"] = processing_mode
        response["full_document"] = is_full_document
        if is_full_document:
            stem = safe_download_stem(original_filename)
            response.setdefault("download_names", {})["md"] = f"{stem}.md"
            response["batch_downloads"] = {
                "full_md": f"/download/{job_id}/md",
                "merged_md": f"/download/batch/{batch_id}/merged.md" if batch_id else None,
                "zip": f"/download/batch/{batch_id}/zip" if batch_id else None,
            }
        elif batch_id:
            response["batch_downloads"] = {
                "merged_md": f"/download/batch/{batch_id}/merged.md",
                "zip": f"/download/batch/{batch_id}/zip",
            }
        write_job_state(job_id, {
            "status": "completed",
            "progress": 100,
            "detail": f"Completado · páginas {pf}-{pt}",
            "pages_done": total_steps,
            "pages_total": total_steps,
            "page_from": pf,
            "page_to": pt,
            "batch_id": batch_id,
            "source_filename": original_filename,
            "lang": lang,
            "mode": effective_mode,
            "result": response,
        })
        if batch_id:
            batches.update_batch_part_status(JOBS_DIR, batch_id, job_id, "completed")
        append_job_history({
            "job_id": job_id,
            "source_filename": (
                original_filename if is_full_document else f"{original_filename} ({pf}-{pt})"
            ),
            "status": "completed",
            "completed_at": datetime.now(timezone.utc).isoformat(),
            "pages_total": total_steps,
            "page_from": pf,
            "page_to": pt,
            "batch_id": batch_id,
            "method": result.get("method"),
            "downloads": response.get("downloads"),
            "download_names": response.get("download_names"),
            "batch_downloads": response.get("batch_downloads"),
            "full_document": is_full_document,
        })
    except JobCancelled:
        write_job_state(job_id, {
            **state,
            "status": "cancelled",
            "progress": state.get("progress", 0),
            "detail": "Cancelado por el usuario",
            "source_filename": original_filename,
        })
        if batch_id:
            batches.update_batch_part_status(JOBS_DIR, batch_id, job_id, "cancelled")
    except Exception as exc:
        write_job_state(job_id, {
            **state,
            "status": "failed",
            "progress": 0,
            "detail": "Error al procesar",
            "error": str(exc),
            "source_filename": original_filename,
        })
        if batch_id:
            batches.update_batch_part_status(JOBS_DIR, batch_id, job_id, "failed")
        append_job_history({
            "job_id": job_id,
            "source_filename": original_filename,
            "status": "failed",
            "completed_at": datetime.now(timezone.utc).isoformat(),
            "pages_total": total_steps,
            "error": str(exc),
        })
    finally:
        _progress_callback.reset(progress_token)
        _job_control_ctx.reset(control_token)
        if delete_slice:
            work_path.unlink(missing_ok=True)
        if upload_path and should_delete_upload_after_job(state, upload_path):
            upload_path.unlink(missing_ok=True)
        gc.collect()


def pause_ocr_job(job_id: str, *, yield_slot: bool = True) -> dict[str, Any]:
    state = read_job_state(job_id)
    if not state:
        raise HTTPException(404, "Trabajo no encontrado")
    status = state.get("status", "")
    if status not in ("processing", "queued"):
        raise HTTPException(400, f"No se puede pausar en estado «{status}»")
    control = _job_controls.get(job_id)
    if not control:
        control = JobControl(job_id)
        _job_controls[job_id] = control
    is_running = False
    with _queue_lock:
        is_running = _running_job_id == job_id
    yield_now = yield_slot and is_running
    control.request_pause(yield_slot=yield_now)
    detail = (
        "Pausando… el siguiente en cola tomará el turno en breve."
        if yield_now
        else "En pausa. Pulsa Reanudar para continuar."
    )
    write_job_state(job_id, {
        **state,
        "status": "paused" if status == "processing" else "paused",
        "detail": detail,
    })
    if status == "queued":
        with _queue_lock:
            if job_id in _pending_job_ids:
                sync_queue_positions()
    return public_job_summary(read_job_state(job_id) or state)


def prioritize_ocr_job(job_id: str) -> dict[str, Any]:
    state = read_job_state(job_id)
    if not state:
        raise HTTPException(404, "Trabajo no encontrado")
    if state.get("status") not in ("queued", "paused"):
        raise HTTPException(400, "Solo se puede priorizar trabajos en cola o en pausa")
    with _queue_lock:
        if _running_job_id == job_id:
            raise HTTPException(400, "Ese trabajo ya está en curso. Pausa primero si quieres cambiar el orden.")
        if job_id not in _pending_job_ids:
            _pending_job_ids.append(job_id)
        _pending_job_ids[:] = [job_id] + [jid for jid in _pending_job_ids if jid != job_id]
    sync_queue_positions()
    kick_next_queue_worker()
    updated = read_job_state(job_id) or state
    return public_job_summary(updated)


def move_ocr_job_in_queue(job_id: str, direction: str) -> dict[str, Any]:
    state = read_job_state(job_id)
    if not state:
        raise HTTPException(404, "Trabajo no encontrado")
    if state.get("status") not in ("queued", "paused"):
        raise HTTPException(400, "Solo se puede reordenar trabajos en cola o en pausa")
    with _queue_lock:
        if _running_job_id == job_id:
            raise HTTPException(400, "No se puede mover el trabajo que está en curso")
        if job_id not in _pending_job_ids:
            raise HTTPException(400, "El trabajo no está en la cola")
        idx = _pending_job_ids.index(job_id)
        if direction == "up" and idx > 0:
            _pending_job_ids[idx - 1], _pending_job_ids[idx] = (
                _pending_job_ids[idx],
                _pending_job_ids[idx - 1],
            )
        elif direction == "down" and idx < len(_pending_job_ids) - 1:
            _pending_job_ids[idx + 1], _pending_job_ids[idx] = (
                _pending_job_ids[idx],
                _pending_job_ids[idx + 1],
            )
        else:
            raise HTTPException(400, "No se puede mover más en esa dirección")
    sync_queue_positions()
    return public_job_summary(read_job_state(job_id) or state)


def resume_ocr_job(job_id: str) -> dict[str, Any]:
    state = read_job_state(job_id)
    if not state:
        raise HTTPException(404, "Trabajo no encontrado")
    if state.get("status") not in ("paused", "queued"):
        raise HTTPException(400, "Solo se puede reanudar un trabajo en pausa o en cola")
    control = _job_controls.get(job_id) or JobControl(job_id)
    _job_controls[job_id] = control
    control.request_resume()
    with _queue_lock:
        is_running = _running_job_id == job_id
        in_pending = job_id in _pending_job_ids
    if state.get("status") == "queued" or in_pending:
        detail = state.get("detail", "En cola")
        new_status = "queued"
    else:
        detail = "Reanudando procesamiento…"
        new_status = "processing" if is_running else "queued"
    write_job_state(job_id, {
        **state,
        "status": new_status,
        "detail": detail,
    })
    spawn_ocr_worker(job_id)
    return public_job_summary(read_job_state(job_id) or state)


def cancel_ocr_job(job_id: str) -> dict[str, Any]:
    state = read_job_state(job_id)
    if not state:
        raise HTTPException(404, "Trabajo no encontrado")
    if state.get("status") in ("completed", "cancelled", "failed"):
        raise HTTPException(400, "El trabajo ya finalizó")
    control = _job_controls.get(job_id) or JobControl(job_id)
    _job_controls[job_id] = control
    control.request_cancel()

    with _queue_lock:
        if job_id in _pending_job_ids:
            _pending_job_ids[:] = [jid for jid in _pending_job_ids if jid != job_id]
        is_running = _running_job_id == job_id

    upload_raw = state.get("upload_path", "")
    if upload_raw and not state.get("keep_source"):
        Path(upload_raw).unlink(missing_ok=True)

    write_job_state(job_id, {
        **state,
        "status": "cancelled",
        "detail": "Cancelado por el usuario",
        "progress": state.get("progress", 0),
    })
    append_job_history({
        "job_id": job_id,
        "source_filename": state.get("source_filename", ""),
        "status": "cancelled",
        "completed_at": datetime.now(timezone.utc).isoformat(),
        "pages_total": state.get("pages_total", 0),
    })
    if is_running:
        pass
    return public_job_summary(read_job_state(job_id) or state)

# ── Exportadores ──────────────────────────────────────────────────────────────
def safe_download_stem(filename: str) -> str:
    stem = Path(filename or "documento").stem
    stem = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", stem).strip()
    return stem[:200] if stem else "documento"


def save_job_meta(job_id: str, original_filename: str) -> None:
    meta_path = OUTPUT_DIR / f"{job_id}.meta.json"
    meta_path.write_text(
        json.dumps({"filename": original_filename, "stem": safe_download_stem(original_filename)}),
        encoding="utf-8",
    )


def load_job_meta(job_id: str) -> dict[str, str]:
    meta_path = OUTPUT_DIR / f"{job_id}.meta.json"
    if not meta_path.exists():
        return {}
    try:
        data = json.loads(meta_path.read_text(encoding="utf-8"))
        return data if isinstance(data, dict) else {}
    except (json.JSONDecodeError, OSError):
        return {}


def save_txt(text: str, path: Path):
    path.write_text(text, encoding="utf-8")

def save_md(markdown: str, path: Path):
    content = normalize_markdown(markdown)
    if not content:
        content = "# Resultado OCR"
    path.write_text(content, encoding="utf-8")


def full_markdown_from_job_result(result: dict[str, Any]) -> str:
    pages = result.get("pages") or []
    if isinstance(pages, list) and len(pages) > 1:
        joined = join_pages_markdown(pages)
        if joined.strip():
            return joined
    markdown = (result.get("markdown") or "").strip()
    return markdown


def ensure_full_job_markdown(job_id: str) -> Path:
    """Regenera el .md desde las páginas guardadas si el archivo quedó incompleto."""
    md_path = OUTPUT_DIR / f"{job_id}.md"
    state = read_job_state(job_id)
    if not state or state.get("status") != "completed":
        return md_path
    result = state.get("result")
    if not isinstance(result, dict):
        return md_path
    full_md = full_markdown_from_job_result(result)
    if not full_md:
        return md_path
    pages = result.get("pages") or []
    page_count = len([p for p in pages if (p.get("markdown") or "").strip()])
    existing = md_path.read_text(encoding="utf-8") if md_path.exists() else ""
    expected_seps = max(0, page_count - 1)
    actual_seps = existing.count("\n\n---\n\n")
    needs_rebuild = (
        page_count > 1
        or not existing
        or len(existing.strip()) < len(full_md.strip()) - 32
        or (page_count > 1 and actual_seps < expected_seps)
    )
    if needs_rebuild:
        save_md(full_md, md_path)
        full_text = result.get("text") or ""
        if page_count > 1 and full_text:
            txt_path = OUTPUT_DIR / f"{job_id}.txt"
            if not txt_path.exists() or txt_path.stat().st_size < len(full_text) - 32:
                save_txt(full_text, txt_path)
    return md_path

def add_markdown_paragraph(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.add_run(text)

def save_docx(markdown: str, path: Path):
    content = normalize_markdown(markdown)
    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10.5)

    lines = content.split("\n")
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        heading = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading:
            level = min(len(heading.group(1)), 3)
            doc.add_heading(heading.group(2).strip(), level=level)
            index += 1
            continue

        if stripped == "---":
            doc.add_page_break()
            index += 1
            continue

        if "|" in stripped and index + 1 < len(lines) and re.match(r"^\s*\|?(?:\s*:?-+:?\s*\|)+\s*$", lines[index + 1].strip()):
            header_cells = [cell.strip() for cell in stripped.strip("|").split("|")]
            rows = []
            index += 2
            while index < len(lines):
                row_line = lines[index].strip()
                if not row_line or "|" not in row_line:
                    break
                rows.append([cell.strip() for cell in row_line.strip("|").split("|")])
                index += 1

            table = doc.add_table(rows=1, cols=len(header_cells))
            table.style = "Table Grid"
            for col, value in enumerate(header_cells):
                table.rows[0].cells[col].text = value
            for row_values in rows:
                row = table.add_row().cells
                for col in range(len(header_cells)):
                    row[col].text = row_values[col] if col < len(row_values) else ""
            continue

        if re.match(r"^\s*[-*]\s+", stripped):
            doc.add_paragraph(re.sub(r"^\s*[-*]\s+", "", stripped), style="List Bullet")
            index += 1
            continue

        if re.match(r"^\s*\d+\.\s+", stripped):
            doc.add_paragraph(re.sub(r"^\s*\d+\.\s+", "", stripped), style="List Number")
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            next_line = lines[index].strip()
            if (
                not next_line
                or next_line == "---"
                or re.match(r"^(#{1,3})\s+", next_line)
                or re.match(r"^\s*[-*]\s+", next_line)
                or re.match(r"^\s*\d+\.\s+", next_line)
                or ("|" in next_line and index + 1 < len(lines) and re.match(r"^\s*\|?(?:\s*:?-+:?\s*\|)+\s*$", lines[index + 1].strip()))
            ):
                break
            paragraph_lines.append(next_line)
            index += 1

        add_markdown_paragraph(doc, " ".join(paragraph_lines))

    doc.save(str(path))

def save_pdf(text: str, path: Path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", size=11)
    for line in text.split("\n"):
        try:
            pdf.multi_cell(0, 8, line)
        except Exception:
            pdf.multi_cell(0, 8, line.encode("latin-1", errors="replace").decode("latin-1"))
    pdf.output(str(path))

# ── Endpoints ─────────────────────────────────────────────────────────────────
@app.get("/", response_class=HTMLResponse)
async def root():
    index = Path("../frontend/index.html")
    if index.exists():
        return HTMLResponse(content=index.read_text(encoding="utf-8"))
    return HTMLResponse("<h1>SCI OCR API</h1><p>Frontend no encontrado.</p>")


@app.get("/manifest.webmanifest")
async def web_manifest():
    path = Path("../frontend/manifest.webmanifest")
    if not path.exists():
        raise HTTPException(404, "Manifest no encontrado")
    return FileResponse(path, media_type="application/manifest+json")


@app.get("/health")
async def health():
    return {"status": "ok", "max_upload_mb": MAX_UPLOAD_MB}

@app.get("/config")
async def get_config():
    vision = vision_llm.vision_status()
    return {
        "max_upload_mb": MAX_UPLOAD_MB,
        "gemini_enabled": bool(GEMINI_API_KEY),
        "api_key_required": api_key_required(),
        "gemini_model": GEMINI_MODEL if GEMINI_API_KEY else None,
        **vision,
        "supplementary_annotations": OCR_SUPPLEMENTARY_PASS,
        "tesseract_only": OCR_TESSERACT_ONLY,
        "ocr_use_layout": OCR_USE_LAYOUT,
        "fix_rotation": OCR_FIX_ROTATION,
        "vision_enabled": (
            not OCR_TESSERACT_ONLY
            and OCR_VISION_ENABLED
            and bool(vision_llm.configured_providers())
        ),
        "local_ocr_only": OCR_TESSERACT_ONLY or not OCR_VISION_ENABLED or not vision_llm.configured_providers(),
        "paddle_enabled": OCR_PADDLE_ENABLED,
        "paddle_available": paddle_ocr.paddle_ocr_available(),
        "paddle_layout": paddle_layout_active(),
        "paddle_lang": OCR_PADDLE_LANG,
        "local_engines": local_ocr_engines(),
        "heavy_async_mb": OCR_HEAVY_ASYNC_MB,
        "heavy_async_pages": OCR_HEAVY_ASYNC_PAGES,
        "max_concurrent_jobs": OCR_MAX_CONCURRENT_JOBS,
        "chunk_pages": OCR_CHUNK_PAGES,
        "ocr_modes": [
            {
                "id": "fast",
                "label": "Rápido",
                "dpi_pdf": ocr_dpi_for_mode("fast"),
                "detail": "Texto impreso principal",
            },
            {
                "id": "balanced",
                "label": "Equilibrado",
                "dpi_pdf": ocr_dpi_for_mode("balanced"),
                "detail": "Recomendado: facturas, recibos y rangos (~2–3× más rápido que máxima calidad)",
            },
            {
                "id": "quality",
                "label": "Máxima calidad",
                "dpi_pdf": ocr_dpi_for_mode("quality"),
                "detail": "Máxima precisión en sellos y manuscritos",
            },
        ],
    }


@app.get("/ocr/queue")
async def ocr_queue(x_api_key: Optional[str] = Header(None)):
    check_api_key(x_api_key)
    stats = reconcile_job_queue()
    jobs = list_active_jobs()
    with _queue_lock:
        running = _running_job_id
        pending_count = len(_pending_job_ids)
    running_jobs = [j for j in jobs if j.get("is_running")]
    queued_jobs = [j for j in jobs if j.get("in_queue")]
    return {
        "running_job_id": running,
        "pending_count": pending_count,
        "max_concurrent": OCR_MAX_CONCURRENT_JOBS,
        "jobs": jobs,
        "running_jobs": running_jobs,
        "queued_jobs": queued_jobs,
        "reconcile": stats,
    }


@app.post("/ocr/queue/cleanup")
async def ocr_queue_cleanup(
    force: bool = Query(False),
    x_api_key: Optional[str] = Header(None),
):
    """Cancela trabajos huérfanos o desbloquea la cola atascada."""
    check_api_key(x_api_key)
    stats = reconcile_job_queue()
    if force:
        stats["unstuck"] = unstick_queue_if_deadlocked(force=True)
    return {"ok": True, **stats}


@app.get("/ocr/history")
async def ocr_history(x_api_key: Optional[str] = Header(None)):
    check_api_key(x_api_key)
    return {"items": get_job_history()}


@app.post("/ocr/jobs/{job_id}/pause")
async def ocr_job_pause(
    job_id: str,
    yield_slot: bool = Form(True),
    x_api_key: Optional[str] = Header(None),
):
    check_api_key(x_api_key)
    return pause_ocr_job(job_id, yield_slot=bool(yield_slot))


@app.post("/ocr/jobs/{job_id}/prioritize")
async def ocr_job_prioritize(job_id: str, x_api_key: Optional[str] = Header(None)):
    check_api_key(x_api_key)
    return prioritize_ocr_job(job_id)


@app.post("/ocr/jobs/{job_id}/queue-move")
async def ocr_job_queue_move(
    job_id: str,
    direction: str = Form(...),
    x_api_key: Optional[str] = Header(None),
):
    check_api_key(x_api_key)
    direction = (direction or "").strip().lower()
    if direction not in ("up", "down"):
        raise HTTPException(400, "direction debe ser «up» o «down»")
    return move_ocr_job_in_queue(job_id, direction)


@app.post("/ocr/jobs/{job_id}/resume")
async def ocr_job_resume(job_id: str, x_api_key: Optional[str] = Header(None)):
    check_api_key(x_api_key)
    return resume_ocr_job(job_id)


@app.post("/ocr/jobs/{job_id}/cancel")
async def ocr_job_cancel(job_id: str, x_api_key: Optional[str] = Header(None)):
    check_api_key(x_api_key)
    return cancel_ocr_job(job_id)


@app.get("/ocr/status/{job_id}")
async def ocr_job_status(job_id: str, x_api_key: Optional[str] = Header(None)):
    check_api_key(x_api_key)
    reconcile_job_queue()
    state = read_job_state(job_id)
    if not state:
        raise HTTPException(404, "Trabajo no encontrado o expirado")
    if state.get("status") in ("queued", "paused") and _job_upload_missing(state):
        write_job_state(job_id, {
            **state,
            "status": "cancelled",
            "detail": "Trabajo huérfano. Vuelve a subir el archivo.",
        })
        return {
            "job_id": job_id,
            "status": "cancelled",
            "detail": "Trabajo huérfano. Vuelve a subir el archivo.",
            "async": True,
        }
    status = state.get("status", "unknown")
    if status == "completed" and state.get("result"):
        return state["result"]
    if status == "failed":
        raise HTTPException(500, state.get("error") or "Error al procesar el documento")
    if status == "cancelled":
        return {
            "job_id": job_id,
            "status": "cancelled",
            "detail": state.get("detail", "Cancelado"),
            "async": True,
        }
    summary = public_job_summary(state)
    summary["async"] = True
    return summary


@app.post("/ocr/source")
async def register_pdf_source(
    file: UploadFile = File(...),
    x_api_key: Optional[str] = Header(None),
):
    """Sube el PDF una vez; luego procesa por rangos sin volver a subir."""
    check_api_key(x_api_key)
    suffix = Path(file.filename or "").suffix.lower()
    if suffix != ".pdf":
        raise HTTPException(400, "Solo PDF en registro por rangos. Usa Procesar completo para imágenes.")

    source_id = str(uuid.uuid4())[:8]
    source_path = (SOURCES_DIR / f"{source_id}.pdf").resolve()
    with open(source_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    total_pages = count_pdf_pages(source_path)
    original_filename = file.filename or "documento.pdf"
    batch_id = batches.create_batch(
        JOBS_DIR,
        source_id=source_id,
        source_filename=original_filename,
        total_pages=total_pages,
    )
    return {
        "source_id": source_id,
        "batch_id": batch_id,
        "total_pages": total_pages,
        "source_filename": original_filename,
        "size_mb": round(source_path.stat().st_size / (1024 * 1024), 2),
        "suggested_ranges": batches.suggest_page_ranges(total_pages, chunk_size=50),
        "preview_url": f"/ocr/source/{source_id}/preview?page=1",
    }


@app.get("/ocr/source/{source_id}/preview")
async def source_pdf_preview(
    source_id: str,
    page: int = Query(1, ge=1),
    scale: float = Query(1.2, ge=0.5, le=3.0),
    x_api_key: Optional[str] = Header(None),
    key: Optional[str] = Query(None),
):
    """Miniatura JPEG de una página del PDF fuente (para la UI)."""
    effective_key = (x_api_key or key or "").strip()
    if api_key_required() and effective_key != API_KEY:
        raise HTTPException(403, "API Key inválida")
    source_path = (SOURCES_DIR / f"{source_id}.pdf").resolve()
    if not source_path.exists():
        raise HTTPException(404, "PDF no encontrado")
    try:
        with fitz.open(str(source_path)) as doc:
            if page > len(doc):
                raise HTTPException(400, f"La página {page} no existe")
            pg = doc.load_page(page - 1)
            matrix = fitz.Matrix(scale, scale).prerotate(pg.rotation)
            pix = pg.get_pixmap(matrix=matrix, alpha=False)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(500, f"No se pudo generar la vista previa: {exc}") from exc
    buffer = BytesIO()
    img.save(buffer, format="JPEG", quality=82, optimize=True)
    return Response(content=buffer.getvalue(), media_type="image/jpeg")


def _resolve_source_pdf(source_id: str) -> Path:
    source_path = (SOURCES_DIR / f"{source_id}.pdf").resolve()
    if not source_path.exists():
        raise HTTPException(404, "PDF no encontrado")
    return source_path


def _batch_for_source(source_id: str) -> tuple[Optional[dict[str, Any]], Optional[str]]:
    for path in batches.batches_dir(JOBS_DIR).glob("*.json"):
        data = batches.load_batch(JOBS_DIR, path.stem)
        if data and data.get("source_id") == source_id:
            return data, path.stem
    return None, None


def _source_download_url(source_id: str) -> str:
    return f"/download/source/{source_id}.pdf"


@app.get("/ocr/source/{source_id}/pages")
async def list_pdf_source_pages(
    source_id: str,
    include_thumbnails: bool = Query(False),
    thumbnail_limit: int = Query(24, ge=1, le=120),
    x_api_key: Optional[str] = Header(None),
):
    """Lista páginas del PDF fuente; opcionalmente URLs de miniatura."""
    check_api_key(x_api_key)
    source_path = _resolve_source_pdf(source_id)
    total_pages = count_pdf_pages(source_path)
    batch = None
    for path in batches.batches_dir(JOBS_DIR).glob("*.json"):
        data = batches.load_batch(JOBS_DIR, path.stem)
        if data and data.get("source_id") == source_id:
            batch = data
            break
    pages: list[dict[str, Any]] = []
    limit = min(thumbnail_limit, total_pages) if include_thumbnails else 0
    for page_num in range(1, total_pages + 1):
        entry: dict[str, Any] = {"page": page_num}
        if page_num <= limit:
            entry["preview_url"] = f"/ocr/source/{source_id}/preview?page={page_num}"
        pages.append(entry)
    return {
        "source_id": source_id,
        "source_filename": (batch or {}).get("source_filename", f"{source_id}.pdf"),
        "total_pages": total_pages,
        "size_mb": round(source_path.stat().st_size / (1024 * 1024), 2),
        "pages": pages,
        "suggested_ranges": batches.suggest_page_ranges(total_pages, chunk_size=50),
    }


@app.post("/ocr/source/{source_id}/extract")
async def extract_pdf_source_pages(
    source_id: str,
    keep_pages: Optional[str] = Form(None),
    remove_pages: Optional[str] = Form(None),
    replace_source: bool = Form(True),
    x_api_key: Optional[str] = Header(None),
):
    """Genera un PDF reducido conservando o excluyendo páginas antes del OCR."""
    check_api_key(x_api_key)
    source_path = _resolve_source_pdf(source_id)
    total_pages = count_pdf_pages(source_path)
    try:
        pages_to_keep = batches.resolve_pages_to_keep(
            total_pages,
            keep_pages=keep_pages,
            remove_pages=remove_pages,
        )
    except ValueError as exc:
        raise HTTPException(400, str(exc)) from exc

    if len(pages_to_keep) == total_pages:
        raise HTTPException(400, "No cambió ninguna página. Ajusta keep_pages o remove_pages.")

    batch = None
    batch_id: Optional[str] = None
    source_filename = f"{source_id}.pdf"
    batch, batch_id = _batch_for_source(source_id)
    if batch:
        source_filename = batch.get("source_filename") or source_filename

    temp_path = (UPLOAD_DIR / f"extract_{uuid.uuid4().hex[:8]}.pdf").resolve()
    try:
        batches.extract_pdf_pages(source_path, pages_to_keep, temp_path)
    except ValueError as exc:
        temp_path.unlink(missing_ok=True)
        raise HTTPException(400, str(exc)) from exc
    except Exception as exc:
        temp_path.unlink(missing_ok=True)
        raise HTTPException(500, f"No se pudo generar el PDF reducido: {exc}") from exc

    new_total = len(pages_to_keep)
    removed_count = total_pages - new_total
    out_source_id = source_id

    if replace_source:
        backup = source_path.with_suffix(".pdf.bak")
        try:
            if backup.exists():
                backup.unlink()
            source_path.replace(backup)
            temp_path.replace(source_path)
            backup.unlink(missing_ok=True)
        except OSError as exc:
            if backup.exists() and not source_path.exists():
                backup.replace(source_path)
            temp_path.unlink(missing_ok=True)
            raise HTTPException(500, f"No se pudo reemplazar el PDF: {exc}") from exc
        if batch_id and batch:
            batch["total_pages"] = new_total
            batch["page_trim"] = {
                "from_total": total_pages,
                "kept_pages": pages_to_keep,
                "removed_count": removed_count,
            }
            batches.save_batch(JOBS_DIR, batch_id, batch)
    else:
        out_source_id = str(uuid.uuid4())[:8]
        dest_path = (SOURCES_DIR / f"{out_source_id}.pdf").resolve()
        temp_path.replace(dest_path)
        batch_id = batches.create_batch(
            JOBS_DIR,
            source_id=out_source_id,
            source_filename=source_filename,
            total_pages=new_total,
        )

    return {
        "source_id": out_source_id,
        "batch_id": batch_id,
        "total_pages": new_total,
        "previous_total_pages": total_pages,
        "pages_kept": pages_to_keep,
        "pages_removed": removed_count,
        "replaced_source": replace_source,
        "source_filename": source_filename,
        "size_mb": round(
            (SOURCES_DIR / f"{out_source_id}.pdf").stat().st_size / (1024 * 1024),
            2,
        ),
        "suggested_ranges": batches.suggest_page_ranges(new_total, chunk_size=50),
        "preview_url": f"/ocr/source/{out_source_id}/preview?page=1",
        "download_url": _source_download_url(out_source_id),
        "detail": (
            f"PDF reducido: {new_total} página(s) conservadas "
            f"({removed_count} eliminada(s) del original de {total_pages})."
        ),
    }


@app.post("/ocr/source/{source_id}/compress")
async def compress_pdf_source(
    source_id: str,
    replace_source: bool = Form(True),
    quality: str = Form("high"),
    x_api_key: Optional[str] = Header(None),
):
    """Comprime el PDF fuente sin cambiar páginas. quality: medium | high."""
    check_api_key(x_api_key)
    source_path = _resolve_source_pdf(source_id)
    previous_size_mb = round(source_path.stat().st_size / (1024 * 1024), 2)
    q = (quality or "high").strip().lower()
    if q not in ("medium", "high"):
        raise HTTPException(400, "quality debe ser «medium» o «high»")

    batch, batch_id = _batch_for_source(source_id)
    source_filename = (batch or {}).get("source_filename") or f"{source_id}.pdf"
    total_pages = count_pdf_pages(source_path)

    temp_path = (UPLOAD_DIR / f"compress_{uuid.uuid4().hex[:8]}.pdf").resolve()
    try:
        batches.save_pdf_compressed(source_path, temp_path, quality=q)
    except Exception as exc:
        temp_path.unlink(missing_ok=True)
        raise HTTPException(500, f"No se pudo comprimir el PDF: {exc}") from exc

    out_source_id = source_id
    if replace_source:
        backup = source_path.with_suffix(".pdf.bak")
        try:
            if backup.exists():
                backup.unlink()
            source_path.replace(backup)
            temp_path.replace(source_path)
            backup.unlink(missing_ok=True)
        except OSError as exc:
            if backup.exists() and not source_path.exists():
                backup.replace(source_path)
            temp_path.unlink(missing_ok=True)
            raise HTTPException(500, f"No se pudo reemplazar el PDF: {exc}") from exc
        if batch_id and batch:
            batch["pdf_compress"] = {"previous_size_mb": previous_size_mb}
            batches.save_batch(JOBS_DIR, batch_id, batch)
    else:
        out_source_id = str(uuid.uuid4())[:8]
        dest_path = (SOURCES_DIR / f"{out_source_id}.pdf").resolve()
        temp_path.replace(dest_path)
        batch_id = batches.create_batch(
            JOBS_DIR,
            source_id=out_source_id,
            source_filename=source_filename,
            total_pages=total_pages,
        )

    new_size_mb = round(
        (SOURCES_DIR / f"{out_source_id}.pdf").stat().st_size / (1024 * 1024),
        2,
    )
    saved_pct = (
        round((1 - new_size_mb / previous_size_mb) * 100, 1)
        if previous_size_mb > 0
        else 0.0
    )
    return {
        "source_id": out_source_id,
        "batch_id": batch_id,
        "total_pages": total_pages,
        "replaced_source": replace_source,
        "source_filename": source_filename,
        "size_mb": new_size_mb,
        "previous_size_mb": previous_size_mb,
        "saved_pct": saved_pct,
        "download_url": _source_download_url(out_source_id),
        "preview_url": f"/ocr/source/{out_source_id}/preview?page=1",
        "quality": q,
        "detail": (
            f"PDF comprimido ({q}): {previous_size_mb} MB → {new_size_mb} MB "
            f"({saved_pct}% menos)."
        ),
    }


@app.post("/pdf/merge")
async def merge_pdf_documents(
    base_source_id: Optional[str] = Form(None),
    extra_source_ids: str = Form(""),
    files: List[UploadFile] = File(default=[]),
    x_api_key: Optional[str] = Header(None),
):
    """Une el PDF base (opcional) con otros PDF del servidor o archivos subidos."""
    check_api_key(x_api_key)
    paths: list[Path] = []
    labels: list[str] = []
    temps: list[Path] = []

    if base_source_id and base_source_id.strip():
        base_path = _resolve_source_pdf(base_source_id.strip())
        paths.append(base_path)
        batch, _ = _batch_for_source(base_source_id.strip())
        labels.append((batch or {}).get("source_filename") or base_source_id)

    for sid in [s.strip() for s in extra_source_ids.split(",") if s.strip()]:
        if base_source_id and sid == base_source_id.strip():
            continue
        paths.append(_resolve_source_pdf(sid))
        batch, _ = _batch_for_source(sid)
        labels.append((batch or {}).get("source_filename") or sid)

    for f in files:
        if not f.filename or not f.filename.lower().endswith(".pdf"):
            continue
        tmp = (UPLOAD_DIR / f"merge_in_{uuid.uuid4().hex[:8]}.pdf").resolve()
        with open(tmp, "wb") as out:
            shutil.copyfileobj(f.file, out)
        paths.append(tmp)
        temps.append(tmp)
        labels.append(f.filename)

    if len(paths) < 2:
        for t in temps:
            t.unlink(missing_ok=True)
        raise HTTPException(400, "Indica al menos 2 PDF (sube varios o usa PDF actual + otro archivo).")

    out_source_id = str(uuid.uuid4())[:8]
    dest_path = (SOURCES_DIR / f"{out_source_id}.pdf").resolve()
    try:
        total_pages = batches.merge_pdf_files(paths, dest_path)
    except ValueError as exc:
        dest_path.unlink(missing_ok=True)
        raise HTTPException(400, str(exc)) from exc
    finally:
        for t in temps:
            t.unlink(missing_ok=True)

    stem = safe_download_stem(labels[0] if labels else "documento")
    out_name = f"{stem}_unido.pdf" if len(labels) == 1 else f"pdf_unido_{len(paths)}_archivos.pdf"
    batch_id = batches.create_batch(
        JOBS_DIR,
        source_id=out_source_id,
        source_filename=out_name,
        total_pages=total_pages,
    )
    size_mb = round(dest_path.stat().st_size / (1024 * 1024), 2)
    return {
        "source_id": out_source_id,
        "batch_id": batch_id,
        "total_pages": total_pages,
        "parts_merged": len(paths),
        "source_filename": out_name,
        "size_mb": size_mb,
        "download_url": _source_download_url(out_source_id),
        "preview_url": f"/ocr/source/{out_source_id}/preview?page=1",
        "detail": f"Se unieron {len(paths)} PDF en uno ({total_pages} páginas, {size_mb} MB).",
    }


@app.get("/download/source/{source_id}.pdf")
async def download_source_pdf(
    source_id: str,
    x_api_key: Optional[str] = Header(None),
    key: Optional[str] = Query(None),
):
    effective_key = (x_api_key or key or "").strip()
    if api_key_required() and effective_key != API_KEY:
        raise HTTPException(403, "API Key inválida")
    source_path = _resolve_source_pdf(source_id)
    batch, _ = _batch_for_source(source_id)
    source_filename = (batch or {}).get("source_filename") or f"{source_id}.pdf"
    stem = safe_download_stem(source_filename)
    return FileResponse(
        source_path,
        media_type="application/pdf",
        filename=f"{stem}.pdf",
    )


@app.post("/ocr/range")
async def process_pdf_range(
    source_id: str = Query(...),
    batch_id: str = Query(...),
    page_from: int = Query(..., ge=1),
    page_to: int = Query(..., ge=1),
    lang: str = Query("spa"),
    mode: str = Query("balanced"),
    vision_provider: str = Query("auto"),
    x_api_key: Optional[str] = Header(None),
):
    check_api_key(x_api_key)
    source_path = (SOURCES_DIR / f"{source_id}.pdf").resolve()
    if not source_path.exists():
        raise HTTPException(404, "PDF fuente no encontrado. Vuelve a subir el archivo.")

    batch = batches.load_batch(JOBS_DIR, batch_id)
    if not batch or batch.get("source_id") != source_id:
        raise HTTPException(400, "Lote o fuente no coinciden.")

    total_pages = int(batch.get("total_pages") or count_pdf_pages(source_path))
    pf, pt = batches.normalize_page_range(page_from, page_to, total_pages)
    range_pages = pt - pf + 1

    for part in batch.get("parts") or []:
        if part.get("status") in ("queued", "processing", "paused") and part.get("page_from") == pf and part.get("page_to") == pt:
            raise HTTPException(400, f"Ya hay un trabajo para las páginas {pf}-{pt}.")

    job_id = str(uuid.uuid4())[:8]
    original_filename = batch.get("source_filename", "documento.pdf")
    effective_mode = normalize_ocr_mode(mode)

    batch["processing_mode"] = "ranges"
    batches.save_batch(JOBS_DIR, batch_id, batch)

    reconcile_job_queue()
    queue_position = enqueue_ocr_job(
        job_id,
        source_path,
        lang=lang,
        mode=effective_mode,
        original_filename=original_filename,
        pages_total=range_pages,
        page_from=pf,
        page_to=pt,
        batch_id=batch_id,
        keep_source=True,
        range_slice=True,
        processing_mode="ranges",
        vision_provider=vision_provider,
    )
    return JSONResponse(
        status_code=202,
        content={
            "job_id": job_id,
            "batch_id": batch_id,
            "async": True,
            "status": "queued",
            "processing_mode": "ranges",
            "page_from": pf,
            "page_to": pt,
            "pages_total": range_pages,
            "queue_position": queue_position,
            "status_url": f"/ocr/status/{job_id}",
            "batch_url": f"/ocr/batch/{batch_id}",
            "detail": f"Procesando páginas {pf}-{pt} ({range_pages} páginas).",
        },
    )


@app.post("/ocr/full")
async def process_pdf_full(
    source_id: str = Query(...),
    batch_id: str = Query(...),
    lang: str = Query("spa"),
    mode: str = Query("balanced"),
    vision_provider: str = Query("auto"),
    x_api_key: Optional[str] = Header(None),
):
    """Procesa todas las páginas del PDF en un solo trabajo OCR."""
    check_api_key(x_api_key)
    source_path = (SOURCES_DIR / f"{source_id}.pdf").resolve()
    if not source_path.exists():
        raise HTTPException(404, "PDF fuente no encontrado. Vuelve a subir el archivo.")

    batch = batches.load_batch(JOBS_DIR, batch_id)
    if not batch or batch.get("source_id") != source_id:
        raise HTTPException(400, "Lote o fuente no coinciden.")

    total_pages = int(batch.get("total_pages") or count_pdf_pages(source_path))
    pf, pt = 1, total_pages

    for part in batch.get("parts") or []:
        if part.get("status") in ("queued", "processing", "paused") and int(part.get("page_from") or 0) == pf and int(part.get("page_to") or 0) == pt:
            raise HTTPException(400, "Ya hay un trabajo en curso para el PDF completo.")

    job_id = str(uuid.uuid4())[:8]
    original_filename = batch.get("source_filename", "documento.pdf")
    effective_mode = normalize_ocr_mode(mode)

    batch["processing_mode"] = "full"
    batches.save_batch(JOBS_DIR, batch_id, batch)

    reconcile_job_queue()
    queue_position = enqueue_ocr_job(
        job_id,
        source_path,
        lang=lang,
        mode=effective_mode,
        original_filename=original_filename,
        pages_total=total_pages,
        page_from=pf,
        page_to=pt,
        batch_id=batch_id,
        keep_source=True,
        vision_provider=vision_provider,
        range_slice=False,
        processing_mode="full",
    )
    return JSONResponse(
        status_code=202,
        content={
            "job_id": job_id,
            "batch_id": batch_id,
            "async": True,
            "status": "queued",
            "processing_mode": "full",
            "page_from": pf,
            "page_to": pt,
            "pages_total": total_pages,
            "queue_position": queue_position,
            "status_url": f"/ocr/status/{job_id}",
            "batch_url": f"/ocr/batch/{batch_id}",
            "full_download_url": f"/download/{job_id}/md",
            "detail": f"Procesando PDF completo ({total_pages} páginas).",
        },
    )


@app.get("/ocr/batch/{batch_id}")
async def get_batch_status(batch_id: str, x_api_key: Optional[str] = Header(None)):
    check_api_key(x_api_key)
    batch = batches.sync_batch_from_jobs(JOBS_DIR, batch_id, read_job_state)
    if not batch:
        raise HTTPException(404, "Lote no encontrado")
    return batch


@app.get("/download/batch/{batch_id}/full.md")
async def download_batch_full_md(
    batch_id: str,
    x_api_key: Optional[str] = Header(None),
    key: Optional[str] = Query(None),
):
    """MD del PDF completo cuando se procesó en modo «full» (una sola parte)."""
    effective_key = (x_api_key or key or "").strip()
    if api_key_required() and effective_key != API_KEY:
        raise HTTPException(403, "API Key inválida")
    batch = batches.sync_batch_from_jobs(JOBS_DIR, batch_id, read_job_state)
    if not batch:
        raise HTTPException(404, "Lote no encontrado")
    job_id = batch.get("full_document_job_id")
    if not job_id:
        raise HTTPException(400, "Aún no hay un PDF completo procesado en este lote.")
    md_path = ensure_full_job_markdown(job_id)
    if not md_path.exists():
        raise HTTPException(404, "Archivo MD no encontrado")
    stem = safe_download_stem(batch.get("source_filename", "documento"))
    return FileResponse(
        md_path,
        media_type="text/markdown",
        filename=f"{stem}.md",
    )


@app.get("/download/batch/{batch_id}/merged.md")
async def download_batch_merged(
    batch_id: str,
    x_api_key: Optional[str] = Header(None),
    key: Optional[str] = Query(None),
):
    effective_key = (x_api_key or key or "").strip()
    if api_key_required() and effective_key != API_KEY:
        raise HTTPException(403, "API Key inválida")
    batch = batches.sync_batch_from_jobs(JOBS_DIR, batch_id, read_job_state) or {}
    for part in batch.get("parts") or []:
        job_id = part.get("job_id")
        if part.get("status") == "completed" and job_id:
            ensure_full_job_markdown(job_id)
    try:
        merged = batches.merge_batch_markdown(JOBS_DIR, batch_id, OUTPUT_DIR, read_job_state)
    except FileNotFoundError:
        raise HTTPException(404, "Lote no encontrado")
    except ValueError as exc:
        raise HTTPException(400, str(exc))
    out_path = OUTPUT_DIR / f"batch_{batch_id}_merged.md"
    out_path.write_text(merged, encoding="utf-8")
    batch = batches.load_batch(JOBS_DIR, batch_id) or {}
    stem = safe_download_stem(batch.get("source_filename", "documento"))
    return FileResponse(
        out_path,
        media_type="text/markdown",
        filename=f"{stem}.md",
    )


@app.get("/download/batch/{batch_id}/zip")
async def download_batch_zip(
    batch_id: str,
    x_api_key: Optional[str] = Header(None),
    key: Optional[str] = Query(None),
):
    effective_key = (x_api_key or key or "").strip()
    if api_key_required() and effective_key != API_KEY:
        raise HTTPException(403, "API Key inválida")
    zip_path = OUTPUT_DIR / f"batch_{batch_id}.zip"
    try:
        batches.build_batch_zip(JOBS_DIR, batch_id, OUTPUT_DIR, zip_path, read_job_state)
    except FileNotFoundError:
        raise HTTPException(404, "Lote no encontrado")
    except ValueError as exc:
        raise HTTPException(400, str(exc))
    batch = batches.load_batch(JOBS_DIR, batch_id) or {}
    stem = safe_download_stem(batch.get("source_filename", "documento"))
    return FileResponse(
        zip_path,
        media_type="application/zip",
        filename=f"{stem}_partes.zip",
    )


def _read_md_from_job(job_id: str) -> str:
    path = OUTPUT_DIR / f"{job_id}.md"
    if not path.exists():
        raise HTTPException(404, f"No hay MD para el trabajo {job_id}")
    return path.read_text(encoding="utf-8")


def _cancel_jobs_for_source(source_id: str) -> int:
    cancelled = 0
    for path in JOBS_DIR.glob("*.json"):
        if path.name == "history.json":
            continue
        state = read_job_state(path.stem)
        if not state or state.get("status") not in ("queued", "paused", "processing"):
            continue
        upload = (state.get("upload_path") or "").replace("\\", "/")
        if f"/sources/{source_id}.pdf" in upload or upload.endswith(f"{source_id}.pdf"):
            try:
                cancel_ocr_job(path.stem)
                cancelled += 1
            except HTTPException:
                pass
    return cancelled


@app.delete("/ocr/source/{source_id}")
async def delete_pdf_source(source_id: str, x_api_key: Optional[str] = Header(None)):
    """Quita el PDF del servidor y cancela trabajos pendientes de esa fuente."""
    check_api_key(x_api_key)
    source_path = (SOURCES_DIR / f"{source_id}.pdf").resolve()
    n = _cancel_jobs_for_source(source_id)
    if source_path.exists():
        source_path.unlink(missing_ok=True)
    elif n == 0:
        raise HTTPException(404, "PDF no encontrado en el servidor")
    return {"ok": True, "cancelled_jobs": n}


@app.get("/ocr/completed-mds")
async def list_completed_mds(x_api_key: Optional[str] = Header(None)):
    """Lista trabajos completados con archivo .md disponible."""
    check_api_key(x_api_key)
    items: list[dict[str, Any]] = []
    seen: set[str] = set()
    for entry in get_job_history():
        if entry.get("status") != "completed":
            continue
        job_id = entry.get("job_id")
        if not job_id or job_id in seen:
            continue
        md_path = ensure_full_job_markdown(job_id)
        if not md_path.exists():
            continue
        seen.add(job_id)
        items.append({
            "job_id": job_id,
            "label": entry.get("source_filename") or job_id,
            "size_kb": round(md_path.stat().st_size / 1024, 1),
        })
    for md_path in sorted(OUTPUT_DIR.glob("*.md"), key=lambda p: p.stat().st_mtime, reverse=True):
        job_id = md_path.stem
        if job_id.startswith("batch_") or job_id.startswith("merge_") or job_id in seen:
            continue
        if read_job_state(job_id) and read_job_state(job_id).get("status") == "completed":
            continue
        if not any(i["job_id"] == job_id for i in items):
            items.append({
                "job_id": job_id,
                "label": md_path.stem,
                "size_kb": round(md_path.stat().st_size / 1024, 1),
            })
    return {"items": items[:30]}


@app.post("/md/merge")
async def merge_markdown_files(
    job_ids: str = Form(""),
    files: List[UploadFile] = File(default=[]),
    x_api_key: Optional[str] = Header(None),
):
    check_api_key(x_api_key)
    chunks: list[str] = []
    labels: list[str] = []
    for jid in [j.strip() for j in job_ids.split(",") if j.strip()]:
        chunks.append(_read_md_from_job(jid))
        labels.append(jid)
    for f in files:
        if not f.filename:
            continue
        raw = await f.read()
        chunks.append(raw.decode("utf-8", errors="replace"))
        labels.append(f.filename)
    try:
        merged = sanitize_document_markdown(md_tools.merge_markdown_chunks(chunks))
    except ValueError as exc:
        raise HTTPException(400, str(exc))
    merge_id = str(uuid.uuid4())[:8]
    out_path = OUTPUT_DIR / f"merge_{merge_id}.md"
    out_path.write_text(merged, encoding="utf-8")
    stats = md_tools.size_stats(merged)
    return {
        "merge_id": merge_id,
        "markdown": merged,
        "parts": len(chunks),
        "labels": labels,
        "download_url": f"/download/merge/{merge_id}.md",
        **stats,
    }


@app.post("/md/clean")
async def clean_markdown(
    file: Optional[UploadFile] = File(None),
    job_id: Optional[str] = Form(None),
    markdown: Optional[str] = Form(None),
    x_api_key: Optional[str] = Header(None),
):
    """Limpia un .md existente (sellos, CFDI, ruido OCR) sin volver a procesar el PDF."""
    check_api_key(x_api_key)
    if job_id:
        text = _read_md_from_job(job_id)
    elif file and file.filename:
        text = (await file.read()).decode("utf-8", errors="replace")
    elif markdown:
        text = markdown
    else:
        raise HTTPException(400, "Sube un .md, indica job_id o envía markdown")
    cleaned = sanitize_document_markdown(text)
    clean_id = str(uuid.uuid4())[:8]
    out_path = OUTPUT_DIR / f"clean_{clean_id}.md"
    out_path.write_text(cleaned, encoding="utf-8")
    before = md_tools.size_stats(text)
    after = md_tools.size_stats(cleaned)
    return {
        "clean_id": clean_id,
        "markdown": cleaned,
        "download_url": f"/download/clean/{clean_id}.md",
        "before": before,
        "after": after,
        "lines_removed": len(text.splitlines()) - len(cleaned.splitlines()),
    }


@app.post("/md/analyze")
async def analyze_markdown(
    file: Optional[UploadFile] = File(None),
    job_id: Optional[str] = Form(None),
    markdown: Optional[str] = Form(None),
    x_api_key: Optional[str] = Header(None),
):
    check_api_key(x_api_key)
    if job_id:
        text = _read_md_from_job(job_id)
    elif file and file.filename:
        text = (await file.read()).decode("utf-8", errors="replace")
    elif markdown:
        text = markdown
    else:
        raise HTTPException(400, "Sube un .md, indica job_id o envía markdown")
    sections = md_tools.split_sections(text)
    stats = md_tools.size_stats(text)
    return {"sections": sections, **stats}


@app.post("/md/process")
async def process_markdown(
    markdown: str = Form(...),
    keep_section_ids: str = Form(""),
    compress: bool = Form(False),
    x_api_key: Optional[str] = Header(None),
):
    check_api_key(x_api_key)
    try:
        if keep_section_ids.strip():
            keep = [s.strip() for s in keep_section_ids.split(",") if s.strip()]
            result = md_tools.apply_section_edits(markdown, keep)
        else:
            result = markdown
        if compress:
            result = md_tools.compress_markdown(result)
    except ValueError as exc:
        raise HTTPException(400, str(exc))
    edit_id = str(uuid.uuid4())[:8]
    out_path = OUTPUT_DIR / f"edit_{edit_id}.md"
    out_path.write_text(result, encoding="utf-8")
    before = md_tools.size_stats(markdown)
    after = md_tools.size_stats(result)
    return {
        "edit_id": edit_id,
        "markdown": result,
        "download_url": f"/download/edit/{edit_id}.md",
        "before_chars": before["chars"],
        "after_chars": after["chars"],
        "saved_pct": round((1 - after["chars"] / before["chars"]) * 100, 1) if before["chars"] else 0,
    }


@app.get("/download/merge/{merge_id}.md")
async def download_merge_md(
    merge_id: str,
    x_api_key: Optional[str] = Header(None),
    key: Optional[str] = Query(None),
):
    effective_key = (x_api_key or key or "").strip()
    if api_key_required() and effective_key != API_KEY:
        raise HTTPException(403, "API Key inválida")
    path = OUTPUT_DIR / f"merge_{merge_id}.md"
    if not path.exists():
        raise HTTPException(404, "Archivo no encontrado")
    return FileResponse(path, media_type="text/markdown", filename=f"documento_unido_{merge_id}.md")


@app.get("/download/edit/{edit_id}.md")
async def download_edit_md(
    edit_id: str,
    x_api_key: Optional[str] = Header(None),
    key: Optional[str] = Query(None),
):
    effective_key = (x_api_key or key or "").strip()
    if api_key_required() and effective_key != API_KEY:
        raise HTTPException(403, "API Key inválida")
    path = OUTPUT_DIR / f"edit_{edit_id}.md"
    if not path.exists():
        raise HTTPException(404, "Archivo no encontrado")
    return FileResponse(path, media_type="text/markdown", filename=f"documento_editado_{edit_id}.md")


@app.get("/download/clean/{clean_id}.md")
async def download_clean_md(
    clean_id: str,
    x_api_key: Optional[str] = Header(None),
    key: Optional[str] = Query(None),
):
    effective_key = (x_api_key or key or "").strip()
    if api_key_required() and effective_key != API_KEY:
        raise HTTPException(403, "API Key inválida")
    path = OUTPUT_DIR / f"clean_{clean_id}.md"
    if not path.exists():
        raise HTTPException(404, "Archivo no encontrado")
    return FileResponse(path, media_type="text/markdown", filename=f"documento_limpio_{clean_id}.md")


@app.post("/ocr")
async def ocr_endpoint(
    file: UploadFile = File(...),
    lang: str = Query("spa"),
    mode: str = Query("balanced"),
    vision_provider: str = Query("auto"),
    async_mode: Optional[bool] = Query(None, alias="async"),
    x_api_key: Optional[str] = Header(None),
):
    check_api_key(x_api_key)

    allowed = {".pdf", ".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp"}
    suffix = Path(file.filename).suffix.lower()
    if suffix not in allowed:
        raise HTTPException(400, f"Tipo no permitido: {suffix}")

    job_id = str(uuid.uuid4())[:8]
    upload_path = UPLOAD_DIR / f"{job_id}{suffix}"

    with open(upload_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    file_size = upload_path.stat().st_size
    original_filename = file.filename or "documento"
    effective_mode = normalize_ocr_mode(mode)
    force_async = async_mode is True
    force_sync = async_mode is False

    if should_process_async(upload_path, file_size, force_async=force_async, force_sync=force_sync):
        reconcile_job_queue()
        pages_total = count_pdf_pages(upload_path) if suffix == ".pdf" else 1
        queue_position = enqueue_ocr_job(
            job_id,
            upload_path,
            lang=lang,
            mode=effective_mode,
            original_filename=original_filename,
            pages_total=pages_total,
            vision_provider=vision_provider,
        )
        return JSONResponse(
            status_code=202,
            content={
                "job_id": job_id,
                "async": True,
                "status": "queued",
                "status_url": f"/ocr/status/{job_id}",
                "queue_url": "/ocr/queue",
                "source_filename": original_filename,
                "pages_total": pages_total,
                "queue_position": queue_position,
                "detail": (
                    f"En cola (posición {queue_position}). "
                    "Abre «Cola de producción» para ver el avance, pausar o cancelar."
                ),
            },
        )

    try:
        result = process_file(
            upload_path,
            lang=lang,
            mode=effective_mode,
            vision_provider=vision_provider,
        )
    except Exception as e:
        raise HTTPException(500, f"Error al procesar: {str(e)}")
    finally:
        upload_path.unlink(missing_ok=True)

    return build_ocr_response(
        job_id,
        result,
        original_filename=original_filename,
        effective_mode=effective_mode,
    )

@app.get("/download/{job_id}/{fmt}")
async def download(
    job_id: str,
    fmt: str,
    x_api_key: Optional[str] = Header(None),
    key: Optional[str] = Query(None),
):
    effective_key = (x_api_key or key or "").strip()
    if api_key_required() and effective_key != API_KEY:
        raise HTTPException(403, "API Key inválida")

    ext_map   = {"txt": ".txt", "md": ".md", "pdf": ".pdf", "docx": ".docx"}
    media_map = {
        "txt": "text/plain",
        "md": "text/markdown",
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }

    if fmt not in ext_map:
        raise HTTPException(400, "Formato inválido")

    if fmt == "md":
        path = ensure_full_job_markdown(job_id)
    else:
        path = OUTPUT_DIR / f"{job_id}{ext_map[fmt]}"
    if not path.exists():
        raise HTTPException(404, "Archivo no encontrado")

    download_name = f"ocr_{job_id}{ext_map[fmt]}"
    if fmt == "md":
        meta = load_job_meta(job_id)
        stem = meta.get("stem") or safe_download_stem(meta.get("filename", ""))
        download_name = f"{stem}.md"

    return FileResponse(path, media_type=media_map[fmt], filename=download_name)

if __name__ == "__main__":
    host = os.getenv("OCR_HOST", "0.0.0.0")
    port = int(os.getenv("OCR_PORT", "8000"))
    reload = os.getenv("OCR_RELOAD", "0").strip().lower() in ("1", "true", "yes")
    uvicorn.run("main:app", host=host, port=port, reload=reload)